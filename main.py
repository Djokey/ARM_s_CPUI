import copy
import threading
import sys
import os
import win32api
import win32print
import datetime
import docx
import res
import imaplib
import email
import py_config as pc

from time import sleep
from subprocess import check_output, call
from PyQt5.QtWidgets import QMessageBox
from PyQt5 import QtCore, QtGui, QtWidgets
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

# My UI includes
from ui import *
from headers_ui import *
from programs_ui import *
from teachers_ui import *
from groups_ui import *
from subjects_ui import *
from students_ui import *
from enrollment_ui import *
from outlay_ui import *
from timetable_edit_ui import *
from outlay_printer_ui import *
from decree_enrollment import *
from note_passes import *
from note_passwords import *
from note_studs_list import *
from contract_ui import *
from docx_creator_ui import *
from settings_ui import *
# My DataBase controller
from arm_db import *


# Class for main window application
class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self._translate = QtCore.QCoreApplication.translate

        self.ui.widget_headers = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_headers)
        self.head_ui = Ui_Headers()
        self.head_ui.setupUi(self.ui.widget_headers)

        self.ui.widget_teachers = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_teachers)
        self.teach_ui = Ui_Teachers()
        self.teach_ui.setupUi(self.ui.widget_teachers)

        self.ui.widget_programs = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_programs)
        self.prog_ui = Ui_Programs()
        self.prog_ui.setupUi(self.ui.widget_programs)

        self.ui.widget_groups = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_groups)
        self.groups_ui = Ui_Groups()
        self.groups_ui.setupUi(self.ui.widget_groups)

        self.ui.widget_subjects = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_subjects)
        self.sub_ui = Ui_Subjects()
        self.sub_ui.setupUi(self.ui.widget_subjects)

        self.ui.widget_students = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_students)
        self.stud_ui = Ui_Students()
        self.stud_ui.setupUi(self.ui.widget_students)

        self.ui.widget_outlay = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_outlay)
        self.outlay_ui = Ui_Outlay()
        self.outlay_ui.setupUi(self.ui.widget_outlay)

        self.ui.widget_enrollment = QtWidgets.QWidget(self.ui.widget_roster_editors)
        self.ui.hL_widget_roster_editors.addWidget(self.ui.widget_enrollment)
        self.enr_ui = Ui_Enrollment()
        self.enr_ui.setupUi(self.ui.widget_enrollment)

        self.ttable = QtWidgets.QDialog(self)
        self.ttable_ui = Ui_TTableEditor()
        self.ttable_ui.setupUi(self.ttable)
        self.ttable.setWindowTitle('Редактор расписания')

        self.ttable_selected_sub = ''
        self.ttable_list = []

        self.outlay_printer = QtWidgets.QDialog(self)
        self.outlay_print_ui = Ui_OutlayPrinter()
        self.outlay_print_ui.setupUi(self.outlay_printer)
        self.outlay_printer.setWindowTitle('Редактор сметы')

        self.docx_creator = QtWidgets.QDialog(self)
        self.docx_ui = Ui_DocxCreator()
        self.docx_ui.setupUi(self.docx_creator)
        self.docx_creator.setWindowTitle('Редактор документов')

        self.docx_ui.widget_enrollment = QtWidgets.QWidget()
        self.docx_ui.not_main.addWidget(self.docx_ui.widget_enrollment)
        self.decree_enr_ui = Ui_DecreeEnrollment()
        self.decree_enr_ui.setupUi(self.docx_ui.widget_enrollment)

        self.docx_ui.widget_notepasses = QtWidgets.QWidget()
        self.docx_ui.not_main.addWidget(self.docx_ui.widget_notepasses)
        self.note_passes_ui = Ui_NotePasses()
        self.note_passes_ui.setupUi(self.docx_ui.widget_notepasses)

        self.docx_ui.widget_notepasswords = QtWidgets.QWidget()
        self.docx_ui.not_main.addWidget(self.docx_ui.widget_notepasswords)
        self.note_passwords_ui = Ui_NotePasswords()
        self.note_passwords_ui.setupUi(self.docx_ui.widget_notepasswords)

        self.docx_ui.widget_notelist = QtWidgets.QWidget()
        self.docx_ui.not_main.addWidget(self.docx_ui.widget_notelist)
        self.note_list_ui = Ui_NoteList()
        self.note_list_ui.setupUi(self.docx_ui.widget_notelist)

        self.docx_ui.widget_contract = QtWidgets.QWidget()
        self.docx_ui.not_main.addWidget(self.docx_ui.widget_contract)
        self.contract_ui = Ui_Contract()
        self.contract_ui.setupUi(self.docx_ui.widget_contract)

        self.disk_dir = os.getenv("SystemDrive")
        self.user = os.environ.get("USERNAME")
        self.dir = os.path.abspath(os.curdir)
        self.clear_for_start()
        self.setup_buttons_funcs()
        self.load_for_start()
        self.resize(1000, 885)

        # MAIN END

    # Func for edit database table Headers
    def headers_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_headers.show()
        self.load_db_headers()

    # Func for edit database table Programs
    def programs_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_programs.show()
        self.load_db_programs()

    # Func for edit database table Teachers
    def teachers_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_teachers.show()
        self.load_db_teachers()

    # Func for edit database table Groups
    def groups_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_groups.show()
        self.load_db_groups()

    # Func for edit database table Subjects
    def subjects_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_subjects.show()
        self.load_db_subjects()

    # Func for edit database table Students
    def students_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_students.show()
        self.load_db_students()

    # Func for edit database table Enrollment
    def enrollment_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_enrollment.show()
        self.load_db_enrollment()

    # Func for create decree enrollment docx
    def decree_enr_win(self):
        self.docx_ui.widget_main.hide()
        self.docx_ui.widget_enrollment.show()
        self.load_db_decree_enr()

    # Func for create note passes docx
    def note_passes_win(self):
        self.docx_ui.widget_main.hide()
        self.docx_ui.widget_notepasses.show()
        self.load_db_note_passes()

    # Func for create note passwords docx
    def note_passwords_win(self):
        self.docx_ui.widget_main.hide()
        self.docx_ui.widget_notepasswords.show()
        self.load_db_note_passwords()

    # Func for create note list studs docx
    def note_list_win(self):
        self.docx_ui.widget_main.hide()
        self.docx_ui.widget_notelist.show()
        self.load_db_note_list()

    # Func for create contract docxs
    def contract_win(self):
        self.docx_ui.widget_main.hide()
        self.docx_ui.widget_contract.show()
        self.load_db_contract()

    # Func for edit database table Outlay
    def outlay_win(self):
        self.ui.widget_roster.hide()
        self.ui.widget_outlay.show()
        self.load_db_outlay()

    # Func for setup all buttons
    def setup_buttons_funcs(self):
        def headers_back():
            self.ui.widget_headers.hide()
            self.ui.widget_roster.show()

        def programs_back():
            self.ui.widget_programs.hide()
            self.ui.widget_roster.show()

        def teachers_back():
            self.ui.widget_teachers.hide()
            self.ui.widget_roster.show()

        def groups_back():
            self.ui.widget_groups.hide()
            self.ui.widget_roster.show()

        def subjects_back():
            self.ui.widget_subjects.hide()
            self.ui.widget_roster.show()

        def students_back():
            self.ui.widget_students.hide()
            self.ui.widget_roster.show()

        def enrollment_back():
            self.ui.widget_enrollment.hide()
            self.ui.widget_roster.show()

        def decree_enr_back():
            self.docx_ui.widget_enrollment.hide()
            self.docx_ui.widget_main.show()

        def note_passes_back():
            self.docx_ui.widget_notepasses.hide()
            self.docx_ui.widget_main.show()

        def note_passwords_back():
            self.docx_ui.widget_notepasswords.hide()
            self.docx_ui.widget_main.show()

        def note_list_back():
            self.docx_ui.widget_notelist.hide()
            self.docx_ui.widget_main.show()

        def contract_back():
            self.docx_ui.widget_contract.hide()
            self.docx_ui.widget_main.show()

        def outlay_back():
            self.ui.widget_outlay.hide()
            self.ui.widget_roster.show()

        def print_selected_doc():
            current_tab = self.ui.tabWidget_docx.currentWidget()
            current_list = current_tab.findChild(
                QtWidgets.QWidget,
                'sAWContent_' + current_tab.objectName().split('_')[-1]).children()
            _set_doc_warning = 1
            if 'decree' in current_tab.objectName():
                folder = 'Приказы'
            elif 'notes' in current_tab.objectName():
                folder = 'Записки'
            elif 'ttable' in current_tab.objectName():
                folder = 'Расписания'
            elif 'outlay' in current_tab.objectName():
                folder = 'Сметы'
            for i in current_list:
                if i.objectName().startswith('clb') and i.isChecked():
                    print_doc(os.path.abspath(os.curdir) + f'/Документы/{folder}/', i.text() + '.docx')
                    _set_doc_warning = 0
                    break
                else:
                    _set_doc_warning = 1
            if _set_doc_warning:
                set_doc_warning("Ошибка (не выбран документ для печати)",
                                'Сначала выберите документ для печати.\n\nНажмите на нужный документ, '
                                'чтобы выбрать его, а потом нажмите на кнопку "Печать"')

        def del_selected_doc():
            current_tab = self.ui.tabWidget_docx.currentWidget()
            current_list = current_tab.findChild(
                QtWidgets.QWidget,
                'sAWContent_' + current_tab.objectName().split('_')[-1]).children()
            _set_doc_warning = 1
            if 'decree' in current_tab.objectName():
                folder = 'Приказы'
            elif 'notes' in current_tab.objectName():
                folder = 'Записки'
            elif 'ttable' in current_tab.objectName():
                folder = 'Расписания'
            elif 'outlay' in current_tab.objectName():
                folder = 'Сметы'
            for i in current_list:
                if i.objectName().startswith('clb') and i.isChecked():
                    os.remove(f'{os.path.abspath(os.curdir)}/Документы/{folder}/{i.text()}.docx')
                    self.load_list()
                    _set_doc_warning = 0
                    break
                else:
                    _set_doc_warning = 1
            if _set_doc_warning:
                set_doc_warning("Ошибка (не выбран документ для удаления)",
                                'Сначала выберите документ для удаления.\n\nНажмите на нужный документ, '
                                'чтобы выбрать его, а потом нажмите на кнопку "Удалить выбранный документ"')

        # But for timetable
        def timetable_print():
            timetable_list = self.ui.sAWContent_timetable.children()
            _set_doc_warning = 1
            i_name = ""
            for i in timetable_list:
                if i.objectName().startswith("clb_"):
                    if i.isChecked():
                        i_name = i.objectName().split("_")[-1]
                        create_timetable(i_name)
                        _set_doc_warning = 0
                        break
                    else:
                        _set_doc_warning = 1
            if _set_doc_warning:
                set_doc_warning("Ошибка (не выбрано расписание для сохранения)",
                                'Сначала выберите расписание для сохранения.\n\nНажмите на нужное расписание, '
                                'чтобы выбрать его, а потом нажмите на кнопку "Сохранить как документ"')
            else:
                path = os.getcwd() + r"/Документы/Расписания/"
                _db = ARMDataBase('arm_db.db')
                _sql = "SELECT sub_name, id_prog FROM subjects WHERE id_sub=" + i_name
                timetable = _db.query(_sql)
                _sql = "SELECT group_name FROM groups WHERE id_prog=" + str(timetable[0][1])
                try:
                    group_name = _db.query(_sql)[0][0]
                except IndexError:
                    group_name = "Нет группы с этой программой"

                copy_index = 0
                filename = "Расписание " + group_name + " " + timetable[0][0] + " №000000.docx"
                desk_list_dir = os.listdir(path)
                indexes_list = []
                for doc_in_dir in desk_list_dir:
                    start_index = doc_in_dir.find('№')
                    if start_index:
                        try:
                            indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
                        except Exception:
                            pass
                copy_index = 0
                while copy_index in indexes_list:
                    copy_index += 1
                    str_copy_index = str(copy_index)
                    while len(str_copy_index) < 6:
                        str_copy_index = "0" + str_copy_index
                    filename = f"Расписание {group_name} {timetable[0][1]} №{str_copy_index}.docx"
                _db.close()
                set_doc_warning("Отправлено",
                                'Документ будет сохранен в расписания.\n'
                                'Имя документа:\n' + filename)

        # But for notes
        def headers_control_db(type_post):
            headers_list = self.head_ui.sAWContent_headers_list.children()
            _set_doc_warning = 1
            headers_selected = ''
            if type_post == 'save':
                for i in headers_list:
                    if i.objectName().startswith('clb_') and i.isChecked():
                        headers_selected = i.objectName().split('_')[-1]
                        _set_doc_warning = 0
                        break
                    else:
                        _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для изменения)",
                                    'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, измените ее содержимое, а потом нажмите на кнопку '
                                    '"Сохранить выбранную запись"')
                else:
                    if len(self.head_ui.textEdit_headers_fullname.toPlainText().split()) == 3:
                        _db = ARMDataBase()
                        _sql = "UPDATE headers SET " \
                               "head_name = '{0}', " \
                               "head_phone = '{1}', " \
                               "head_mail = '{2}', " \
                               "head_web = '{3}', " \
                               "head_prof = '{4}' " \
                               "WHERE id_head = '{5}'".format(self.head_ui.textEdit_headers_fullname.toPlainText(),
                                                              self.head_ui.textEdit_headers_phone.toPlainText(),
                                                              self.head_ui.textEdit_headers_mail.toPlainText(),
                                                              self.head_ui.textEdit_headers_web.toPlainText(),
                                                              self.head_ui.textEdit_headers_prof.toPlainText(),
                                                              headers_selected)
                        _db.query(_sql)
                        _db.close()
                        self.load_db_headers()
                    else:
                        set_doc_warning("Ошибка (не правильный формат имени)",
                                        'Необходимо ввести имя состоящее из 3-х частей:\n\n'
                                        '"Иванов Иван Андреевич"\n'
                                        '"Иванов И А"\n'
                                        '"Иванов И. А."\n\n'
                                        'Обязательно используйте пробелы!')
            elif type_post == 'add':
                if len(self.head_ui.textEdit_headers_fullname.toPlainText().split()) == 3:
                    _db = ARMDataBase()
                    _sql = "INSERT INTO headers VALUES(" \
                           "NULL," \
                           "'{0}'," \
                           "'{1}'," \
                           "'{2}'," \
                           "'{3}'," \
                           "'{4}')".format(self.head_ui.textEdit_headers_fullname.toPlainText(),
                                           self.head_ui.textEdit_headers_phone.toPlainText(),
                                           self.head_ui.textEdit_headers_mail.toPlainText(),
                                           self.head_ui.textEdit_headers_web.toPlainText(),
                                           self.head_ui.textEdit_headers_prof.toPlainText())
                    _db.query(_sql)
                    _db.close()
                    self.load_db_headers()
                else:
                    set_doc_warning("Ошибка (не правильный формат имени)",
                                    'Необходимо ввести имя состоящее из 3-х частей:\n\n'
                                    '"Иванов Иван Андреевич"\n'
                                    '"Иванов И А"\n'
                                    '"Иванов И. А."\n\n'
                                    'Обязательно используйте пробелы!')
            elif type_post == "del":
                for i in headers_list:
                    if i.objectName() == 'vL_sAWContent_headers_list':
                        pass
                    else:
                        if i.isChecked():
                            headers_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для удаления)",
                                    'Сначала выберите запись для удаления.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, а потом нажмите на кнопку '
                                    '"Удалить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "DELETE FROM headers WHERE id_head={0}".format(headers_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_headers()

        def programs_control_db(type_post):
            programs_list = self.prog_ui.sAWContent_programs_list.children()
            _set_doc_warning = 1
            programs_selected = ''
            if type_post == 'save':
                for i in programs_list:
                    if i.objectName() == 'vL_sAWContent_programs_list':
                        pass
                    else:
                        if i.isChecked():
                            programs_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для изменения)",
                                    'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, измените ее содержимое, а потом нажмите на кнопку '
                                    '"Сохранить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "UPDATE programs SET " \
                           "prog_name = '{0}', " \
                           "prog_range = '{1}', " \
                           "prog_range_dates = '{2}' " \
                           "WHERE id_prog = '{3}'".format(self.prog_ui.textEdit_prog_name.toPlainText(),
                                                          self.prog_ui.textEdit_prog_range.toPlainText(),
                                                          self.prog_ui.dateEdit_start_program.date().
                                                          toString('dd.MM.yyyy')
                                                          + "|" +
                                                          self.prog_ui.dateEdit_end_program.date().
                                                          toString('dd.MM.yyyy'),
                                                          programs_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_programs()
            elif type_post == 'add':
                _db = ARMDataBase()
                _sql = "INSERT INTO programs VALUES(" \
                       "NULL," \
                       "'{0}'," \
                       "'{1}'," \
                       "'{2}')".format(self.prog_ui.textEdit_prog_name.toPlainText(),
                                       self.prog_ui.textEdit_prog_range.toPlainText(),
                                       self.prog_ui.dateEdit_start_program.date().
                                       toString('dd.MM.yyyy')
                                       + "|" +
                                       self.prog_ui.dateEdit_end_program.date().
                                       toString('dd.MM.yyyy'))
                _db.query(_sql)
                _db.close()
                self.load_db_programs()
            elif type_post == "del":
                for i in programs_list:
                    if i.objectName() == 'vL_sAWContent_programs_list':
                        pass
                    else:
                        if i.isChecked():
                            programs_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для удаления)",
                                    'Сначала выберите запись для удаления.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, а потом нажмите на кнопку '
                                    '"Удалить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "SELECT id_group FROM groups WHERE id_prog=" + programs_selected
                    groups = _db.query(_sql)
                    for group in groups:
                        _sql = "SELECT id_student FROM students WHERE id_group=" + group
                        studs = _db.query(_sql)
                        for stud in studs:
                            _sql = "UPDATE subs_in_studs SET status='0' WHERE id_student=" + stud[0]
                            _db.query(_sql)
                    _sql = "UPDATE groups SET id_prog='8' WHERE id_prog=" + programs_selected
                    _db.query(_sql)
                    _sql = "UPDATE subjects SET id_prog='8' WHERE id_prog=" + programs_selected
                    _db.query(_sql)
                    _sql = "DELETE FROM programs WHERE id_prog={0}".format(programs_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_programs()

        def teachers_control_db(type_post):
            teachers_list = self.teach_ui.sAWContent_teachers_list.children()
            _set_doc_warning = 1
            teachers_selected = ''
            if type_post == 'save':
                for i in teachers_list:
                    if i.objectName().startswith('clb_') and i.isChecked():
                        teachers_selected = i.objectName().split('_')[-1]
                        _set_doc_warning = 0
                        break
                    else:
                        _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для изменения)",
                                    'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, измените ее содержимое, а потом нажмите на кнопку '
                                    '"Сохранить выбранную запись"')
                else:
                    if len(self.teach_ui.textEdit_teachers_fullname.toPlainText()) == 3:
                        _db = ARMDataBase()
                        _sql = "UPDATE teachers SET " \
                               "teacher_name = '{0}', " \
                               "teacher_phone = '{1}', " \
                               "teacher_mail = '{2}', " \
                               "teacher_web = '{3}', " \
                               "teacher_prof = '{4}' " \
                               "WHERE id_teacher = '{5}'".format(self.teach_ui.textEdit_teachers_fullname.toPlainText(),
                                                               self.teach_ui.textEdit_teachers_phone.toPlainText(),
                                                               self.teach_ui.textEdit_teachers_mail.toPlainText(),
                                                               self.teach_ui.textEdit_teachers_web.toPlainText(),
                                                               self.teach_ui.textEdit_teachers_prof.toPlainText(),
                                                               teachers_selected)
                        _db.query(_sql)
                        _db.close()
                        self.load_db_teachers()
                    else:
                        set_doc_warning("Ошибка (не правильный формат имени)",
                                        'Необходимо ввести имя состоящее из 3-х частей:\n\n'
                                        '"Иванов Иван Андреевич"\n'
                                        '"Иванов И А"\n'
                                        '"Иванов И. А."\n\n'
                                        'Обязательно используйте пробелы!')
            elif type_post == 'add':
                if len(self.teach_ui.textEdit_teachers_fullname.toPlainText()) == 3:
                    _db = ARMDataBase()
                    _sql = "INSERT INTO teachers VALUES(" \
                           "NULL," \
                           "'{0}'," \
                           "'{1}'," \
                           "'{2}'," \
                           "'{3}'," \
                           "'{4}')".format(self.teach_ui.textEdit_teachers_fullname.toPlainText(),
                                           self.teach_ui.textEdit_teachers_phone.toPlainText(),
                                           self.teach_ui.textEdit_teachers_mail.toPlainText(),
                                           self.teach_ui.textEdit_teachers_web.toPlainText(),
                                           self.teach_ui.textEdit_teachers_prof.toPlainText())
                    _db.query(_sql)
                    _db.close()
                    self.load_db_teachers()
                else:
                    set_doc_warning("Ошибка (не правильный формат имени)",
                                    'Необходимо ввести имя состоящее из 3-х частей:\n\n'
                                    '"Иванов Иван Андреевич"\n'
                                    '"Иванов И А"\n'
                                    '"Иванов И. А."\n\n'
                                    'Обязательно используйте пробелы!')
            elif type_post == "del":
                for i in teachers_list:
                    if i.objectName() == 'vL_sAWContent_teachers_list':
                        pass
                    else:
                        if i.isChecked():
                            teachers_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для удаления)",
                                    'Сначала выберите запись для удаления.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, а потом нажмите на кнопку '
                                    '"Удалить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "UPDATE subjects SET id_teacher='1' WHERE id_teacher={0}".format(teachers_selected)
                    _db.query(_sql)
                    _sql = "DELETE FROM teachers WHERE id_teacher={0}".format(teachers_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_teachers()

        def groups_control_db(type_post):
            groups_list = self.groups_ui.sAWContent_groups_list.children()
            _set_doc_warning = 1
            groups_selected = ''
            if type_post == 'save':
                for i in groups_list:
                    if i.objectName() == 'vL_sAWContent_groups_list':
                        pass
                    else:
                        if i.isChecked():
                            groups_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для изменения)",
                                    'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, измените ее содержимое, а потом нажмите на кнопку '
                                    '"Сохранить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "SELECT id_prog FROM groups WHERE id_group=" + groups_selected
                    prog_checker = str(_db.query(_sql)[0][0])
                    if prog_checker != str(self.groups_ui.comboBox_groups_prog.currentData()):
                        _sql = "SELECT id_student FROM students WHERE id_group=" + groups_selected
                        studs = _db.query(_sql)
                        for stud in studs:
                            _sql = "UPDATE subs_in_studs SET status='0' WHERE id_student=" + stud[0]
                            _db.query(_sql)
                    _sql = "UPDATE groups SET " \
                           "group_name = '{0}', " \
                           "class = '{1}', " \
                           "id_prog = '{2}' " \
                           "WHERE id_group = '{3}'".format(self.groups_ui.textEdit_groups_name.toPlainText(),
                                                           self.groups_ui.textEdit_groups_class.toPlainText(),
                                                           str(self.groups_ui.comboBox_groups_prog.currentData()),
                                                           groups_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_groups()
            elif type_post == 'add':
                _db = ARMDataBase()
                _sql = "INSERT INTO groups VALUES(" \
                       "NULL," \
                       "'{0}'," \
                       "'{1}'," \
                       "'{2}')".format(self.groups_ui.textEdit_groups_class.toPlainText(),
                                       self.groups_ui.textEdit_groups_name.toPlainText(),
                                       str(self.groups_ui.comboBox_groups_prog.currentData()))
                _db.query(_sql)
                _db.close()
                self.load_db_groups()
            elif type_post == "del":
                for i in groups_list:
                    if i.objectName() == 'vL_sAWContent_groups_list':
                        pass
                    else:
                        if i.isChecked():
                            groups_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для удаления)",
                                    'Сначала выберите запись для удаления.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, а потом нажмите на кнопку '
                                    '"Удалить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "SELECT id_student FROM students WHERE id_group=" + groups_selected
                    studs = _db.query(_sql)
                    for stud in studs:
                        _sql = "UPDATE students SET id_group='1' WHERE id_student=" + str(stud[0])
                        _db.query(_sql)
                        _sql = "UPDATE subs_in_studs SET status='0' WHERE id_student=" + str(stud[0])
                        _db.query(_sql)
                    _sql = "DELETE FROM groups WHERE id_group={0}".format(groups_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_groups()

        def subjects_control_db(type_post):
            subjects_list = self.sub_ui.sAWContent_sub_list.children()
            _set_doc_warning = 1
            subjects_selected = ''
            if type_post == 'save':
                for i in subjects_list:
                    if i.objectName() == 'vL_sAWContent_sub_list':
                        pass
                    else:
                        if i.isChecked():
                            subjects_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для изменения)",
                                    'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, измените ее содержимое, а потом нажмите на кнопку '
                                    '"Сохранить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "SELECT id_prog FROM subjects WHERE id_sub=" + subjects_selected
                    prog_checker = str(_db.query(_sql)[0][0])
                    if prog_checker != str(self.sub_ui.comboBox_sub_prog.currentData()):
                        _sql = "UPDATE subs_in_studs SET status='0' WHERE id_sub=" + subjects_selected
                        _db.query(_sql)
                    _sql = "UPDATE subjects SET " \
                           "sub_name = '{0}', " \
                           "sub_price_hour = '{1}', " \
                           "id_teacher = '{2}', " \
                           "sub_price_month = '{3}', " \
                           "id_prog = '{4}', " \
                           "sub_hours_need = '{5}' " \
                           "WHERE id_sub = '{6}'".format(self.sub_ui.textEdit_sub_name.toPlainText(),
                                                         self.sub_ui.lineEdit_sub_tax.text(),
                                                         str(self.sub_ui.comboBox_sub_teach.currentData()),
                                                         self.sub_ui.lineEdit_sub_price.text(),
                                                         str(self.sub_ui.comboBox_sub_prog.currentData()),
                                                         self.sub_ui.lineEdit_sub_hours.text(),
                                                         subjects_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_subjects()
            elif type_post == 'add':
                _db = ARMDataBase()
                _sql = "INSERT INTO subjects VALUES(" \
                       "NULL," \
                       "'{0}'," \
                       "'{1}'," \
                       "'{2}'," \
                       "'{3}'," \
                       "'{4}'," \
                       "NULL," \
                       "'0'," \
                       "'{5}')".format(self.sub_ui.textEdit_sub_name.toPlainText(),
                                       self.sub_ui.lineEdit_sub_tax.text(),
                                       str(self.sub_ui.comboBox_sub_teach.currentData()),
                                       self.sub_ui.lineEdit_sub_price.text(),
                                       str(self.sub_ui.comboBox_sub_prog.currentData()),
                                       self.sub_ui.lineEdit_sub_hours.text())
                _db.query(_sql)
                _db.close()
                self.load_db_subjects()
            elif type_post == "del":
                for i in subjects_list:
                    if i.objectName() == 'vL_sAWContent_sub_list':
                        pass
                    else:
                        if i.isChecked():
                            subjects_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для удаления)",
                                    'Сначала выберите запись для удаления.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, а потом нажмите на кнопку '
                                    '"Удалить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "DELETE FROM subs_in_studs WHERE id_sub=" + subjects_selected
                    _db.query(_sql)
                    _sql = "DELETE FROM subjects WHERE id_sub=" + subjects_selected
                    _db.query(_sql)
                    _db.close()
                    self.load_db_subjects()

        def students_control_db(type_post):
            students_list = self.stud_ui.sAWContent_stud_list.children()
            _set_doc_warning = 1
            students_selected = ''
            if type_post == 'save':
                for i in students_list:
                    if i.objectName().startswith('clb_') and i.isChecked():
                        students_selected = i.objectName().split('_')[-1]
                        _set_doc_warning = 0
                        break
                    else:
                        _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для изменения)",
                                    'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, измените ее содержимое, а потом нажмите на кнопку '
                                    '"Сохранить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    gender = 'male' if self.stud_ui.radioButton_stud_gender_male.isChecked() else 'female'
                    _sql = "SELECT id_group FROM students WHERE id_student=" + students_selected
                    group_checker = str(_db.query(_sql)[0][0])
                    if group_checker != str(self.stud_ui.comboBox_stud_group.currentData()):
                        _sql = "UPDATE subs_in_studs SET status='0' WHERE id_student=" + students_selected
                        _db.query(_sql)
                    _sql = "UPDATE students SET " \
                           "student_name = '{0}', " \
                           "id_group = '{1}', " \
                           "student_birthday = '{2}', " \
                           "student_phone = '{3}', " \
                           "student_gender = '{4}', " \
                           "student_city = '{5}', " \
                           "student_einst = '{6}', " \
                           "student_mail = '{7}', " \
                           "student_web = '{8}' " \
                           "WHERE id_student = '{9}'".format(self.stud_ui.textEdit_stud_fullname.toPlainText(),
                                                             str(self.stud_ui.comboBox_stud_group.currentData()),
                                                             self.stud_ui.dateEdit_stud_birthday.date().toString(
                                                                 'dd.MM.yyyy'),
                                                             self.stud_ui.lineEdit_stud_phone.text(),
                                                             gender,
                                                             self.stud_ui.lineEdit_stud_city.text(),
                                                             self.stud_ui.lineEdit_stud_einst.text(),
                                                             self.stud_ui.lineEdit_stud_mail.text(),
                                                             self.stud_ui.lineEdit_stud_web.text(),
                                                             students_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_students()
            elif type_post == 'add':
                gender = 'male' if self.stud_ui.radioButton_stud_gender_male.isChecked() else 'female'
                _db = ARMDataBase()
                _sql = "INSERT INTO students VALUES(" \
                       "NULL," \
                       "'{0}'," \
                       "'{1}'," \
                       "'{2}'," \
                       "'{3}'," \
                       "'{4}'," \
                       "'{5}'," \
                       "'{6}'," \
                       "'{7}'," \
                       "'{8}')".format(self.stud_ui.textEdit_stud_fullname.toPlainText(),
                                       str(self.stud_ui.comboBox_stud_group.currentData()),
                                       self.stud_ui.dateEdit_stud_birthday.date().toString(
                                           'dd.MM.yyyy'),
                                       self.stud_ui.lineEdit_stud_phone.text(),
                                       gender,
                                       self.stud_ui.lineEdit_stud_city.text(),
                                       self.stud_ui.lineEdit_stud_einst.text(),
                                       self.stud_ui.lineEdit_stud_mail.text(),
                                       self.stud_ui.lineEdit_stud_web.text())
                _db.query(_sql)
                _db.close()
                self.load_db_students()
            elif type_post == "del":
                for i in students_list:
                    if i.objectName() == 'vL_sAWContent_stud_list':
                        pass
                    else:
                        if i.isChecked():
                            students_selected = i.objectName().split('_')[-1]
                            _set_doc_warning = 0
                            break
                        else:
                            _set_doc_warning = 1
                if _set_doc_warning:
                    set_doc_warning("Ошибка (не выбрана запись для удаления)",
                                    'Сначала выберите запись для удаления.\n\nНажмите на нужную запись, '
                                    'чтобы выбрать ее, а потом нажмите на кнопку '
                                    '"Удалить выбранную запись"')
                else:
                    _db = ARMDataBase()
                    _sql = "DELETE FROM subs_in_studs WHERE id_student=" + students_selected
                    _db.query(_sql)
                    _sql = "DELETE FROM students WHERE id_student={0}".format(students_selected)
                    _db.query(_sql)
                    _db.close()
                    self.load_db_students()

        def enrollment_control_db():
            _db = ARMDataBase()
            enrollments_list = self.enr_ui.sAWContent_enr_list.children()
            _set_doc_warning = 1
            enrollment_selected = ''
            for i in enrollments_list:
                if i.objectName() == 'vL_sAWContent_enr_list':
                    pass
                else:
                    if i.isChecked():
                        enrollment_selected = i.objectName().split('_')[-1]
                        _set_doc_warning = 0
                        break
                    else:
                        _set_doc_warning = 1
            if _set_doc_warning:
                set_doc_warning("Ошибка (не выбрана запись для изменения)",
                                'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                                'чтобы выбрать ее, измените ее содержимое, а потом нажмите на кнопку '
                                '"Сохранить выбранную запись"')
            else:
                for enrollment in self.enr_ui.list_cb_checked:
                    _sql = "SELECT id_sis FROM subs_in_studs WHERE id_student=" + enrollment_selected + \
                           " AND id_sub=" + enrollment[0]
                    check_sis = _db.query(_sql)
                    if check_sis:
                        _sql = "UPDATE subs_in_studs SET " \
                               "student_numcontract = '{0}', " \
                               "student_datecontract = '{1}', " \
                               "status = '{2}' " \
                               "WHERE id_sis = '{3}'". \
                            format(self.enr_ui.groupBox_sis_contracts.findChild(QtWidgets.QGroupBox,
                                                                                "grB_" + enrollment[0]).
                                   findChild(QtWidgets.QLineEdit, "ledit_" + "grB_" + enrollment[0]).text(),
                                   self.enr_ui.groupBox_sis_contracts.findChild(QtWidgets.QGroupBox,
                                                                                "grB_" + enrollment[0]).
                                   findChild(QtWidgets.QDateEdit, "dedit_" + "grB_" + enrollment[0]).date().
                                   toString('dd.MM.yyyy'),
                                   "1" if enrollment[1] else "0",
                                   check_sis[0][0])
                    else:
                        if enrollment[1]:
                            _sql = "INSERT INTO subs_in_studs VALUES(" \
                                   "NULL, " \
                                   "'{0}', " \
                                   "'{1}', " \
                                   "'{2}', " \
                                   "'{3}', " \
                                   "'{4}')".format(enrollment_selected,
                                                   enrollment[0],
                                                   self.enr_ui.groupBox_sis_contracts.findChild(QtWidgets.QGroupBox,
                                                                                                "grB_" + enrollment[0]).
                                                   findChild(QtWidgets.QLineEdit,
                                                             "ledit_" + "grB_" + enrollment[0]).text(),
                                                   self.enr_ui.groupBox_sis_contracts.findChild(QtWidgets.QGroupBox,
                                                                                                "grB_" + enrollment[0]).
                                                   findChild(QtWidgets.QDateEdit, "dedit_" + "grB_" + enrollment[0]).
                                                   date().toString('dd.MM.yyyy'),
                                                   "1" if enrollment[1] else "0")
                    _db.query(_sql)
                _db.close()
                self.load_db_enrollment()

        def outlay_control_db():
            self.save_calculate_values()
            self.outlay_ui.calcs_before = len(self.outlay_ui.widget_calcs.children()) - 1
            clear_widget(self.outlay_ui.widget_calcs.children())
            if self.outlay_ui.comboBox_progs.currentData() == "None" or self.outlay_ui.comboBox_progs.currentData() is None:
                self.outlay_ui.widget_col_subs.setEnabled(True)
                rb = 0
                if self.outlay_ui.radio_col_1.isChecked():
                    rb = 1
                elif self.outlay_ui.radio_col_2.isChecked():
                    rb = 2
                elif self.outlay_ui.radio_col_3.isChecked():
                    rb = 3
                elif self.outlay_ui.radio_col_4.isChecked():
                    rb = 4
                for r in range(rb):
                    if r + 1 == 1:
                        self.add_calculate_box(False, str(r + 1), 0, 0)
                    elif r + 1 == 2:
                        self.add_calculate_box(False, str(r + 1), 1, 0)
                    elif r + 1 == 3:
                        self.add_calculate_box(False, str(r + 1), 0, 1)
                    elif r + 1 == 4:
                        self.add_calculate_box(False, str(r + 1), 1, 1)
                self.load_calculate_values()
            else:
                self.outlay_ui.widget_col_subs.setEnabled(False)
                _db = ARMDataBase()
                _sql = "SELECT id_sub FROM subjects WHERE id_prog=" + str(self.outlay_ui.comboBox_progs.currentData())
                subs = _db.query(_sql)
                try:
                    self.outlay_ui.widget_col_subs.findChild(QtWidgets.QRadioButton,
                                                             "radio_col_" + str(len(subs))).setChecked(True)
                except Exception:
                    pass
                _db.close()
                a = 0
                b = 0
                for i in subs:
                    self.add_calculate_box(True, str(i[0]), a, b)
                    if a == 0 and b == 0:
                        a = 1
                        b = 0
                    elif a == 1 and b == 0:
                        a = 0
                        b = 1
                    elif a == 0 and b == 1:
                        a = 1
                        b = 1

        def open_selected_docx():
            current_tab = self.ui.tabWidget_docx.currentWidget()
            content_list = current_tab.findChild(
                QtWidgets.QWidget, "sAWContent_" + current_tab.objectName().split("_")[-1]).children()
            docx_folder = ""
            if "decree" in current_tab.objectName():
                docx_folder = r'Приказы'
            elif "note" in current_tab.objectName():
                docx_folder = r'Записки'
            elif "ttable" in current_tab.objectName():
                docx_folder = r'Расписания'
            elif "outlay" in current_tab.objectName():
                docx_folder = r'Сметы'
            elif "contracts" in current_tab.objectName():
                docx_folder = r'Договора'
            for clb in content_list:
                if clb.objectName().startswith("clb_") and clb.isChecked():
                    command = f'"{os.path.abspath(os.curdir)}\\Документы\\{docx_folder}\\{clb.text()}.docx"'
                    open_file(command)

        # SETUP BUTS
        self.ui.pushButton_print_docx.clicked.connect(lambda: print_selected_doc())
        self.ui.pushButton_edit_docx.clicked.connect(lambda: open_selected_docx())
        self.ui.pushButton_del_docx.clicked.connect(lambda: del_selected_doc())
        self.ui.pushButton_print_timetable.clicked.connect(lambda: timetable_print())
        self.ui.pushButton_update_docx.clicked.connect(
            lambda: self.load_list())
        self.ui.lineEdit_search_docx.textEdited.connect(lambda: self.load_list())
        self.ui.pushButton_create_docx.clicked.connect(lambda: self.decree_creator_exec())

        self.ui.pushButton_update_timetable.clicked.connect(
            lambda: self.load_db_timetable(self.ui.lineEdit_search_timetable.text()))
        self.ui.pushButton_edit_timetable.clicked.connect(lambda: self.timetable_list_exec())
        self.ui.lineEdit_search_timetable.textEdited.connect(
            lambda: self.load_db_timetable(self.ui.lineEdit_search_timetable.text()))
        self.ttable_ui.btn_set_hours.clicked.connect(lambda: self.set_hours())
        self.ttable_ui.btn_del_hours.clicked.connect(lambda: self.del_hours())
        self.ttable_ui.calendar.clicked.connect(lambda: self.select_list_el())

        self.ui.pushButton_headers_roster.clicked.connect(lambda: self.headers_win())
        self.ui.pushButton_programs_roster.clicked.connect(lambda: self.programs_win())
        self.ui.pushButton_teachers_roster.clicked.connect(lambda: self.teachers_win())
        self.ui.pushButton_groups_roster.clicked.connect(lambda: self.groups_win())
        self.ui.pushButton_subjects_roster.clicked.connect(lambda: self.subjects_win())
        self.ui.pushButton_students_roster.clicked.connect(lambda: self.students_win())
        self.ui.pushButton_enrollment_roster.clicked.connect(lambda: self.enrollment_win())
        self.ui.pushButton_outlay.clicked.connect(lambda: self.outlay_win())
        self.ui.settings.triggered.connect(lambda: self.settings_window())

        self.head_ui.pushButton_headers_add.clicked.connect(lambda: headers_control_db('add'))
        self.head_ui.pushButton_headers_save.clicked.connect(lambda: headers_control_db('save'))
        self.head_ui.pushButton_headers_delete.clicked.connect(lambda: headers_control_db('del'))
        self.head_ui.pushButton_headers_back.clicked.connect(lambda: headers_back())
        self.head_ui.lineEdit_search_headers.textEdited.connect(
            lambda: self.load_db_headers(self.head_ui.lineEdit_search_headers.text()))

        self.prog_ui.pushButton_programs_add.clicked.connect(lambda: programs_control_db('add'))
        self.prog_ui.pushButton_programs_save.clicked.connect(lambda: programs_control_db('save'))
        self.prog_ui.pushButton_programs_delete.clicked.connect(lambda: programs_control_db('del'))
        self.prog_ui.pushButton_programs_back.clicked.connect(lambda: programs_back())
        self.prog_ui.lineEdit_search_programs.textEdited.connect(
            lambda: self.load_db_programs(self.prog_ui.lineEdit_search_programs.text()))

        self.teach_ui.pushButton_teachers_add.clicked.connect(lambda: teachers_control_db('add'))
        self.teach_ui.pushButton_teachers_save.clicked.connect(lambda: teachers_control_db('save'))
        self.teach_ui.pushButton_teachers_delete.clicked.connect(lambda: teachers_control_db('del'))
        self.teach_ui.pushButton_teachers_back.clicked.connect(lambda: teachers_back())
        self.teach_ui.lineEdit_search_teachers.textEdited.connect(
            lambda: self.load_db_teachers(self.teach_ui.lineEdit_search_teachers.text()))

        self.groups_ui.pushButton_groups_add.clicked.connect(lambda: groups_control_db('add'))
        self.groups_ui.pushButton_groups_save.clicked.connect(lambda: groups_control_db('save'))
        self.groups_ui.pushButton_groups_delete.clicked.connect(lambda: groups_control_db('del'))
        self.groups_ui.pushButton_groups_back.clicked.connect(lambda: groups_back())
        self.groups_ui.lineEdit_search_groups.textEdited.connect(
            lambda: self.load_db_groups(self.groups_ui.lineEdit_search_groups.text()))

        self.sub_ui.pushButton_sub_add.clicked.connect(lambda: subjects_control_db('add'))
        self.sub_ui.pushButton_sub_save.clicked.connect(lambda: subjects_control_db('save'))
        self.sub_ui.pushButton_sub_delete.clicked.connect(lambda: subjects_control_db('del'))
        self.sub_ui.pushButton_sub_back.clicked.connect(lambda: subjects_back())
        self.sub_ui.lineEdit_search_sub.textEdited.connect(
            lambda: self.load_db_subjects(self.sub_ui.lineEdit_search_sub.text()))

        self.stud_ui.pushButton_stud_add.clicked.connect(lambda: students_control_db('add'))
        self.stud_ui.pushButton_stud_save.clicked.connect(lambda: students_control_db('save'))
        self.stud_ui.pushButton_stud_delete.clicked.connect(lambda: students_control_db('del'))
        self.stud_ui.pushButton_stud_back.clicked.connect(lambda: students_back())
        self.stud_ui.lineEdit_search_stud.textEdited.connect(
            lambda: self.load_db_students(self.stud_ui.lineEdit_search_stud.text()))

        self.enr_ui.pushButton_enr_save.clicked.connect(lambda: enrollment_control_db())
        self.enr_ui.pushButton_enr_back.clicked.connect(lambda: enrollment_back())
        self.enr_ui.lineEdit_search_enr.textEdited.connect(
            lambda: self.load_db_enrollment(self.enr_ui.lineEdit_search_enr.text()))

        self.outlay_ui.pushButton_outlay_back.clicked.connect(lambda: outlay_back())
        self.outlay_ui.comboBox_progs.currentIndexChanged.connect(lambda: outlay_control_db())
        self.outlay_ui.pushButton_outlay_next.clicked.connect(lambda: self.outlay_printer_exec())
        self.outlay_ui.radio_col_1.clicked.connect(lambda: outlay_control_db())
        self.outlay_ui.radio_col_2.clicked.connect(lambda: outlay_control_db())
        self.outlay_ui.radio_col_3.clicked.connect(lambda: outlay_control_db())
        self.outlay_ui.radio_col_4.clicked.connect(lambda: outlay_control_db())
        self.outlay_print_ui.pushButton_save_doc.clicked.connect(lambda: self.create_outlay())

        self.docx_ui.pushButton_decree_enrollment.clicked.connect(lambda: self.decree_enr_win())
        self.docx_ui.pushButton_note_passes.clicked.connect(lambda: self.note_passes_win())
        self.docx_ui.pushButton_note_passwords.clicked.connect(lambda: self.note_passwords_win())
        self.docx_ui.pushButton_note_studs_list.clicked.connect(lambda: self.note_list_win())
        self.docx_ui.pushButton_contract.clicked.connect(lambda: self.contract_win())

        self.decree_enr_ui.pushButton_back.clicked.connect(lambda: decree_enr_back())
        self.decree_enr_ui.pushButton_save_doc.clicked.connect(lambda: self.create_decree_enr())

        self.note_passes_ui.pushButton_back.clicked.connect(lambda: note_passes_back())
        self.note_passes_ui.pushButton_save_doc.clicked.connect(lambda: self.create_note_passes())

        self.note_passwords_ui.pushButton_back.clicked.connect(lambda: note_passwords_back())
        self.note_passwords_ui.pushButton_save_doc.clicked.connect(lambda: self.create_note_passwords())

        self.note_list_ui.pushButton_back.clicked.connect(lambda: note_list_back())
        self.note_list_ui.pushButton_save_doc.clicked.connect(lambda: self.create_note_list())

        self.contract_ui.pushButton_back.clicked.connect(lambda: contract_back())
        self.contract_ui.pushButton_save_doc.clicked.connect(lambda: self.create_contract())

    # END BUTTONS
    def clear_for_start(self):
        self.ui.widget_headers.hide()
        self.ui.widget_programs.hide()
        self.ui.widget_teachers.hide()
        self.ui.widget_groups.hide()
        self.ui.widget_subjects.hide()
        self.ui.widget_students.hide()
        self.ui.widget_enrollment.hide()
        self.ui.widget_outlay.hide()
        self.docx_ui.widget_enrollment.hide()

    def load_list(self):
        clear_list(self.ui.sAWContent_decree.children())
        clear_list(self.ui.sAWContent_outlay.children())
        clear_list(self.ui.sAWContent_notes.children())
        clear_list(self.ui.sAWContent_ttable.children())
        clear_list(self.ui.sAWContent_contracts.children())
        list_docx = [['decree', self.ui.sAWContent_decree, 'Приказы'],
                     ['notes', self.ui.sAWContent_notes, 'Записки'],
                     ['ttable', self.ui.sAWContent_ttable, 'Расписания'],
                     ['contacts', self.ui.sAWContent_contracts, 'Договора'],
                     ['outlay', self.ui.sAWContent_outlay, 'Сметы']]
        for docxs in list_docx:
            list_dir = os.listdir(os.path.abspath(os.curdir) + r"/Документы/" + docxs[2])
            items = []
            for _docx in list_dir:
                if not _docx.startswith("~$"):
                    pos1 = _docx.find('№')
                    doc_id = _docx[pos1 + 1:pos1 + 7]
                    items.append(self.create_list_el("clb_" + docxs[0] + doc_id, _docx[:-5], docxs[1]))

            _search_text = self.ui.tabWidget_Main.findChild(
                QtWidgets.QWidget, "lineEdit_search_docx").text().lower()
            for i in items:
                searcher = i.text().lower()
                if _search_text in searcher:
                    i.show()
                else:
                    i.hide()

    def load_for_start(self):
        # Loading doc's for lists with doc's
        self.load_list()
        self.load_db_timetable()
        # Add normal icon
        self.ui.icon = QtGui.QIcon()
        self.ui.icon.addPixmap(QtGui.QPixmap(":/sfu_logo.ico"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.setWindowIcon(self.ui.icon)
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(":/sfu_logo.ico"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.ttable.setWindowIcon(icon)
        self.outlay_printer.setWindowIcon(icon)

    def create_list_el(self, name, text, ls, auto_exclusive=True):
        a = QtWidgets.QCommandLinkButton(ls)
        font = QtGui.QFont()
        font.setFamily("Segoe UI")
        font.setPointSize(10)
        a.setFont(font)
        a.setCheckable(True)
        a.setChecked(False)
        a.setAutoExclusive(auto_exclusive)
        a.setAutoDefault(False)
        a.setDefault(False)
        a.setObjectName(name)
        ls_Layout = ls.children()[0]
        a.setMinimumSize(len(max(text.split("\n"), key=lambda i: len(i))) * 6 + 100, text.count("\n") * 15 + 40)
        ls_Layout.addWidget(a)
        a.setText(self._translate("MainWindow", text))
        return a

    def create_combo_box_el(self, ls, data, text):
        ls.addItem(text, data)
        return ls

    def create_check_box_el(self, ls, name, text, checked=False):
        cb = QtWidgets.QCheckBox(ls)
        font = QtGui.QFont()
        font.setPointSize(11)
        cb.setFont(font)
        cb.setObjectName(name)
        cb.setChecked(checked)
        ls_Layout = ls.children()[0]
        ls_Layout.addWidget(cb)
        cb.setText(self._translate("MainWindow", text))
        return cb

    def create_sub_groupbox(self, ls, name, text, _hide):
        # GroupBox for Subject
        gb = QtWidgets.QGroupBox(ls)
        font = QtGui.QFont()
        font.setPointSize(11)
        gb.setFont(font)
        gb.setObjectName(name)
        gL_gb = QtWidgets.QGridLayout(gb)
        gL_gb.setObjectName("gL_" + name)
        gb.setTitle(self._translate("Enrollment", text))
        ls_Layout = ls.children()[0]
        ls_Layout.addWidget(gb)
        # Label for Num Contract
        lb_num = QtWidgets.QLabel(gb)
        font = QtGui.QFont()
        font.setPointSize(12)
        lb_num.setFont(font)
        lb_num.setObjectName("lb_num_" + name)
        lb_num.setText(self._translate("Enrollment", "Номер контракта"))
        gL_gb.addWidget(lb_num, 0, 0)
        # LineEdit for Num Contract
        ledit = QtWidgets.QLineEdit(gb)
        ledit.setObjectName("ledit_" + name)
        gL_gb.addWidget(ledit, 0, 1)
        # Label for Date
        lb_date = QtWidgets.QLabel(gb)
        font = QtGui.QFont()
        font.setPointSize(12)
        lb_date.setFont(font)
        lb_date.setObjectName("lb_date_" + name)
        lb_date.setText(self._translate("Enrollment", "Дата заключения контракта"))
        gL_gb.addWidget(lb_date, 1, 0)
        # DateEdit for Date
        dedit = QtWidgets.QDateEdit(gb)
        font = QtGui.QFont()
        font.setPointSize(10)
        dedit.setFont(font)
        dedit.setDate(QtCore.QDate(2000, 1, 1))
        dedit.setObjectName("dedit_" + name)
        gL_gb.addWidget(dedit, 1, 1)
        if _hide:
            gb.hide()
        return gb

    # Loader database for header
    def load_db_headers(self, search_text=None):
        clear_list(self.head_ui.sAWContent_headers_list.children())

        def loader_headers_edits():
            selected_header = ''
            headers_list = self.head_ui.sAWContent_headers_list.children()
            if len(headers_list) != 2:
                for head in headers_list:
                    if head.objectName() != 'vL_sAWContent_headers_list':
                        if head.isChecked():
                            selected_header = head.objectName().split('_')[-1]
                            break
            else:
                selected_header = self.head_ui.sAWContent_headers_list.children()[-1].objectName().split('_')[-1]
            _db1 = ARMDataBase()
            _sql1 = "SELECT * FROM headers WHERE id_head=" + selected_header
            _head = _db1.query(_sql1)
            _db1.close()
            self.head_ui.textEdit_headers_fullname.setText(_head[0][1])
            self.head_ui.textEdit_headers_prof.setText(_head[0][5])

            if _head[0][2] is not None and _head[0][2] != '':
                self.head_ui.textEdit_headers_phone.setText(_head[0][2])
            else:
                self.head_ui.textEdit_headers_phone.setText('')

            if _head[0][3] is not None and _head[0][3] != '':
                self.head_ui.textEdit_headers_mail.setText(_head[0][3])
            else:
                self.head_ui.textEdit_headers_mail.setText('')

            if _head[0][4] is not None and _head[0][4] != '':
                self.head_ui.textEdit_headers_web.setText(_head[0][4])
            else:
                self.head_ui.textEdit_headers_web.setText('')

        _db = ARMDataBase()
        _sql = "SELECT * FROM headers"
        headers = _db.query(_sql)
        _db.close()
        head_loader = []
        for i in range(len(headers)):
            heads = []
            for h in headers[i]:
                heads.append(h)
            head_loader.append(str(heads[0])[:])
            heads[0] = 'clb_head_' + str(heads[0])
            heads[1] = 'ФИО: ' + heads[1] + '\n'
            heads[2] = 'Телефоны: ' + heads[2] + '\n' if heads[2] is not None and heads[2] != '' else ''
            heads[3] = 'Электронные почты: ' + heads[3] + '\n' if heads[3] is not None and heads[3] != '' else ''
            heads[4] = 'Социальные сети: ' + heads[4] + '\n' if heads[4] is not None and heads[4] != '' else ''
            heads[5] = 'Должность: ' + heads[5] + '\n' if heads[5] is not None and heads[5] != '' else ''
            searcher = ''
            if search_text is None or search_text == "":
                _search_text = self.head_ui.lineEdit_search_headers.text().lower()
            elif search_text is not None:
                _search_text = search_text.lower()
            else:
                _search_text = search_text
            for h in heads:
                if h is not None and h != '':
                    searcher = searcher + h.lower()
            if _search_text is not None and _search_text != '':
                if _search_text in searcher:
                    head_but = self.create_list_el(heads[0],
                                                   heads[1] + heads[5] + heads[2] + heads[3] + heads[4],
                                                   self.head_ui.sAWContent_headers_list)
                    head_but.clicked.connect(lambda: loader_headers_edits())
            else:
                head_but = self.create_list_el(heads[0],
                                               heads[1] + heads[5] + heads[2] + heads[3] + heads[4],
                                               self.head_ui.sAWContent_headers_list)
                head_but.clicked.connect(lambda: loader_headers_edits())

    # Loader database for programs
    def load_db_programs(self, search_text=None):
        clear_list(self.prog_ui.sAWContent_programs_list.children())

        def loader_programs_edits():
            selected_program = ''
            programs_list = self.prog_ui.sAWContent_programs_list.children()
            if len(programs_list) != 2:
                for prog in programs_list:
                    if prog.objectName() != 'vL_sAWContent_programs_list':
                        if prog.isChecked():
                            selected_program = prog.objectName().split('_')[-1]
                            break
            else:
                selected_program = self.prog_ui.sAWContent_programs_list.children()[-1].objectName().split('_')[-1]
            _db1 = ARMDataBase()
            _sql1 = "SELECT * FROM programs WHERE id_prog=" + selected_program
            _prog = _db1.query(_sql1)
            _db1.close()
            self.prog_ui.textEdit_prog_name.setText(_prog[0][1])
            self.prog_ui.textEdit_prog_range.setText(_prog[0][2])
            dates = _prog[0][3].split("|")
            self.prog_ui.dateEdit_start_program.setDate(datetime.date(int(dates[0].split('.')[2]),
                                                                      int(dates[0].split('.')[1]),
                                                                      int(dates[0].split('.')[0])))
            self.prog_ui.dateEdit_end_program.setDate(datetime.date(int(dates[1].split('.')[2]),
                                                                    int(dates[1].split('.')[1]),
                                                                    int(dates[1].split('.')[0])))

        _db = ARMDataBase()
        _sql = "SELECT * FROM programs"
        programs = _db.query(_sql)
        _db.close()
        prog_loader = []
        for i in range(len(programs)):
            progs = []
            for h in programs[i]:
                progs.append(h)
            if str(progs[0]) != "8":
                prog_loader.append(str(progs[0])[:])
                progs[0] = 'clb_prog_' + str(progs[0])
                progs[1] = 'Программа: ' + progs[1] + '\n'
                progs[2] = 'Продолжительность: в течении ' + progs[2] + ' месяцев\n' if progs[2] is not None and progs[
                    2] != '' else ''
                progs[3] = 'Даты проведения: с ' + progs[3].split("|")[0] + " по " + progs[3].split("|")[1] + '\n'
                searcher = ''
                if search_text is None or search_text == "":
                    _search_text = self.prog_ui.lineEdit_search_programs.text().lower()
                elif search_text is not None:
                    _search_text = search_text.lower()
                else:
                    _search_text = search_text
                for h in progs:
                    if h is not None and h != '':
                        searcher = searcher + h.lower()
                if _search_text is not None and _search_text != '':
                    if _search_text in searcher:
                        prog_but = self.create_list_el(progs[0],
                                                       progs[1] + progs[2] + progs[3],
                                                       self.prog_ui.sAWContent_programs_list)
                        prog_but.clicked.connect(lambda: loader_programs_edits())
                else:
                    prog_but = self.create_list_el(progs[0],
                                                   progs[1] + progs[2] + progs[3],
                                                   self.prog_ui.sAWContent_programs_list)
                    prog_but.clicked.connect(lambda: loader_programs_edits())

    # Loader database for teachers
    def load_db_teachers(self, search_text=None):
        clear_list(self.teach_ui.sAWContent_teachers_list.children())

        def loader_teachers_edits():
            selected_teacher = ''
            teachers_list = self.teach_ui.sAWContent_teachers_list.children()
            if len(teachers_list) != 2:
                for teach in teachers_list:
                    if teach.objectName() != 'vL_sAWContent_teachers_list':
                        if teach.isChecked():
                            selected_teacher = teach.objectName().split('_')[-1]
                            break
            else:
                selected_teacher = self.teach_ui.sAWContent_teachers_list.children()[-1].objectName().split('_')[-1]
            _db1 = ARMDataBase()
            _sql1 = "SELECT * FROM teachers WHERE id_teacher=" + selected_teacher
            _teach = _db1.query(_sql1)
            _db1.close()
            self.teach_ui.textEdit_teachers_fullname.setText(_teach[0][1])
            self.teach_ui.textEdit_teachers_prof.setText(_teach[0][5])

            if _teach[0][2] is not None and _teach[0][2] != '':
                self.teach_ui.textEdit_teachers_phone.setText(_teach[0][2])
            else:
                self.teach_ui.textEdit_teachers_phone.setText('')

            if _teach[0][3] is not None and _teach[0][3] != '':
                self.teach_ui.textEdit_teachers_mail.setText(_teach[0][3])
            else:
                self.teach_ui.textEdit_teachers_mail.setText('')

            if _teach[0][4] is not None and _teach[0][4] != '':
                self.teach_ui.textEdit_teachers_web.setText(_teach[0][4])
            else:
                self.teach_ui.textEdit_teachers_web.setText('')

        _db = ARMDataBase()
        _sql = "SELECT * FROM teachers"
        teachers = _db.query(_sql)
        _db.close()
        teach_loader = []
        for i in range(len(teachers)):
            teachs = []
            for h in teachers[i]:
                teachs.append(h)
            if str(teachs[0]) != "1":
                teach_loader.append(str(teachs[0])[:])
                teachs[0] = 'clb_teach_' + str(teachs[0])
                teachs[1] = 'ФИО: ' + teachs[1] + '\n'
                teachs[2] = 'Телефоны: ' + teachs[2] + '\n' if teachs[2] is not None and teachs[2] != '' else ''
                teachs[3] = 'Электронные почты: ' + teachs[3] + '\n' if teachs[3] is not None and teachs[3] != '' else ''
                teachs[4] = 'Социальные сети: ' + teachs[4] + '\n' if teachs[4] is not None and teachs[4] != '' else ''
                teachs[5] = 'Должность: ' + teachs[5] + '\n' if teachs[5] is not None and teachs[5] != '' else ''
                searcher = ''
                if search_text is None or search_text == "":
                    _search_text = self.teach_ui.lineEdit_search_teachers.text().lower()
                elif search_text is not None:
                    _search_text = search_text.lower()
                else:
                    _search_text = search_text
                for h in teachs:
                    if h is not None and h != '':
                        searcher = searcher + h.lower()
                if _search_text is not None and _search_text != '':
                    if _search_text in searcher:
                        teach_but = self.create_list_el(teachs[0],
                                                        teachs[1] + teachs[5] + teachs[2] + teachs[3] + teachs[4],
                                                        self.teach_ui.sAWContent_teachers_list)
                        teach_but.clicked.connect(lambda: loader_teachers_edits())
                else:
                    teach_but = self.create_list_el(teachs[0],
                                                    teachs[1] + teachs[5] + teachs[2] + teachs[3] + teachs[4],
                                                    self.teach_ui.sAWContent_teachers_list)
                    teach_but.clicked.connect(lambda: loader_teachers_edits())

    # Loader database for groups
    def load_db_groups(self, search_text=None):
        clear_list(self.groups_ui.sAWContent_groups_list.children())

        def loader_groups_edits():
            selected_group = ''
            groups_list = self.groups_ui.sAWContent_groups_list.children()
            if len(groups_list) != 2:
                for grp in groups_list:
                    if grp.objectName() != 'vL_sAWContent_groups_list':
                        if grp.isChecked():
                            selected_group = grp.objectName().split('_')[-1]
                            break
            else:
                selected_group = self.groups_ui.sAWContent_groups_list.children()[-1].objectName().split('_')[-1]
            _db1 = ARMDataBase()
            _sql1 = "SELECT * FROM groups WHERE id_group=" + selected_group
            _grp = _db1.query(_sql1)
            _db1.close()
            self.groups_ui.textEdit_groups_name.setText(_grp[0][2])
            self.groups_ui.comboBox_groups_prog.setCurrentIndex(
                self.groups_ui.comboBox_groups_prog.findData(_grp[0][3]))
            if _grp[0][2] is not None and _grp[0][2] != '':
                self.groups_ui.textEdit_groups_class.setText(_grp[0][1])
            else:
                self.groups_ui.textEdit_groups_class.setText('')

        _db = ARMDataBase()
        _sql = "SELECT * FROM groups"
        groups = _db.query(_sql)

        for i in range(len(groups)):
            grps = []
            for h in groups[i]:
                grps.append(h)

            _sql = "SELECT prog_name FROM programs WHERE id_prog=" + str(grps[3])
            group_prog = _db.query(_sql)
            if not group_prog:
                group_prog = [["Отсутствует"]]

            grps[0] = 'clb_grp_' + str(grps[0])
            grps[1] = 'Класс: ' + grps[1] + '\n' if grps[1] is not None and grps[1] != '' else ''
            grps[2] = 'Группа: ' + grps[2] + '\n'
            grps[3] = 'Программа: ' + group_prog[0][0] + '\n' if group_prog[0][0] is not None and group_prog[0][
                0] != '' else ''
            searcher = ''
            if search_text is None or search_text == "":
                _search_text = self.groups_ui.lineEdit_search_groups.text().lower()
            elif search_text is not None:
                _search_text = search_text.lower()
            else:
                _search_text = search_text
            for h in grps:
                if h is not None and h != '':
                    searcher = searcher + h.lower()
            if grps[0].split("_")[-1] != "1":
                if _search_text is not None and _search_text != '':
                    if _search_text in searcher:
                        grp_but = self.create_list_el(grps[0],
                                                      grps[2] + grps[3] + grps[1],
                                                      self.groups_ui.sAWContent_groups_list)
                        grp_but.clicked.connect(lambda: loader_groups_edits())
                else:
                    grp_but = self.create_list_el(grps[0],
                                                  grps[2] + grps[3] + grps[1],
                                                  self.groups_ui.sAWContent_groups_list)
                    grp_but.clicked.connect(lambda: loader_groups_edits())

        _sql = "SELECT * FROM programs"
        programs = _db.query(_sql)
        _db.close()

        self.groups_ui.comboBox_groups_prog.clear()
        for prog in programs:
            self.create_combo_box_el(self.groups_ui.comboBox_groups_prog, prog[0], str(prog[1]))

    # Loader database for subjects
    def load_db_subjects(self, search_text=None):
        clear_list(self.sub_ui.sAWContent_sub_list.children())

        def loader_sub_edits():
            selected_subject = ''
            subjects_list = self.sub_ui.sAWContent_sub_list.children()
            if len(subjects_list) != 2:
                for sub in subjects_list:
                    if sub.objectName() != 'vL_sAWContent_sub_list':
                        if sub.isChecked():
                            selected_subject = sub.objectName().split('_')[-1]
                            break
            else:
                selected_subject = self.sub_ui.sAWContent_sub_list.children()[-1].objectName().split('_')[-1]
            _db1 = ARMDataBase()
            _sql1 = "SELECT * FROM subjects WHERE id_sub=" + selected_subject
            _sub = _db1.query(_sql1)
            _db1.close()

            self.sub_ui.textEdit_sub_name.setText(_sub[0][1])
            self.sub_ui.lineEdit_sub_tax.setText(_sub[0][2])
            self.sub_ui.comboBox_sub_teach.setCurrentIndex(
                self.sub_ui.comboBox_sub_teach.findData(_sub[0][3]))
            self.sub_ui.lineEdit_sub_price.setText(_sub[0][4])
            self.sub_ui.comboBox_sub_prog.setCurrentIndex(
                self.sub_ui.comboBox_sub_prog.findData(_sub[0][5]))
            self.sub_ui.lineEdit_sub_hours.setText(_sub[0][8])

        _db = ARMDataBase()
        _sql = "SELECT * FROM subjects"
        subjects = _db.query(_sql)
        _db.close()
        sub_loader = []

        _db = ARMDataBase()
        for i in range(len(subjects)):
            subs = []
            for h in subjects[i]:
                subs.append(h)
            sub_loader.append(str(subs[0])[:])

            _sql = "SELECT teacher_name FROM teachers WHERE id_teacher=" + str(subs[3])
            sub_teach = _db.query(_sql)

            _sql = "SELECT prog_name FROM programs WHERE id_prog=" + str(subs[5])
            sub_prog = _db.query(_sql)

            subs[0] = 'clb_sub_' + str(subs[0])
            subs[1] = 'Название: ' + subs[1] + '\n' if subs[1] is not None and subs[1] != '' else ''
            subs[2] = 'Почасовая оплата: ' + subs[2] + '\n' if subs[2] is not None and subs[2] != '' else ''
            try:
                subs[3] = 'Преподаватель: ' + sub_teach[0][0] + '\n' if sub_teach[0][0] is not None and sub_teach[0][
                    0] != '' else ''
            except IndexError:
                subs[3] = ""
            subs[4] = 'Стоимость: ' + subs[4] + '\n' if subs[4] is not None and subs[4] != '' else ''
            subs[5] = 'Программа: ' + sub_prog[0][0] + '\n' if sub_prog[0][0] is not None and sub_prog[0][
                0] != '' else ''
            subs[8] = 'Всего часов: ' + subs[8] + '\n' if subs[8] is not None and subs[8] != '' else ''
            searcher = ''
            if search_text is None or search_text == "":
                _search_text = self.sub_ui.lineEdit_search_sub.text().lower()
            elif search_text is not None:
                _search_text = search_text.lower()
            else:
                _search_text = search_text
            for h in subs:
                if h is not None and h != '':
                    searcher = searcher + h.lower()
            if _search_text is not None and _search_text != '':
                if _search_text in searcher:
                    sub_but = self.create_list_el(subs[0],
                                                  subs[1] + subs[2] + subs[3] + subs[4] + subs[5] + subs[8],
                                                  self.sub_ui.sAWContent_sub_list)
                    sub_but.clicked.connect(lambda: loader_sub_edits())
            else:
                sub_but = self.create_list_el(subs[0],
                                              subs[1] + subs[2] + subs[3] + subs[4] + subs[5] + subs[8],
                                              self.sub_ui.sAWContent_sub_list)
                sub_but.clicked.connect(lambda: loader_sub_edits())
        _db.close()

        _db = ARMDataBase()

        _sql = "SELECT * FROM programs"
        programs = _db.query(_sql)
        _programs = []
        self.sub_ui.comboBox_sub_prog.clear()
        for prog in programs:
            self.create_combo_box_el(self.sub_ui.comboBox_sub_prog, prog[0], str(prog[1]))

        _sql = "SELECT * FROM teachers"
        teachers = _db.query(_sql)
        _teachers = []
        self.sub_ui.comboBox_sub_teach.clear()
        for teach in teachers:
            self.create_combo_box_el(self.sub_ui.comboBox_sub_teach, teach[0], str(teach[1]))

        _db.close()

    # Loader database for students
    def load_db_students(self, search_text=None):
        clear_list(self.stud_ui.sAWContent_stud_list.children())

        def loader_stud_edits():
            selected_students = ''
            students_list = self.stud_ui.sAWContent_stud_list.children()
            if len(students_list) != 2:
                for stud in students_list:
                    if stud.objectName() != 'vL_sAWContent_stud_list':
                        if stud.isChecked():
                            selected_students = stud.objectName().split('_')[-1]
                            break
            else:
                selected_students = self.stud_ui.sAWContent_stud_list.children()[-1].objectName().split('_')[-1]
            _db1 = ARMDataBase()
            _sql1 = "SELECT * FROM students WHERE id_student=" + selected_students
            _stud = _db1.query(_sql1)
            _db1.close()

            self.stud_ui.textEdit_stud_fullname.setText(_stud[0][1])
            self.stud_ui.dateEdit_stud_birthday.setDate(datetime.date(int(_stud[0][3].split('.')[2]),
                                                                      int(_stud[0][3].split('.')[1]),
                                                                      int(_stud[0][3].split('.')[0])))
            self.stud_ui.lineEdit_stud_phone.setText(_stud[0][4])
            if _stud[0][5] == 'male':
                self.stud_ui.radioButton_stud_gender_male.setChecked(1)
                self.stud_ui.radioButton_stud_gender_female.setChecked(0)
            else:
                self.stud_ui.radioButton_stud_gender_female.setChecked(1)
                self.stud_ui.radioButton_stud_gender_male.setChecked(0)
            self.stud_ui.lineEdit_stud_city.setText(_stud[0][6])
            self.stud_ui.lineEdit_stud_einst.setText(_stud[0][7])
            self.stud_ui.lineEdit_stud_mail.setText(_stud[0][8])
            self.stud_ui.lineEdit_stud_web.setText(_stud[0][9])
            self.stud_ui.comboBox_stud_group.setCurrentIndex(
                self.stud_ui.comboBox_stud_group.findData(_stud[0][2]))

        _db = ARMDataBase()
        _sql = "SELECT * FROM students"
        students = _db.query(_sql)
        _db.close()

        _db = ARMDataBase()
        for i in range(len(students)):
            studs = []
            for h in students[i]:
                studs.append(h)

            if studs[2] is not None and studs[2] != '':
                _sql = "SELECT group_name FROM groups WHERE id_group=" + str(studs[2])
                stud_group = _db.query(_sql)
            else:
                stud_group = [['']]

            if studs[5] == "male":
                studs[5] = "Мужской"
            else:
                studs[5] = "Женский"

            studs[0] = 'clb_stud_' + str(studs[0])
            studs[1] = 'ФИО: ' + studs[1] + '\n' if studs[1] is not None and studs[1] != '' else ''
            studs[2] = 'Группа: ' + stud_group[0][0] + '\n' if stud_group[0][0] is not None and stud_group[0][
                0] != '' else ''
            studs[3] = 'День рождения: ' + studs[3] + '\n' if studs[3] is not None and studs[3] != '' else ''
            studs[4] = 'Телефон: ' + studs[4] + '\n' if studs[4] is not None and studs[4] != '' else ''
            studs[5] = 'Пол: ' + studs[5] + '\n' if studs[5] is not None and studs[5] != '' else ''
            studs[6] = 'Место проживания: ' + studs[6] + '\n' if studs[6] is not None and studs[6] != '' else ''
            studs[7] = 'Место обучения: ' + studs[7] + '\n' if studs[7] is not None and studs[7] != '' else ''
            studs[8] = 'Электронная почта: ' + studs[8] + '\n' if studs[8] is not None and studs[8] != '' else ''
            studs[9] = 'Социальные сети: ' + studs[9] + '\n' if studs[9] is not None and studs[9] != '' else ''
            searcher = ''
            if search_text is None or search_text == "":
                _search_text = self.stud_ui.lineEdit_search_stud.text().lower()
            elif search_text is not None:
                _search_text = search_text.lower()
            else:
                _search_text = search_text
            for h in studs:
                if h is not None and h != '':
                    searcher = searcher + h.lower()
            if _search_text is not None and _search_text != '':
                if _search_text in searcher:
                    stud_but = self.create_list_el(studs[0],
                                                   studs[1] + studs[2] + studs[3] + studs[4] + studs[5] + studs[6] +
                                                   studs[7] + studs[8] + studs[9],
                                                   self.stud_ui.sAWContent_stud_list)
                    stud_but.clicked.connect(lambda: loader_stud_edits())
            else:
                stud_but = self.create_list_el(studs[0],
                                               studs[1] + studs[2] + studs[3] + studs[4] + studs[5] + studs[6] + studs[
                                                   7] + studs[8] + studs[9],
                                               self.stud_ui.sAWContent_stud_list)
                stud_but.clicked.connect(lambda: loader_stud_edits())
        _db.close()

        _db = ARMDataBase()
        _sql = "SELECT * FROM groups"
        groups = _db.query(_sql)
        _db.close()
        _groups = []
        self.stud_ui.comboBox_stud_group.clear()
        for group in groups:
            self.create_combo_box_el(self.stud_ui.comboBox_stud_group, group[0], str(group[2]))

    # Loader database for enrollment
    def load_db_enrollment(self, search_text=None):
        clear_list(self.enr_ui.sAWContent_enr_list.children())

        def check_box_clicked():
            self.enr_ui.list_cb_check = []
            for check in self.enr_ui.groupBox_stud_subs.children():
                if "el_" in check.objectName():
                    self.enr_ui.list_cb_check.append([check.objectName().split('_')[-1], check.isChecked()])
            for o in range(len(self.enr_ui.list_cb_check)):
                if self.enr_ui.list_cb_check[o][1] is not self.enr_ui.list_cb_checked[o][1]:
                    for check in self.enr_ui.groupBox_stud_subs.children():
                        if "el_" + self.enr_ui.list_cb_check[o][0] in check.objectName():
                            if check.isChecked():
                                for child in self.enr_ui.groupBox_sis_contracts.children():
                                    if "grB_" in child.objectName() and \
                                            self.enr_ui.list_cb_check[o][0] in child.objectName():
                                        child.show()
                                        self.enr_ui.list_cb_checked[o][1] = True
                            else:
                                for child in self.enr_ui.groupBox_sis_contracts.children():
                                    if "grB_" in child.objectName() and \
                                            self.enr_ui.list_cb_check[o][0] in child.objectName():
                                        child.hide()
                                        self.enr_ui.list_cb_checked[o][1] = False

        def loader_enr_edits():
            clear_group_box(self.enr_ui.groupBox_stud_subs.children())
            clear_group_box(self.enr_ui.groupBox_sis_contracts.children())
            selected_enrollment = ''
            enrollment_list = self.enr_ui.sAWContent_enr_list.children()
            if len(enrollment_list) != 2:
                for enr in enrollment_list:
                    if enr.objectName() != 'vL_sAWContent_enr_list':
                        if enr.isChecked():
                            selected_enrollment = enr.objectName().split('_')[-1]
                            break
            else:
                selected_enrollment = self.enr_ui.sAWContent_enr_list.children()[-1].objectName().split('_')[-1]
            _db1 = ARMDataBase()
            _sql3 = "SELECT id_sub, status FROM subs_in_studs WHERE id_student=" + selected_enrollment
            list_active_subs = _db1.query(_sql3)
            _sql3 = "SELECT id_group FROM students WHERE id_student=" + selected_enrollment
            id_group_stud = _db1.query(_sql3)
            _sql3 = "SELECT id_prog FROM groups WHERE id_group=" + str(id_group_stud[0][0])
            id_prog_stud = _db1.query(_sql3)
            _sql3 = "SELECT id_sub, sub_name, id_teacher FROM subjects WHERE id_prog=" + str(id_prog_stud[0][0])
            list_subs = _db1.query(_sql3)
            _list_active_subs = []
            self.enr_ui.list_cb_checked = []
            for sub in range(len(list_active_subs)):
                if list_active_subs[sub][1] == "1":
                    _list_active_subs.append(str(list_active_subs[sub][0])[:])
            for sub in range(len(list_subs)):
                _sql2 = "SELECT student_numcontract, student_datecontract, status FROM subs_in_studs WHERE id_sub=" + \
                        str(list_subs[sub][0]) + " AND id_student=" + selected_enrollment
                contracts1 = _db1.query(_sql2)
                try:
                    _sql2 = "SELECT teacher_name FROM teachers WHERE id_teacher=" + str(list_subs[sub][2])
                    teacher_fam = _db1.query(_sql2)[0][0].split(" ")[0]
                except IndexError:
                    teacher_fam = ""
                cb = self.create_check_box_el(self.enr_ui.groupBox_stud_subs, 'el_' + str(list_subs[sub][0]),
                                              f"{list_subs[sub][1]} ({teacher_fam})",
                                              True if str(list_subs[sub][0]) in _list_active_subs else False)
                self.enr_ui.list_cb_checked.append(
                    [str(list_subs[sub][0]), True if str(list_subs[sub][0]) in _list_active_subs else False])
                cb.clicked.connect(lambda: check_box_clicked())
                gb = self.create_sub_groupbox(self.enr_ui.groupBox_sis_contracts,
                                              "grB_" + str(list_subs[sub][0]),
                                              f"{list_subs[sub][1]} ({teacher_fam})",
                                              False if str(list_subs[sub][0]) in _list_active_subs else True)
                for chield in gb.children():
                    if "gL_" not in chield.objectName() and contracts1 != []:
                        if "ledit_grB_" in chield.objectName():
                            chield.setText(contracts1[0][0])
                        elif "dedit_grB_" in chield.objectName():
                            chield.setDate(datetime.date(int(contracts1[0][1].split('.')[2]),
                                                         int(contracts1[0][1].split('.')[1]),
                                                         int(contracts1[0][1].split('.')[0])))
            _db1.close()

        _db = ARMDataBase()
        _sql = "SELECT * FROM students"
        students = _db.query(_sql)
        for i in range(len(students)):
            studs = []
            for h in students[i]:
                studs.append(h)

            if studs[2] is not None and studs[2] != '':
                _sql = "SELECT group_name FROM groups WHERE id_group=" + str(studs[2])
                stud_group = _db.query(_sql)
            else:
                stud_group = [['']]

            if stud_group != [['']] and stud_group is not None:
                _sql = "SELECT id_prog FROM groups WHERE id_group=" + str(studs[2])
                prog_id = _db.query(_sql)
                _sql = "SELECT prog_name FROM programs WHERE id_prog=" + str(prog_id[0][0])
                prog_name = _db.query(_sql)
                _sql = "SELECT prog_range FROM programs WHERE id_prog=" + str(prog_id[0][0])
                stud_group_prog_range = _db.query(_sql)
            else:
                prog_name = [['']]
                stud_group_prog_range = [['']]

            _sql = "SELECT id_sub FROM subs_in_studs WHERE id_student=" + str(studs[0])
            stud_subs = _db.query(_sql)
            subjects = ''
            for s in range(len(stud_subs)):
                _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(stud_subs[s][0])
                _sql1 = "SELECT student_numcontract, student_datecontract, status FROM subs_in_studs WHERE id_sub=" + str(
                    stud_subs[s][0]) + " AND id_student=" + str(studs[0])
                contracts = _db.query(_sql1)
                if s > 0:
                    if contracts[0][2] == "1":
                        subjects += ', ' + _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0], contracts[0][1])
                else:
                    if contracts[0][2] == "1":
                        subjects += _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0], contracts[0][1])

            studs[0] = 'clb_sis_' + str(studs[0])
            studs[1] = 'ФИО: ' + studs[1] + '\n' if studs[1] is not None and studs[1] != '' else ''
            studs[2] = 'Группа: ' + stud_group[0][0] + '\n' if stud_group[0][0] is not None and stud_group[0][
                0] != '' else ''
            studs[3] = 'Программа: ' + prog_name[0][0] + '\n' if prog_name[0][0] is not None and \
                       prog_name[0][0] != '' else ''
            studs[4] = 'Продолжительность обучения: ' + stud_group_prog_range[0][0] + ' месяцев\n' if \
                stud_group_prog_range[0][0] is not None and stud_group_prog_range[0][0] != '' else ''
            studs[5] = 'Предметы: ' + subjects + '\n' \
                if subjects is not None and subjects != '' else 'Предметы: Отсутствуют\n'

            searcher = ''
            if search_text is None or search_text == "":
                _search_text = self.enr_ui.lineEdit_search_enr.text().lower()
            elif search_text is not None:
                _search_text = search_text.lower()
            else:
                _search_text = search_text
            for h in studs:
                if h is not None and h != '':
                    searcher = searcher + h.lower()
            if _search_text is not None and _search_text != '':
                if _search_text in searcher:
                    stud_but = self.create_list_el(studs[0],
                                                   studs[1] + studs[2] + studs[3] + studs[4] + studs[5],
                                                   self.enr_ui.sAWContent_enr_list)
                    stud_but.clicked.connect(lambda: loader_enr_edits())
            else:
                stud_but = self.create_list_el(studs[0],
                                               studs[1] + studs[2] + studs[3] + studs[4] + studs[5],
                                               self.enr_ui.sAWContent_enr_list)
                stud_but.clicked.connect(lambda: loader_enr_edits())
        _db.close()

    # Loader database for enrollment in decree
    def load_db_decree_enr(self):
        def setup_prog_info():
            uncheck_clb()
            _db = ARMDataBase()
            selected_prog = self.decree_enr_ui.comboBox_prog.currentData()
            if selected_prog is not None:
                _sql = "SELECT id_group FROM groups WHERE id_prog=" + str(selected_prog)
                try:
                    groups = _db.query(_sql)[0][0]
                except IndexError:
                    groups = 1
                self.decree_enr_ui.comboBox_group.setCurrentIndex(self.decree_enr_ui.comboBox_group.findData(groups))
                _sql = "SELECT prog_range_dates FROM programs WHERE id_prog=" + \
                       str(self.decree_enr_ui.comboBox_prog.currentData())
                try:
                    dates = _db.query(_sql)[0][0].split("|")
                    self.decree_enr_ui.dateEdit_date_start.setDate(datetime.date(int(dates[0].split('.')[2]),
                                                                                 int(dates[0].split('.')[1]),
                                                                                 int(dates[0].split('.')[0])))
                    self.decree_enr_ui.dateEdit_date_end.setDate(datetime.date(int(dates[1].split('.')[2]),
                                                                               int(dates[1].split('.')[1]),
                                                                               int(dates[1].split('.')[0])))
                except AttributeError:
                    pass
            _db.close()

        def setup_group_info():
            uncheck_clb()
            clear_list(self.decree_enr_ui.sAWContent_outlay_studs.children())
            selected_group = self.decree_enr_ui.comboBox_group.currentData()
            _db = ARMDataBase()
            if selected_group is not None:
                _sql = "SELECT id_student, student_name FROM students WHERE id_group=" + str(selected_group)
                students = _db.query(_sql)
            else:
                students = []
            for i in range(len(students)):
                studs = []
                for h in students[i]:
                    studs.append(h)
                _sql = "SELECT id_sub FROM subs_in_studs WHERE id_student=" + str(studs[0])
                stud_subs = _db.query(_sql)
                subjects = ''
                for s in range(len(stud_subs)):
                    _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(stud_subs[s][0])
                    _sql1 = "SELECT student_numcontract, student_datecontract, status FROM subs_in_studs WHERE id_sub=" + str(
                        stud_subs[s][0]) + " AND id_student=" + str(studs[0])
                    contracts = _db.query(_sql1)
                    if s > 0:
                        if contracts[0][2] == "1":
                            subjects += ', ' + _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0],
                                                                                         contracts[0][1])
                    else:
                        if contracts[0][2] == "1":
                            subjects += _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0], contracts[0][1])
                if subjects != "":
                    studs[0] = 'clb_decree_enr_' + str(studs[0])
                    studs[1] = 'ФИО: ' + studs[1] + '\n' if studs[1] is not None and studs[1] != '' else ''
                    studs.append('Предметы: ' + subjects)
                    stud_but = self.create_list_el(studs[0],
                                                   studs[1] + studs[2],
                                                   self.decree_enr_ui.sAWContent_outlay_studs,
                                                   auto_exclusive=False)
                    stud_but.clicked.connect(lambda: uncheck_clb())
            _db.close()

        def check_all_group():
            for clb in self.decree_enr_ui.sAWContent_outlay_studs.children():
                if clb.objectName().startswith("clb_"):
                    clb.setChecked(True) if self.decree_enr_ui.checkBox_all_group.isChecked() else clb.setChecked(False)

        def uncheck_clb():
            self.decree_enr_ui.checkBox_all_group.setChecked(False)

        self.decree_enr_ui.comboBox_head.clear()
        self.decree_enr_ui.comboBox_prog.clear()
        self.decree_enr_ui.comboBox_pfs.clear()
        self.decree_enr_ui.comboBox_office.clear()
        self.decree_enr_ui.comboBox_manager_cpui.clear()
        self.decree_enr_ui.comboBox_group.clear()
        clear_list(self.decree_enr_ui.sAWContent_outlay_studs.children())
        _db = ARMDataBase()

        self.decree_enr_ui.checkBox_all_group.clicked.connect(lambda: check_all_group())

        # comboBox for Head load
        _sql = "SELECT id_head, head_name, head_prof FROM headers"
        headers = _db.query(_sql)
        for head in headers:
            self.create_combo_box_el(self.decree_enr_ui.comboBox_head, head[0], str(head[2]) + " | " + str(head[1]))
            if "директор хти" in str(head[2]).lower():
                self.decree_enr_ui.comboBox_head.setCurrentIndex(
                    self.decree_enr_ui.comboBox_head.findData(head[0]))

        # comboBox for Programs load
        _sql = "SELECT id_prog, prog_name, prog_range_dates FROM programs"
        programs = _db.query(_sql)
        for prog in programs:
            self.create_combo_box_el(self.decree_enr_ui.comboBox_prog, prog[0], str(prog[1]))
        self.decree_enr_ui.comboBox_prog.currentIndexChanged.connect(lambda: setup_prog_info())

        # comboBox for Groups load
        _sql = "SELECT id_group, group_name FROM groups"
        groups = _db.query(_sql)
        for group in groups:
            self.create_combo_box_el(self.decree_enr_ui.comboBox_group, group[0], str(group[1]))
        self.decree_enr_ui.comboBox_group.currentIndexChanged.connect(lambda: setup_group_info())

        # comboBox for PFS load
        for head in headers:
            self.create_combo_box_el(self.decree_enr_ui.comboBox_pfs, head[0], str(head[2]) + " | " + str(head[1]))
            if "пфс" in str(head[2]).lower():
                self.decree_enr_ui.comboBox_pfs.setCurrentIndex(
                    self.decree_enr_ui.comboBox_pfs.findData(head[0]))

        # comboBox for head of office load
        for head in headers:
            self.create_combo_box_el(self.decree_enr_ui.comboBox_office, head[0],
                                     str(head[2]) + " | " + str(head[1]))
            if "канцеляр" in str(head[2]).lower():
                self.decree_enr_ui.comboBox_office.setCurrentIndex(
                    self.decree_enr_ui.comboBox_office.findData(head[0]))

        # comboBox for manager CPUI load
        for head in headers:
            self.create_combo_box_el(self.decree_enr_ui.comboBox_manager_cpui, head[0],
                                     str(head[2]) + " | " + str(head[1]))
            if "цпюи" in str(head[2]).lower() and 'зав' in str(head[2]).lower():
                self.decree_enr_ui.comboBox_manager_cpui.setCurrentIndex(
                    self.decree_enr_ui.comboBox_manager_cpui.findData(head[0]))

        _db.close()

    # Loader database for passes in note
    def load_db_note_passes(self):
        def setup_prog_info():
            uncheck_clb()
            _db = ARMDataBase()
            selected_prog = self.note_passes_ui.comboBox_prog.currentData()
            if selected_prog is not None:
                _sql = "SELECT id_group FROM groups WHERE id_prog=" + str(selected_prog)
                try:
                    groups = _db.query(_sql)[0][0]
                except IndexError:
                    groups = 1
                self.note_passes_ui.comboBox_group.setCurrentIndex(self.note_passes_ui.comboBox_group.findData(groups))
                _sql = "SELECT prog_range_dates FROM programs WHERE id_prog=" + \
                       str(self.note_passes_ui.comboBox_prog.currentData())
                try:
                    dates = _db.query(_sql)[0][0].split("|")
                    self.note_passes_ui.dateEdit_date_start.setDate(datetime.date(int(dates[0].split('.')[2]),
                                                                                 int(dates[0].split('.')[1]),
                                                                                 int(dates[0].split('.')[0])))
                    self.note_passes_ui.dateEdit_date_end.setDate(datetime.date(int(dates[1].split('.')[2]),
                                                                               int(dates[1].split('.')[1]),
                                                                               int(dates[1].split('.')[0])))
                except AttributeError:
                    pass
            _db.close()

        def setup_group_info():
            uncheck_clb()
            clear_list(self.note_passes_ui.sAWContent_outlay_studs.children())
            selected_group = self.note_passes_ui.comboBox_group.currentData()
            _db = ARMDataBase()
            if selected_group is not None:
                _sql = "SELECT id_student, student_name FROM students WHERE id_group=" + str(selected_group)
                students = _db.query(_sql)
            else:
                students = []
            for i in range(len(students)):
                studs = []
                for h in students[i]:
                    studs.append(h)
                _sql = "SELECT id_sub FROM subs_in_studs WHERE id_student=" + str(studs[0])
                stud_subs = _db.query(_sql)
                subjects = ''
                for s in range(len(stud_subs)):
                    _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(stud_subs[s][0])
                    _sql1 = "SELECT student_numcontract, student_datecontract, status FROM subs_in_studs WHERE id_sub=" + str(
                        stud_subs[s][0]) + " AND id_student=" + str(studs[0])
                    contracts = _db.query(_sql1)
                    if s > 0:
                        if contracts[0][2] == "1":
                            subjects += ', ' + _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0],
                                                                                         contracts[0][1])
                    else:
                        if contracts[0][2] == "1":
                            subjects += _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0], contracts[0][1])
                if subjects != "":
                    studs[0] = 'clb_decree_enr_' + str(studs[0])
                    studs[1] = 'ФИО: ' + studs[1] + '\n' if studs[1] is not None and studs[1] != '' else ''
                    studs.append('Предметы: ' + subjects)
                    stud_but = self.create_list_el(studs[0],
                                                   studs[1] + studs[2],
                                                   self.note_passes_ui.sAWContent_outlay_studs,
                                                   auto_exclusive=False)
                    stud_but.clicked.connect(lambda: uncheck_clb())
            _db.close()

        def check_all_group():
            for clb in self.note_passes_ui.sAWContent_outlay_studs.children():
                if clb.objectName().startswith("clb_"):
                    clb.setChecked(True) if self.note_passes_ui.checkBox_all_group.isChecked() else clb.setChecked(False)

        def uncheck_clb():
            self.note_passes_ui.checkBox_all_group.setChecked(False)

        self.note_passes_ui.comboBox_head.clear()
        self.note_passes_ui.comboBox_prog.clear()
        self.note_passes_ui.comboBox_manager_cpui.clear()
        self.note_passes_ui.comboBox_group.clear()
        clear_list(self.note_passes_ui.sAWContent_outlay_studs.children())
        _db = ARMDataBase()

        self.note_passes_ui.checkBox_all_group.clicked.connect(lambda: check_all_group())

        # comboBox for Head load
        _sql = "SELECT id_head, head_name, head_prof FROM headers"
        headers = _db.query(_sql)
        for head in headers:
            self.create_combo_box_el(self.note_passes_ui.comboBox_head, head[0], str(head[2]) + " | " + str(head[1]))
            if "директор хти" in str(head[2]).lower():
                self.note_passes_ui.comboBox_head.setCurrentIndex(
                    self.note_passes_ui.comboBox_head.findData(head[0]))

        # comboBox for Programs load
        _sql = "SELECT id_prog, prog_name, prog_range_dates FROM programs"
        programs = _db.query(_sql)
        for prog in programs:
            self.create_combo_box_el(self.note_passes_ui.comboBox_prog, prog[0], str(prog[1]))
        self.note_passes_ui.comboBox_prog.currentIndexChanged.connect(lambda: setup_prog_info())

        # comboBox for Groups load
        _sql = "SELECT id_group, group_name FROM groups"
        groups = _db.query(_sql)
        for group in groups:
            self.create_combo_box_el(self.note_passes_ui.comboBox_group, group[0], str(group[1]))
        self.note_passes_ui.comboBox_group.currentIndexChanged.connect(lambda: setup_group_info())

        # comboBox for manager CPUI load
        for head in headers:
            self.create_combo_box_el(self.note_passes_ui.comboBox_manager_cpui, head[0],
                                     str(head[2]) + " | " + str(head[1]))
            if "цпюи" in str(head[2]).lower() and 'зав' in str(head[2]).lower():
                self.note_passes_ui.comboBox_manager_cpui.setCurrentIndex(
                    self.note_passes_ui.comboBox_manager_cpui.findData(head[0]))

        # dateEdit for date note
        self.note_passes_ui.dateEdit_date.setDate(datetime.datetime.today())

        _db.close()

    # Loader database for passwords in note
    def load_db_note_passwords(self):
        def setup_prog_info():
            uncheck_clb()
            _db = ARMDataBase()
            selected_prog = self.note_passwords_ui.comboBox_prog.currentData()
            if selected_prog is not None:
                _sql = "SELECT id_group FROM groups WHERE id_prog=" + str(selected_prog)
                try:
                    groups = _db.query(_sql)[0][0]
                except IndexError:
                    groups = 1
                self.note_passwords_ui.comboBox_group.setCurrentIndex(self.note_passwords_ui.comboBox_group.findData(groups))
                _sql = "SELECT prog_range_dates FROM programs WHERE id_prog=" + \
                       str(self.note_passwords_ui.comboBox_prog.currentData())
                try:
                    dates = _db.query(_sql)[0][0].split("|")
                    self.note_passwords_ui.dateEdit_date_start.setDate(datetime.date(int(dates[0].split('.')[2]),
                                                                                 int(dates[0].split('.')[1]),
                                                                                 int(dates[0].split('.')[0])))
                    self.note_passwords_ui.dateEdit_date_end.setDate(datetime.date(int(dates[1].split('.')[2]),
                                                                               int(dates[1].split('.')[1]),
                                                                               int(dates[1].split('.')[0])))
                except AttributeError:
                    pass
            _db.close()

        def setup_group_info():
            uncheck_clb()
            clear_list(self.note_passwords_ui.sAWContent_outlay_studs.children())
            selected_group = self.note_passwords_ui.comboBox_group.currentData()
            _db = ARMDataBase()
            if selected_group is not None:
                _sql = "SELECT id_student, student_name FROM students WHERE id_group=" + str(selected_group)
                students = _db.query(_sql)
            else:
                students = []
            for i in range(len(students)):
                studs = []
                for h in students[i]:
                    studs.append(h)
                _sql = "SELECT id_sub FROM subs_in_studs WHERE id_student=" + str(studs[0])
                stud_subs = _db.query(_sql)
                subjects = ''
                for s in range(len(stud_subs)):
                    _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(stud_subs[s][0])
                    _sql1 = "SELECT student_numcontract, student_datecontract, status FROM subs_in_studs WHERE id_sub=" + str(
                        stud_subs[s][0]) + " AND id_student=" + str(studs[0])
                    contracts = _db.query(_sql1)
                    if s > 0:
                        if contracts[0][2] == "1":
                            subjects += ', ' + _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0],
                                                                                         contracts[0][1])
                    else:
                        if contracts[0][2] == "1":
                            subjects += _db.query(_sql)[0][0] + '({}, {})'.format(contracts[0][0], contracts[0][1])
                if subjects != "":
                    studs[0] = 'clb_decree_enr_' + str(studs[0])
                    studs[1] = 'ФИО: ' + studs[1] + '\n' if studs[1] is not None and studs[1] != '' else ''
                    studs.append('Предметы: ' + subjects)
                    stud_but = self.create_list_el(studs[0],
                                                   studs[1] + studs[2],
                                                   self.note_passwords_ui.sAWContent_outlay_studs,
                                                   auto_exclusive=False)
                    stud_but.clicked.connect(lambda: uncheck_clb())
            _db.close()

        def check_all_group():
            for clb in self.note_passwords_ui.sAWContent_outlay_studs.children():
                if clb.objectName().startswith("clb_"):
                    clb.setChecked(True) if self.note_passwords_ui.checkBox_all_group.isChecked() else clb.setChecked(False)

        def uncheck_clb():
            self.note_passwords_ui.checkBox_all_group.setChecked(False)

        self.note_passwords_ui.comboBox_head.clear()
        self.note_passwords_ui.comboBox_prog.clear()
        self.note_passwords_ui.comboBox_manager_cpui.clear()
        self.note_passwords_ui.comboBox_group.clear()
        clear_list(self.note_passwords_ui.sAWContent_outlay_studs.children())
        _db = ARMDataBase()

        self.note_passwords_ui.checkBox_all_group.clicked.connect(lambda: check_all_group())

        # comboBox for Head load
        _sql = "SELECT id_head, head_name, head_prof FROM headers"
        headers = _db.query(_sql)
        for head in headers:
            self.create_combo_box_el(self.note_passwords_ui.comboBox_head, head[0], str(head[2]) + " | " + str(head[1]))
            if "информац" in str(head[2]).lower() \
            and 'технол' in str(head[2]).lower() \
            and 'департа' in str(head[2]).lower() \
            and 'руковод' in str(head[2]).lower():
                self.note_passwords_ui.comboBox_head.setCurrentIndex(
                    self.note_passwords_ui.comboBox_head.findData(head[0]))

        # comboBox for Programs load
        _sql = "SELECT id_prog, prog_name, prog_range_dates FROM programs"
        programs = _db.query(_sql)
        for prog in programs:
            self.create_combo_box_el(self.note_passwords_ui.comboBox_prog, prog[0], str(prog[1]))
        self.note_passwords_ui.comboBox_prog.currentIndexChanged.connect(lambda: setup_prog_info())

        # comboBox for Groups load
        _sql = "SELECT id_group, group_name FROM groups"
        groups = _db.query(_sql)
        for group in groups:
            self.create_combo_box_el(self.note_passwords_ui.comboBox_group, group[0], str(group[1]))
        self.note_passwords_ui.comboBox_group.currentIndexChanged.connect(lambda: setup_group_info())

        # comboBox for manager CPUI load
        for head in headers:
            self.create_combo_box_el(self.note_passwords_ui.comboBox_manager_cpui, head[0],
                                     str(head[2]) + " | " + str(head[1]))
            if "цпюи" in str(head[2]).lower() and 'зав' in str(head[2]).lower():
                self.note_passwords_ui.comboBox_manager_cpui.setCurrentIndex(
                    self.note_passwords_ui.comboBox_manager_cpui.findData(head[0]))

        # dateEdit for date note
        self.note_passwords_ui.dateEdit_date_before.setDate(datetime.datetime.today())
        self.note_passwords_ui.dateEdit_date.setDate(datetime.datetime.today())

        _db.close()

    # Loader database for list studs in note
    def load_db_note_list(self):
        clear_list(self.note_list_ui.sAWContent_groups.children())
        _db = ARMDataBase()

        # list for Groups load
        _sql = "SELECT id_group, group_name FROM groups"
        groups = _db.query(_sql)
        for group in groups:
            if str(group[0]) != '1':
                self.create_list_el('clb_group_' + str(group[0]), str(group[1]), self.note_list_ui.sAWContent_groups, auto_exclusive=False)

        _db.close()

    # Loader database for contract docxs
    def load_db_contract(self):
        def setup_prog_info():
            clear_list(self.contract_ui.sAWContent_subs.children())
            _db = ARMDataBase()
            selected_prog = self.contract_ui.comboBox_prog.currentData()
            if selected_prog is not None and selected_prog != 8:
                _sql = "SELECT prog_range_dates FROM programs WHERE id_prog=" + \
                       str(self.contract_ui.comboBox_prog.currentData())
                dates = _db.query(_sql)[0][0].split("|")
                self.contract_ui.dateEdit_date_start.setDate(datetime.date(int(dates[0].split('.')[2]),
                                                                             int(dates[0].split('.')[1]),
                                                                             int(dates[0].split('.')[0])))
                self.contract_ui.dateEdit_date_end.setDate(datetime.date(int(dates[1].split('.')[2]),
                                                                           int(dates[1].split('.')[1]),
                                                                           int(dates[1].split('.')[0])))

                if str(self.contract_ui.comboBox_prog.currentData()) != "8":
                    _sql = "SELECT id_sub, sub_name FROM subjects WHERE id_prog=" + \
                           str(self.contract_ui.comboBox_prog.currentData())
                    subs = _db.query(_sql)
                    for sub in subs:
                        _sql = "SELECT id_teacher FROM subjects WHERE id_sub=" + str(sub[0])
                        id_teacher = _db.query(_sql)[0][0]
                        _sql = "SELECT teacher_name FROM teachers WHERE id_teacher=" + str(id_teacher)
                        teacher = _db.query(_sql)[0][0]
                        self.create_list_el(
                            'clb_sub_' + str(sub[0]),
                            f"{str(sub[1])} ({teacher})",
                            self.contract_ui.sAWContent_subs,
                            auto_exclusive=False
                        )

            _db.close()

        _db = ARMDataBase()

        clear_list(self.contract_ui.sAWContent_subs.children())
        self.contract_ui.comboBox_prog.clear()
        self.contract_ui.comboBox_manager_cpui.clear()
        self.contract_ui.comboBox_head.clear()
        self.contract_ui.comboBox_head_ls.clear()

        # comboBox for Head load
        _sql = "SELECT id_head, head_name, head_prof FROM headers"
        headers = _db.query(_sql)
        for head in headers:
            self.create_combo_box_el(self.contract_ui.comboBox_head, head[0], str(head[2]) + " | " + str(head[1]))
            if "директор хти" in str(head[2]).lower():
                self.contract_ui.comboBox_head.setCurrentIndex(
                    self.contract_ui.comboBox_head.findData(head[0]))

        # comboBox for Programs load
        _sql = "SELECT id_prog, prog_name, prog_range_dates FROM programs"
        programs = _db.query(_sql)
        for prog in programs:
            self.create_combo_box_el(self.contract_ui.comboBox_prog, prog[0], str(prog[1]))
        self.contract_ui.comboBox_prog.currentIndexChanged.connect(lambda: setup_prog_info())

        # comboBox for manager CPUI load
        for head in headers:
            self.create_combo_box_el(self.contract_ui.comboBox_manager_cpui, head[0],
                                     str(head[2]) + " | " + str(head[1]))
            if "цпюи" in str(head[2]).lower() and 'зав' in str(head[2]).lower():
                self.contract_ui.comboBox_manager_cpui.setCurrentIndex(
                    self.contract_ui.comboBox_manager_cpui.findData(head[0]))

        # comboBox for manager Legal Sector load
        for head in headers:
            self.create_combo_box_el(self.contract_ui.comboBox_head_ls, head[0],
                                     str(head[2]) + " | " + str(head[1]))
            if "правов" in str(head[2]).lower() and 'зав' in str(head[2]).lower() and 'сектор' in str(head[2]).lower():
                self.contract_ui.comboBox_head_ls.setCurrentIndex(
                    self.contract_ui.comboBox_head_ls.findData(head[0]))

        # list for Subjects load
        if str(self.contract_ui.comboBox_prog.currentData()) != "8":
            _sql = "SELECT id_sub, sub_name FROM subjects WHERE id_prog=" + \
                   str(self.contract_ui.comboBox_prog.currentData())
            subs = _db.query(_sql)
            for sub in subs:
                _sql = "SELECT id_teacher FROM subjects WHERE id_sub=" + str(sub[0])
                id_teacher = _db.query(_sql)[0][0]
                _sql = "SELECT teacher_name FROM teachers WHERE id_teacher=" + str(id_teacher)
                teacher = _db.query(_sql)[0][0].split()[0]
                self.create_list_el(
                    'clb_sub_' + str(sub[0]),
                    f"{str(sub[1])} ({teacher})",
                    self.contract_ui.sAWContent_subs,
                    auto_exclusive=False
                )

        _db.close()

    # Loader database for Outlay Calculator
    def load_db_outlay(self):
        _db = ARMDataBase()
        _sql = "SELECT * FROM programs"
        programs = _db.query(_sql)
        _db.close()

        self.outlay_ui.comboBox_progs.clear()
        self.create_combo_box_el(self.outlay_ui.comboBox_progs, "None", "Отсутствует")
        for prog in programs:
            if str(prog[0]) != "8":
                self.create_combo_box_el(self.outlay_ui.comboBox_progs, prog[0], str(prog[1]))
        clear_widget(self.outlay_ui.widget_calcs.children())

        self.outlay_ui.radio_col_1.setChecked(True)
        self.add_calculate_box(False, "1", 0, 0)
        # self.add_calculate_box(False, "2", 1, 0)
        # self.add_calculate_box(False, "3", 0, 1)
        # self.add_calculate_box(False, "4", 1, 1)

    # Loader database for Timetable list
    def load_db_timetable(self, search_text=None):
        clear_list(self.ui.sAWContent_timetable.children())
        _db = ARMDataBase()
        _sql = "SELECT id_sub, sub_name, id_teacher, id_prog, sub_hours, sub_hours_need FROM subjects"
        timetable_info = _db.query(_sql)
        for i in range(len(timetable_info)):
            text = "Предмет: {}\n".format(timetable_info[i][1])
            _sql = "SELECT prog_name, prog_range FROM programs WHERE id_prog=" + str(timetable_info[i][3])
            prog_info = _db.query(_sql)
            text += "Программа: {}\n".format(prog_info[0][0])
            text += "Продолжительность программы: в течении {} месяцев\n".format(prog_info[0][1])
            _sql = "SELECT teacher_name FROM teachers WHERE id_teacher=" + str(timetable_info[i][2])
            teacher_info = _db.query(_sql)
            text += "Преподаватель: {}\n".format(teacher_info[0][0])
            text += "Необходимо часов: {}\n".format(timetable_info[i][5])
            text += "Часы: {}\n".format(timetable_info[i][4])
            _search_text = search_text
            searcher = text.lower()
            if _search_text is not None and _search_text != '':
                _search_text = search_text.lower()
                if _search_text in searcher:
                    self.create_list_el("clb_ttible_" + str(timetable_info[i][0]),
                                        text,
                                        self.ui.sAWContent_timetable)
            else:
                self.create_list_el("clb_ttible_" + str(timetable_info[i][0]),
                                    text,
                                    self.ui.sAWContent_timetable)
        _db.close()

    # Loader database for exe_Timetable list
    def load_db_timetable_list(self):
        def setup_date():
            for date_ in self.ttable_ui.sAWContent_hours_list.children():
                if date_.objectName().startswith("clb_"):
                    if date_.isChecked():
                        self.ttable_ui.calendar.setSelectedDate(
                            datetime.datetime.fromtimestamp(float(date_.objectName().split("_")[-1])))
                        break

        self.ttable_list = []
        clear_list(self.ttable_ui.sAWContent_hours_list.children())
        _db = ARMDataBase()
        _sql = "SELECT sub_ttable, sub_hours, sub_hours_need FROM subjects WHERE id_sub=" + self.ttable_selected_sub
        timetable_sub = _db.query(_sql)
        self.ttable_ui.lab_sum.setText("Сумма часов: " + timetable_sub[0][1])
        self.ttable_ui.lab_need.setText("Необходимо часов: " + timetable_sub[0][2])
        parse_timetable = []
        if timetable_sub[0][0] is not None and timetable_sub[0][0] != '':
            for date in timetable_sub[0][0].split(","):
                parse_timetable.append(date)
            for i in range(len(parse_timetable)):
                parse_timetable[i] = parse_timetable[i].split("|")
                parse_timetable[i][0] = datetime.datetime.strptime(parse_timetable[i][0], "%d.%m.%Y")
                if parse_timetable[i][0].weekday() == 0:
                    parse_timetable[i].append('Понедельник')
                elif parse_timetable[i][0].weekday() == 1:
                    parse_timetable[i].append('Вторник')
                elif parse_timetable[i][0].weekday() == 2:
                    parse_timetable[i].append('Среда')
                elif parse_timetable[i][0].weekday() == 3:
                    parse_timetable[i].append('Четверг')
                elif parse_timetable[i][0].weekday() == 4:
                    parse_timetable[i].append('Пятница')
                elif parse_timetable[i][0].weekday() == 5:
                    parse_timetable[i].append('Суббота')
                elif parse_timetable[i][0].weekday() == 6:
                    parse_timetable[i].append('Воскресенье')

                if parse_timetable[i][0].strftime("%m") == "01":
                    parse_timetable[i].append('Января')
                elif parse_timetable[i][0].strftime("%m") == "02":
                    parse_timetable[i].append('Февраля')
                elif parse_timetable[i][0].strftime("%m") == "03":
                    parse_timetable[i].append('Марта')
                elif parse_timetable[i][0].strftime("%m") == "04":
                    parse_timetable[i].append('Апреля')
                elif parse_timetable[i][0].strftime("%m") == "05":
                    parse_timetable[i].append('Мая')
                elif parse_timetable[i][0].strftime("%m") == "06":
                    parse_timetable[i].append('Июня')
                elif parse_timetable[i][0].strftime("%m") == "07":
                    parse_timetable[i].append('Июля')
                elif parse_timetable[i][0].strftime("%m") == "08":
                    parse_timetable[i].append('Августа')
                elif parse_timetable[i][0].strftime("%m") == "09":
                    parse_timetable[i].append('Сентября')
                elif parse_timetable[i][0].strftime("%m") == "10":
                    parse_timetable[i].append('Октября')
                elif parse_timetable[i][0].strftime("%m") == "11":
                    parse_timetable[i].append('Ноября')
                elif parse_timetable[i][0].strftime("%m") == "12":
                    parse_timetable[i].append('Декабря')
                btn = self.create_list_el(
                    "clb_" + str(parse_timetable[i][0].timestamp()),
                    parse_timetable[i][0].strftime("%d") + " " +
                    parse_timetable[i][3] + ", " +
                    parse_timetable[i][2] + ", " +
                    parse_timetable[i][0].strftime("%Y") + ", " +
                    parse_timetable[i][1] + " часов",
                    self.ttable_ui.sAWContent_hours_list
                )
                btn.clicked.connect(lambda: setup_date())
            self.ttable_list = copy.deepcopy(parse_timetable)

    def del_hours(self):
        del_hour = datetime.datetime.strptime(self.ttable_ui.calendar.selectedDate().toString('dd.MM.yyyy'),
                                              "%d.%m.%Y")
        parsed_hours = ''
        sum_hours = 0

        for _date in self.ttable_list:
            if _date[0] == del_hour:
                self.ttable_list.remove(_date)

        for hour in self.ttable_list:
            parsed_hours += hour[0].strftime("%d.%m.%Y") + "|" + hour[1] + ','
            sum_hours += int(hour[1])
        parsed_hours = parsed_hours[:-1]
        _db = ARMDataBase()
        _sql = "UPDATE subjects SET " \
               "sub_ttable = '{0}', " \
               "sub_hours = '{1}' " \
               "WHERE id_sub = '{2}'".format(parsed_hours,
                                             str(sum_hours),
                                             self.ttable_selected_sub)
        _db.query(_sql)
        _db.close()
        self.load_db_timetable_list()

    def set_hours(self):
        parse_hour = datetime.datetime.strptime(self.ttable_ui.calendar.selectedDate().toString('dd.MM.yyyy'),
                                                "%d.%m.%Y")
        search_status = 0
        sum_hours = 0
        parsed_hours = ''
        for hour in self.ttable_list:
            if hour[0] == parse_hour:
                hour[1] = self.ttable_ui.lEdit_hours.text()
                search_status = 1
                break
            else:
                search_status = 0
        if not search_status:
            self.ttable_list.append([parse_hour, self.ttable_ui.lEdit_hours.text()])
            self.ttable_list = sorted(
                self.ttable_list,
                key=lambda x: x[0], reverse=False
            )
        for hour in self.ttable_list:
            parsed_hours += hour[0].strftime("%d.%m.%Y") + "|" + hour[1] + ','
            sum_hours += int(hour[1])
        parsed_hours = parsed_hours[:-1]
        _db = ARMDataBase()
        _sql = "UPDATE subjects SET " \
               "sub_ttable = '{0}', " \
               "sub_hours = '{1}' " \
               "WHERE id_sub = '{2}'".format(parsed_hours,
                                             str(sum_hours),
                                             self.ttable_selected_sub)
        _db.query(_sql)
        _db.close()
        self.load_db_timetable_list()

    def timetable_list_exec(self):
        _set_doc_warning = 1
        for i in self.ui.sAWContent_timetable.children():
            if not i.objectName().startswith("clb_ttible_"):
                pass
            else:
                if i.isChecked():
                    self.ttable_selected_sub = i.objectName().split('_')[-1]
                    _set_doc_warning = 0
                    break
                else:
                    _set_doc_warning = 1
        if _set_doc_warning:
            set_doc_warning("Ошибка (не выбрана запись для изменения)",
                            'Сначала выберите запись для изменения.\n\nНажмите на нужную запись, '
                            'чтобы выбрать ее, а потом нажмите на кнопку '
                            '"Изменить расписание"')
        else:
            self.load_db_timetable_list()
            self.ttable.exec_()

    # Loader database for exe_OutlayPrinter
    def load_db_outlay_printer(self):
        def setup_prog_info():
            # lEdit for class load
            _db1 = ARMDataBase()
            try:
                _sql1 = "SELECT class FROM groups WHERE id_prog=" + str(self.outlay_print_ui.comboBox_prog.currentData())
                group_class = _db1.query(_sql1)
                self.outlay_print_ui.lEdit_class.setText(group_class[0][0])
            except Exception:
                self.outlay_print_ui.lEdit_class.setText("10")
            try:
                _sql1 = "SELECT prog_range_dates FROM programs WHERE id_prog=" + str(
                    self.outlay_print_ui.comboBox_prog.currentData())
                dates = _db1.query(_sql1)[0][0].split("|")
                self.outlay_print_ui.dateEdit_date_start.setDate(datetime.date(int(dates[0].split('.')[2]),
                                                                               int(dates[0].split('.')[1]),
                                                                               int(dates[0].split('.')[0])))
                self.outlay_print_ui.dateEdit_date_end.setDate(datetime.date(int(dates[1].split('.')[2]),
                                                                             int(dates[1].split('.')[1]),
                                                                             int(dates[1].split('.')[0])))
            except Exception:
                pass
            _db1.close()

        if self.outlay_ui.comboBox_progs.currentData() == "None" or self.outlay_ui.comboBox_progs.currentData() is None:
            selected = False
        else:
            selected = True

        self.outlay_print_ui.comboBox_head.clear()
        self.outlay_print_ui.comboBox_prog.clear()
        self.outlay_print_ui.lEdit_class.setText("")
        self.outlay_print_ui.comboBox_pfs.clear()
        self.outlay_print_ui.comboBox_bookkeeper.clear()
        self.outlay_print_ui.comboBox_manager_cpui.clear()
        clear_widget(self.outlay_print_ui.widget_subs_teachs.children())
        _db = ARMDataBase()

        # comboBox for Head load
        _sql = "SELECT id_head, head_name, head_prof FROM headers"
        headers = _db.query(_sql)
        for head in headers:
            self.create_combo_box_el(self.outlay_print_ui.comboBox_head, head[0], str(head[2]) + " | " + str(head[1]))
            if "директор хти" in str(head[2]).lower():
                self.outlay_print_ui.comboBox_head.setCurrentIndex(
                    self.outlay_print_ui.comboBox_head.findData(head[0]))

        # DateEdit for Date Confirm
        self.outlay_print_ui.dateEdit_date_confirm.setDate(datetime.datetime.today())

        # comboBox for Programs load
        _sql = "SELECT id_prog, prog_name, prog_range_dates FROM programs"
        programs = _db.query(_sql)
        for prog in programs:
            self.create_combo_box_el(self.outlay_print_ui.comboBox_prog, prog[0], str(prog[1]))
        if selected:
            self.outlay_print_ui.comboBox_prog.setCurrentIndex(
                self.outlay_print_ui.comboBox_prog.findData(
                    self.outlay_ui.comboBox_progs.currentData()
                ))
        self.outlay_print_ui.comboBox_prog.currentIndexChanged.connect(lambda: setup_prog_info())
        setup_prog_info()

        # comboBox for PFS load
        for head in headers:
            self.create_combo_box_el(self.outlay_print_ui.comboBox_pfs, head[0], str(head[2]) + " | " + str(head[1]))
            if "пфс" in str(head[2]).lower():
                self.outlay_print_ui.comboBox_pfs.setCurrentIndex(
                    self.outlay_print_ui.comboBox_pfs.findData(head[0]))

        # comboBox for Bookkeeper load
        for head in headers:
            self.create_combo_box_el(self.outlay_print_ui.comboBox_bookkeeper, head[0], str(head[2]) + " | " + str(head[1]))
            if "бухгалтер" in str(head[2]).lower():
                self.outlay_print_ui.comboBox_bookkeeper.setCurrentIndex(
                    self.outlay_print_ui.comboBox_bookkeeper.findData(head[0]))

        # comboBox for manager CPUI load
        for head in headers:
            self.create_combo_box_el(self.outlay_print_ui.comboBox_manager_cpui, head[0], str(head[2]) + " | " + str(head[1]))
            if "цпюи" in str(head[2]).lower() and 'зав' in str(head[2]).lower():
                self.outlay_print_ui.comboBox_manager_cpui.setCurrentIndex(
                    self.outlay_print_ui.comboBox_manager_cpui.findData(head[0]))

        # Widgets for subs and teachers load
        if not selected:
            rb = 0
            if self.outlay_ui.radio_col_1.isChecked():
                rb = 1
            elif self.outlay_ui.radio_col_2.isChecked():
                rb = 2
            elif self.outlay_ui.radio_col_3.isChecked():
                rb = 3
            elif self.outlay_ui.radio_col_4.isChecked():
                rb = 4
            for r in range(rb):
                if r + 1 == 1:
                    self.add_sub_teach_box(False, str(r + 1), 0, 0)
                elif r + 1 == 2:
                    self.add_sub_teach_box(False, str(r + 1), 1, 0)
                elif r + 1 == 3:
                    self.add_sub_teach_box(False, str(r + 1), 0, 1)
                elif r + 1 == 4:
                    self.add_sub_teach_box(False, str(r + 1), 1, 1)
        else:
            _sql = "SELECT id_sub FROM subjects WHERE id_prog=" + str(self.outlay_ui.comboBox_progs.currentData())
            subs = _db.query(_sql)
            a = 0
            b = 0
            for i in subs:
                self.add_sub_teach_box(True, str(i[0]), a, b)
                if a == 0 and b == 0:
                    a = 1
                    b = 0
                elif a == 1 and b == 0:
                    a = 0
                    b = 1
                elif a == 0 and b == 1:
                    a = 1
                    b = 1

        _db.close()

    def outlay_printer_exec(self):
        self.load_db_outlay_printer()
        self.outlay_printer.exec_()

    def decree_creator_exec(self):
        self.docx_ui.widget_enrollment.hide()
        self.docx_ui.widget_notepasses.hide()
        self.docx_ui.widget_notepasswords.hide()
        self.docx_ui.widget_notelist.hide()
        self.docx_ui.widget_contract.hide()
        self.docx_ui.widget_main.show()
        self.docx_creator.exec_()

    def select_list_el(self):
        if self.ttable_ui.sAWContent_hours_list.findChild(
                QtWidgets.QCommandLinkButton, "clb_" +
                                              str(datetime.datetime.timestamp(datetime.datetime.combine(
                                                  self.ttable_ui.calendar.selectedDate().toPyDate(),
                                                  datetime.datetime.min.time())))
        ) is not None:
            self.ttable_ui.sAWContent_hours_list.findChild(
                QtWidgets.QCommandLinkButton, "clb_" +
                                              str(datetime.datetime.timestamp(datetime.datetime.combine(
                                                  self.ttable_ui.calendar.selectedDate().toPyDate(),
                                                  datetime.datetime.min.time())))
            ).setChecked(1)
        else:
            self.load_db_timetable_list()

    def add_calculate_box(self, sel_sub=False, id_sub="", i=0, j=0):
        widget_calcbox = QtWidgets.QGroupBox(self.outlay_ui.widget_calcs)
        widget_calcbox.setObjectName("widget_calcbox_" + id_sub)
        gL_widget_calcbox = QtWidgets.QGridLayout(widget_calcbox)
        gL_widget_calcbox.setObjectName("gL_widget_calcbox_" + id_sub)
        studs_col = [['0']]

        sub_info = []

        if sel_sub:
            _db = ARMDataBase()
            _sql = "SELECT sub_name, sub_price_hour, sub_price_month, sub_hours_need FROM subjects WHERE id_sub=" + id_sub
            sub_info = _db.query(_sql)

            widget_calcbox.setTitle(sub_info[0][0])

            _sql = "SELECT COUNT(*) FROM subs_in_studs WHERE id_sub={0} AND status=1".format(id_sub)
            studs_col = _db.query(_sql)

            _db.close()
        else:
            widget_calcbox.setTitle("Предмет " + id_sub)

        lab_studs = QtWidgets.QLabel(widget_calcbox)
        lab_studs.setObjectName("lab_studs_" + id_sub)
        lab_studs.setText(self._translate("Outlay", "Количество слушателей: "))
        gL_widget_calcbox.addWidget(lab_studs, 1, 0)

        check_auto = QtWidgets.QCheckBox(widget_calcbox)
        check_auto.setObjectName("check_auto_" + id_sub)
        check_auto.setChecked(True)
        check_auto.setText(self._translate("Outlay", "Автоматический расчет"))
        gL_widget_calcbox.addWidget(check_auto, 0, 2)

        spin_studs = QtWidgets.QSpinBox(widget_calcbox)
        spin_studs.setObjectName("spin_studs_" + id_sub)
        spin_studs.setMaximum(100000)
        spin_studs.setSingleStep(1)
        gL_widget_calcbox.addWidget(spin_studs, 1, 1)

        radio_variability_studs = QtWidgets.QRadioButton(widget_calcbox)
        radio_variability_studs.setObjectName("radio_variability_studs_" + id_sub)
        gL_widget_calcbox.addWidget(radio_variability_studs, 1, 2)

        lab_price = QtWidgets.QLabel(widget_calcbox)
        lab_price.setObjectName("lab_price_" + id_sub)
        lab_price.setText(self._translate("Outlay", "Стоимость курса: "))
        gL_widget_calcbox.addWidget(lab_price, 2, 0)

        spin_price = QtWidgets.QSpinBox(widget_calcbox)
        spin_price.setObjectName("spin_price_" + id_sub)
        spin_price.setMaximum(1000000000)
        spin_price.setSingleStep(100)
        gL_widget_calcbox.addWidget(spin_price, 2, 1)

        radio_variability_price = QtWidgets.QRadioButton(widget_calcbox)
        radio_variability_price.setObjectName("radio_variability_price_" + id_sub)
        radio_variability_price.setChecked(True)
        gL_widget_calcbox.addWidget(radio_variability_price, 2, 2)

        lab_tax = QtWidgets.QLabel(widget_calcbox)
        lab_tax.setObjectName("lab_tax_" + id_sub)
        lab_tax.setText(self._translate("Outlay", "Почасовая оплата: "))
        gL_widget_calcbox.addWidget(lab_tax, 3, 0)

        spin_tax = QtWidgets.QSpinBox(widget_calcbox)
        spin_tax.setObjectName("spin_tax_" + id_sub)
        spin_tax.setMaximum(1000000000)
        spin_tax.setSingleStep(50)
        gL_widget_calcbox.addWidget(spin_tax, 3, 1)

        radio_variability_tax = QtWidgets.QRadioButton(widget_calcbox)
        radio_variability_tax.setObjectName("radio_variability_tax_" + id_sub)
        gL_widget_calcbox.addWidget(radio_variability_tax, 3, 2)

        lab_hours = QtWidgets.QLabel(widget_calcbox)
        lab_hours.setObjectName("lab_hours_" + id_sub)
        lab_hours.setText(self._translate("Outlay", "Количество часов: "))
        gL_widget_calcbox.addWidget(lab_hours, 4, 0)

        spin_hours = QtWidgets.QSpinBox(widget_calcbox)
        spin_hours.setObjectName("spin_hours_" + id_sub)
        spin_hours.setMaximum(1000)
        spin_hours.setSingleStep(2)
        gL_widget_calcbox.addWidget(spin_hours, 4, 1)

        # noinspection PyUnresolvedReferences
        spin_studs.valueChanged.connect(lambda: self.calculate_values())
        # noinspection PyUnresolvedReferences
        spin_hours.valueChanged.connect(lambda: self.calculate_values())
        # noinspection PyUnresolvedReferences
        spin_tax.valueChanged.connect(lambda: self.calculate_values())
        # noinspection PyUnresolvedReferences
        spin_price.valueChanged.connect(lambda: self.calculate_values())
        check_auto.clicked.connect(lambda: self.outlay_check_click())

        if sel_sub:
            spin_studs.setValue(int(studs_col[0][0]))
            try:
                spin_price.setValue(int(sub_info[0][2]))
            except ValueError:
                spin_price.setValue(int(0))
            try:
                spin_tax.setValue(int(sub_info[0][1]))
            except ValueError:
                spin_tax.setValue(int(0))
            try:
                spin_hours.setValue(int(sub_info[0][3]))
            except ValueError:
                spin_hours.setValue(int(0))

        self.outlay_ui.gL_widget_calcs.addWidget(widget_calcbox, i, j)
        return widget_calcbox

    def add_sub_teach_box(self, sel_sub=False, id_sub="", i=0, j=0):
        widget_sub_teach = QtWidgets.QWidget(self.outlay_print_ui.widget_subs_teachs)
        widget_sub_teach.setObjectName("widget_sub_teach_" + id_sub)
        gL_widget_sub_teach = QtWidgets.QGridLayout(widget_sub_teach)
        gL_widget_sub_teach.setObjectName("gL_widget_sub_teach_" + id_sub)

        lab_sub_name = QtWidgets.QLabel(widget_sub_teach)
        lab_sub_name.setObjectName("lab_sub_name_" + id_sub)
        if sel_sub:
            lab_sub_name.setText(self._translate("OutlayPrinter", "Предмет: "))
        else:
            lab_sub_name.setText(self._translate("OutlayPrinter", "Предмет {}: ".format(id_sub)))
        gL_widget_sub_teach.addWidget(lab_sub_name, 0, 0)

        lEdit_sub_name = QtWidgets.QLineEdit(widget_sub_teach)
        lEdit_sub_name.setObjectName("lEdit_sub_name_" + id_sub)
        gL_widget_sub_teach.addWidget(lEdit_sub_name, 0, 1)

        lab_teach_name = QtWidgets.QLabel(widget_sub_teach)
        lab_teach_name.setObjectName("lab_teach_name_" + id_sub)
        if sel_sub:
            lab_teach_name.setText(self._translate("OutlayPrinter", "Преподаватель: "))
        else:
            lab_teach_name.setText(self._translate("OutlayPrinter", "Преподаватель {}: ".format(id_sub)))
        gL_widget_sub_teach.addWidget(lab_teach_name, 1, 0)

        lEdit_teach_name = QtWidgets.QLineEdit(widget_sub_teach)
        lEdit_teach_name.setObjectName("lEdit_teach_name_" + id_sub)
        gL_widget_sub_teach.addWidget(lEdit_teach_name, 1, 1)

        if sel_sub:
            _db = ARMDataBase()
            _sql = "SELECT sub_name, id_teacher FROM subjects WHERE id_sub=" + id_sub
            sub_info = _db.query(_sql)
            lEdit_sub_name.setText(sub_info[0][0])

            _sql = "SELECT teacher_name FROM teachers WHERE id_teacher=" + str(sub_info[0][1])
            teach_info = _db.query(_sql)
            lEdit_teach_name.setText(teach_info[0][0])

            _db.close()

        self.outlay_print_ui.gL_widget_subs_teachs.addWidget(widget_sub_teach, i, j)

        return widget_sub_teach

    def save_calculate_values(self):
        i = 0
        for widget_calcbox in self.outlay_ui.widget_calcs.children():
            if widget_calcbox.objectName().startswith("widget_calcbox_") and i < 4:
                for child in widget_calcbox.children():
                    if child.objectName().startswith("spin_studs_"):
                        self.outlay_ui.variability_list[i][0] = child.value()
                    elif child.objectName().startswith("spin_price_"):
                        self.outlay_ui.variability_list[i][1] = child.value()
                    elif child.objectName().startswith("spin_tax_"):
                        self.outlay_ui.variability_list[i][2] = child.value()
                    elif child.objectName().startswith("spin_hours_"):
                        self.outlay_ui.variability_list[i][3] = child.value()
                    elif child.objectName().startswith("radio_variability_studs_") and child.isChecked():
                        self.outlay_ui.variability_list[i][4] = "radio_variability_studs_"
                    elif child.objectName().startswith("radio_variability_price_") and child.isChecked():
                        self.outlay_ui.variability_list[i][4] = "radio_variability_price_"
                    elif child.objectName().startswith("radio_variability_tax_") and child.isChecked():
                        self.outlay_ui.variability_list[i][4] = "radio_variability_tax_"
                    elif child.objectName().startswith("check_auto_"):
                        self.outlay_ui.variability_list[i][5] = child.isChecked()
                i += 1

    def load_calculate_values(self):
        i = 0
        for widget_calcbox in self.outlay_ui.widget_calcs.children():
            if widget_calcbox.objectName().startswith("widget_calcbox_"):
                for child in widget_calcbox.children():
                    if child.objectName().startswith("spin_studs_") \
                            and self.outlay_ui.variability_list[i][0] != "":
                        child.setValue(self.outlay_ui.variability_list[i][0])
                    elif child.objectName().startswith("spin_price_") \
                            and self.outlay_ui.variability_list[i][1] != "":
                        child.setValue(self.outlay_ui.variability_list[i][1])
                    elif child.objectName().startswith("spin_tax_") \
                            and self.outlay_ui.variability_list[i][2] != "":
                        child.setValue(self.outlay_ui.variability_list[i][2])
                    elif child.objectName().startswith("spin_hours_") \
                            and self.outlay_ui.variability_list[i][3] != "":
                        child.setValue(self.outlay_ui.variability_list[i][3])
                    elif self.outlay_ui.variability_list[i][4] in child.objectName() \
                            and self.outlay_ui.variability_list[i][4] != "":
                        child.setChecked(True)
                    elif child.objectName().startswith("check_auto_") \
                            and self.outlay_ui.variability_list[i][5] != "":
                        child.setChecked(self.outlay_ui.variability_list[i][5])
                i += 1
        self.outlay_check_click()

    def calculate_values(self):
        price = [0, 0, 0, 0]
        studs = [0, 0, 0, 0]
        hours = [0, 0, 0, 0]
        tax = [0, 0, 0, 0]
        profit = 0
        cost = 0
        i = 0
        for widget_calcbox in self.outlay_ui.widget_calcs.children():
            if widget_calcbox.objectName().startswith("widget_calcbox_") and i < 4:
                if widget_calcbox.findChild(QtWidgets.QCheckBox,
                                            "check_auto_" + widget_calcbox.objectName().split("_")[-1]).isChecked():
                    for child in widget_calcbox.children():
                        if child.objectName().startswith("radio_variability_tax_") \
                                and child.isChecked() \
                                and widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                             "spin_hours_" + widget_calcbox.objectName().split("_")[
                                                                 -1]).value() != 0:
                            widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                     "spin_tax_" + widget_calcbox.objectName().split("_")[-1]).setValue(
                                round(
                                    (1 / 2 *
                                     widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                              "spin_price_" + widget_calcbox.objectName().split("_")[
                                                                  -1]).value() *
                                     widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                              "spin_studs_" + widget_calcbox.objectName().split("_")[
                                                                  -1]).value()) /
                                    widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                             "spin_hours_" + widget_calcbox.objectName().split("_")[
                                                                 -1]).value()))
                        elif child.objectName().startswith("radio_variability_price_") \
                                and child.isChecked() \
                                and widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                             "spin_studs_" + widget_calcbox.objectName().split("_")[
                                                                 -1]).value() != 0:
                            widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                     "spin_price_" + widget_calcbox.objectName().split("_")[
                                                         -1]).setValue(round(
                                (2 *
                                 widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                          "spin_tax_" + widget_calcbox.objectName().split("_")[
                                                              -1]).value() *
                                 widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                          "spin_hours_" + widget_calcbox.objectName().split("_")[
                                                              -1]).value()) /
                                widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                         "spin_studs_" + widget_calcbox.objectName().split("_")[
                                                             -1]).value()))
                        elif child.objectName().startswith("radio_variability_studs_") \
                                and child.isChecked() \
                                and widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                             "spin_price_" + widget_calcbox.objectName().split("_")[
                                                                 -1]).value() != 0:
                            widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                     "spin_studs_" + widget_calcbox.objectName().split("_")[
                                                         -1]).setValue(round(
                                (2 *
                                 widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                          "spin_hours_" + widget_calcbox.objectName().split("_")[
                                                              -1]).value() *
                                 widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                          "spin_tax_" + widget_calcbox.objectName().split("_")[
                                                              -1]).value()) /
                                widget_calcbox.findChild(QtWidgets.QSpinBox,
                                                         "spin_price_" + widget_calcbox.objectName().split("_")[
                                                             -1]).value()))
                i += 1
        i = 0
        for widget_calcbox in self.outlay_ui.widget_calcs.children():
            if widget_calcbox.objectName().startswith("widget_calcbox_") and i < 4:
                for child in widget_calcbox.children():
                    if child.objectName().startswith("spin_studs_"):
                        studs[i] = child.value()
                    elif child.objectName().startswith("spin_price_"):
                        price[i] = child.value()
                    elif child.objectName().startswith("spin_tax_"):
                        tax[i] = child.value()
                    elif child.objectName().startswith("spin_hours_"):
                        hours[i] = child.value()
                i += 1
        for j in range(i):
            profit += price[j] * studs[j]
            cost += (hours[j] * tax[j]) * 1.302
        self.outlay_ui.label_profit.setText("Доходы = " + str(profit) + " рублей")
        self.outlay_ui.label_cost.setText("Оплата часов + ФОТ = " + str(round(cost, 2)) + " рублей")
        if profit != 0:
            self.outlay_ui.label_otfot.setText("Оплата часов + ФОТ % = " + str(round(cost / profit * 100, 2)) + "%")

    def outlay_check_click(self):
        for widget_calcbox in self.outlay_ui.widget_calcs.children():
            if widget_calcbox.objectName().startswith("widget_calcbox_"):
                for child in widget_calcbox.children():
                    if child.objectName().startswith("check_auto_") \
                            and child.isChecked():
                        widget_calcbox.findChild(QtWidgets.QRadioButton,
                                                 "radio_variability_tax_" + widget_calcbox.objectName().split("_")[
                                                     -1]).setEnabled(True)
                        widget_calcbox.findChild(QtWidgets.QRadioButton,
                                                 "radio_variability_price_" + widget_calcbox.objectName().split("_")[
                                                     -1]).setEnabled(True)
                        widget_calcbox.findChild(QtWidgets.QRadioButton,
                                                 "radio_variability_studs_" + widget_calcbox.objectName().split("_")[
                                                     -1]).setEnabled(True)
                    elif child.objectName().startswith("check_auto_") \
                            and not child.isChecked():
                        widget_calcbox.findChild(QtWidgets.QRadioButton,
                                                 "radio_variability_tax_" + widget_calcbox.objectName().split("_")[
                                                     -1]).setEnabled(False)
                        widget_calcbox.findChild(QtWidgets.QRadioButton,
                                                 "radio_variability_price_" + widget_calcbox.objectName().split("_")[
                                                     -1]).setEnabled(False)
                        widget_calcbox.findChild(QtWidgets.QRadioButton,
                                                 "radio_variability_studs_" + widget_calcbox.objectName().split("_")[
                                                     -1]).setEnabled(False)

    def create_outlay(self):
        rb = 0
        if self.outlay_ui.radio_col_1.isChecked():
            rb = 1
        elif self.outlay_ui.radio_col_2.isChecked():
            rb = 2
        elif self.outlay_ui.radio_col_3.isChecked():
            rb = 3
        elif self.outlay_ui.radio_col_4.isChecked():
            rb = 4
        outlay_data = [{}, {}, {}, {}, {
            "count": rb,
            "head": self.outlay_print_ui.comboBox_head.currentText(),
            "date_confirm": self.outlay_print_ui.dateEdit_date_confirm.date().toString('dd.MM.yyyy'),
            "program": self.outlay_print_ui.comboBox_prog.currentText(),
            "class": self.outlay_print_ui.lEdit_class.text(),
            "date_start": self.outlay_print_ui.dateEdit_date_start.date().toString('dd.MM.yyyy'),
            "date_end": self.outlay_print_ui.dateEdit_date_end.date().toString('dd.MM.yyyy'),
            "manager_cpui": self.outlay_print_ui.comboBox_manager_cpui.currentText(),
            "bookkeeper": self.outlay_print_ui.comboBox_bookkeeper.currentText(),
            "pfc": self.outlay_print_ui.comboBox_pfs.currentText()}]
        j = 0
        for i in self.outlay_ui.widget_calcs.children():
            if i.objectName().startswith("widget_calcbox_"):
                for spin in i.children():
                    if spin.objectName().startswith("spin_studs_"):
                        outlay_data[j]["studs"] = spin.value()
                    elif spin.objectName().startswith("spin_price_"):
                        outlay_data[j]["price"] = spin.value()
                    elif spin.objectName().startswith("spin_tax_"):
                        outlay_data[j]["tax"] = spin.value()
                    elif spin.objectName().startswith("spin_hours_"):
                        outlay_data[j]["hours"] = spin.value()
                j += 1
        j = 0
        for i in self.outlay_print_ui.widget_subs_teachs.children():
            if i.objectName().startswith("widget_sub_teach_"):
                for line in i.children():
                    if line.objectName().startswith("lEdit_sub_name_"):
                        outlay_data[j]["subject"] = line.text()
                    elif line.objectName().startswith("lEdit_teach_name_"):
                        outlay_data[j]["teacher"] = line.text()
                j += 1

        path = os.getcwd() + r"/Документы/Сметы/"
        filename = f"Смета {outlay_data[4]['class']} класс на курсы {outlay_data[4]['program']}," \
                   f" {outlay_data[4]['date_start'][-2:]}-{outlay_data[4]['date_end'][-4:]} №000000.docx"
        desk_list_dir = os.listdir(path)
        indexes_list = []
        for doc_in_dir in desk_list_dir:
            start_index = doc_in_dir.find('№')
            if start_index:
                try:
                    indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
                except Exception:
                    pass
        copy_index = 0
        while copy_index in indexes_list:
            copy_index += 1
            str_copy_index = str(copy_index)
            while len(str_copy_index) < 6:
                str_copy_index = "0" + str_copy_index
            filename = f"Смета {outlay_data[4]['class']} класс на курсы {outlay_data[4]['program']}," \
                       f" {outlay_data[4]['date_start'][-2:]}-{outlay_data[4]['date_end'][-4:]} №{str_copy_index}.docx"

        self.outlay_printer.close()
        set_doc_warning("Отправлено",
                        'Документ будет сохранен в сметы.\n'
                        'Вы сможете найти его во вкладке "Документы"->"Сметы"\n'
                        'Имя документа:\n' +
                        filename)

        thread_list = []
        task = threading.Thread(target=OutlayCreate(), args=(outlay_data,))
        thread_list.append(task)
        task.deamon = True
        task.start()

    def create_decree_enr(self):
        _db = ARMDataBase()
        _sql = "SELECT prog_range FROM programs WHERE id_prog=" + str(self.decree_enr_ui.comboBox_prog.currentData())
        try:
            prog_range = str(_db.query(_sql)[0][0])
        except IndexError:
            prog_range = "8"
        decree_enr_data = [{
            "head": self.decree_enr_ui.comboBox_head.currentText(),
            "program": self.decree_enr_ui.comboBox_prog.currentText(),
            "prog_range": prog_range,
            "group": self.decree_enr_ui.comboBox_group.currentText(),
            "date_start": self.decree_enr_ui.dateEdit_date_start.date().toString('dd.MM.yyyy'),
            "date_end": self.decree_enr_ui.dateEdit_date_end.date().toString('dd.MM.yyyy'),
            "manager_cpui": self.decree_enr_ui.comboBox_manager_cpui.currentText(),
            "office": self.decree_enr_ui.comboBox_office.currentText(),
            "pfc": self.decree_enr_ui.comboBox_pfs.currentText()},
        []]
        for i in self.decree_enr_ui.sAWContent_outlay_studs.children():
            if i.objectName().startswith("clb_") and i.isChecked():
                _sql = "SELECT student_name FROM students WHERE id_student=" + i.objectName().split("_")[-1]
                student_name = _db.query(_sql)[0][0]
                _sql = "SELECT id_sub FROM subs_in_studs WHERE status='1' AND id_student=" + i.objectName().split("_")[-1]
                subs_id = _db.query(_sql)
                subs_list = ''
                j = 0
                for sub in range(len(subs_id)):
                    _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(subs_id[sub][0])
                    if sub == 0:
                        subs_list += _db.query(_sql)[0][0].lower()
                    elif sub > 0:
                        subs_list += ', ' + _db.query(_sql)[0][0].lower()
                decree_enr_data[1].append([student_name, subs_list])
        _db.close()

        path = os.getcwd() + r"/Документы/Приказы/"
        filename = f"Приказ на зачисление {decree_enr_data[0]['program']}, " \
                   f"{decree_enr_data[0]['date_start'][-2:]}-{decree_enr_data[0]['date_end'][-2:]}, " \
                   f"{decree_enr_data[0]['group']} №000000.docx"
        desk_list_dir = os.listdir(path)
        indexes_list = []
        for doc_in_dir in desk_list_dir:
            start_index = doc_in_dir.find('№')
            if start_index:
                try:
                    indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
                except Exception:
                    pass
        copy_index = 0
        while copy_index in indexes_list:
            copy_index += 1
            str_copy_index = str(copy_index)
            while len(str_copy_index) < 6:
                str_copy_index = "0" + str_copy_index
            filename = f"Приказ на зачисление {decree_enr_data[0]['program']}, " \
                       f"{decree_enr_data[0]['date_start'][-2:]}-{decree_enr_data[0]['date_end'][-2:]}, " \
                       f"{decree_enr_data[0]['group']} №{str_copy_index}.docx"

        self.docx_creator.close()
        set_doc_warning("Отправлено",
                        'Документ будет сохранен в приказы.\n'
                        'Вы сможете найти его во вкладке "Документы"->"Приказы"\n'
                        'Имя документа:\n' +
                        filename)

        thread_list = []
        task = threading.Thread(target=DecreeEnrollmentCreate(), args=(decree_enr_data,))
        thread_list.append(task)
        task.deamon = True
        task.start()

    def create_note_passes(self):
        _db = ARMDataBase()
        _sql = "SELECT prog_range FROM programs WHERE id_prog=" + str(self.note_passes_ui.comboBox_prog.currentData())
        try:
            prog_range = str(_db.query(_sql)[0][0])
        except IndexError:
            prog_range = "8"
        data = [{
            "head": self.note_passes_ui.comboBox_head.currentText(),
            "program": self.note_passes_ui.comboBox_prog.currentText(),
            "group": self.note_passes_ui.comboBox_group.currentText(),
            "date_start": self.note_passes_ui.dateEdit_date_start.date().toString('dd.MM.yyyy'),
            "date_end": self.note_passes_ui.dateEdit_date_end.date().toString('dd.MM.yyyy'),
            "date": self.note_passes_ui.dateEdit_date.date().toString('dd.MM.yyyy'),
            "manager_cpui": self.note_passes_ui.comboBox_manager_cpui.currentText()},
        []]
        for i in self.note_passes_ui.sAWContent_outlay_studs.children():
            if i.objectName().startswith("clb_") and i.isChecked():
                _sql = "SELECT student_name FROM students WHERE id_student=" + i.objectName().split("_")[-1]
                student_name = _db.query(_sql)[0][0]
                _sql = "SELECT id_sub FROM subs_in_studs WHERE status='1' AND id_student=" + i.objectName().split("_")[-1]
                subs_id = _db.query(_sql)
                subs_list = ''
                j = 0
                for sub in range(len(subs_id)):
                    _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(subs_id[sub][0])
                    if sub == 0:
                        subs_list += _db.query(_sql)[0][0].lower()
                    elif sub > 0:
                        subs_list += ', ' + _db.query(_sql)[0][0].lower()
                data[1].append([student_name, subs_list])
        _db.close()

        path = os.getcwd() + r"/Документы/Записки/"
        filename = f"Служебка на пропуска {data[0]['program']}, " \
                   f"{data[0]['group']} №000000.docx"
        desk_list_dir = os.listdir(path)
        indexes_list = []
        for doc_in_dir in desk_list_dir:
            start_index = doc_in_dir.find('№')
            if start_index:
                try:
                    indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
                except Exception:
                    pass
        copy_index = 0
        while copy_index in indexes_list:
            copy_index += 1
            str_copy_index = str(copy_index)
            while len(str_copy_index) < 6:
                str_copy_index = "0" + str_copy_index
            filename = f"Служебка на пропуска {data[0]['program']}, " \
                       f"{data[0]['group']} №{str_copy_index}.docx"

        self.docx_creator.close()
        set_doc_warning("Отправлено",
                        'Документ будет сохранен в служебные записки.\n'
                        'Вы сможете найти его во вкладке "Документы"->"Служебные записки"\n'
                        'Имя документа:\n' +
                        filename)

        thread_list = []
        task = threading.Thread(target=NotePassesCreate(), args=(data,))
        thread_list.append(task)
        task.deamon = True
        task.start()

    def create_note_passwords(self):
        _db = ARMDataBase()
        _sql = "SELECT prog_range FROM programs WHERE id_prog=" + str(self.note_passwords_ui.comboBox_prog.currentData())
        try:
            prog_range = str(_db.query(_sql)[0][0])
        except IndexError:
            prog_range = "8"
        data = [{
            "depo": self.note_passwords_ui.comboBox_head.currentText(),
            "program": self.note_passwords_ui.comboBox_prog.currentText(),
            "prog_range": prog_range,
            "group": self.note_passwords_ui.comboBox_group.currentText(),
            "date_before": self.note_passwords_ui.dateEdit_date_before.date().toString('dd.MM.yyyy'),
            "date": self.note_passwords_ui.dateEdit_date.date().toString('dd.MM.yyyy'),
            "manager_cpui": self.note_passwords_ui.comboBox_manager_cpui.currentText()},
        []]
        for i in self.note_passwords_ui.sAWContent_outlay_studs.children():
            if i.objectName().startswith("clb_") and i.isChecked():
                _sql = "SELECT student_name FROM students WHERE id_student=" + i.objectName().split("_")[-1]
                student_name = _db.query(_sql)[0][0]
                _sql = "SELECT id_sub FROM subs_in_studs WHERE status='1' AND id_student=" + i.objectName().split("_")[-1]
                subs_id = _db.query(_sql)
                subs_list = ''
                j = 0
                for sub in range(len(subs_id)):
                    _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(subs_id[sub][0])
                    if sub == 0:
                        subs_list += _db.query(_sql)[0][0].lower()
                    elif sub > 0:
                        subs_list += ', ' + _db.query(_sql)[0][0].lower()
                data[1].append([student_name, subs_list])
        _db.close()

        path = os.getcwd() + r"/Документы/Записки/"
        filename = f"Служебка на пароли {data[0]['program']}, " \
                   f"{data[0]['group']} №000000.docx"
        desk_list_dir = os.listdir(path)
        indexes_list = []
        for doc_in_dir in desk_list_dir:
            start_index = doc_in_dir.find('№')
            if start_index:
                try:
                    indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
                except Exception:
                    pass
        copy_index = 0
        while copy_index in indexes_list:
            copy_index += 1
            str_copy_index = str(copy_index)
            while len(str_copy_index) < 6:
                str_copy_index = "0" + str_copy_index
            filename = f"Служебка на пароли {data[0]['program']}, " \
                       f"{data[0]['group']} №{str_copy_index}.docx"

        self.docx_creator.close()
        set_doc_warning("Отправлено",
                        'Документ будет сохранен в служебные записки.\n'
                        'Вы сможете найти его во вкладке "Документы"->"Служебные записки"\n'
                        'Имя документа:\n' +
                        filename)

        thread_list = []
        task = threading.Thread(target=NotePasswordCreate(), args=(data,))
        thread_list.append(task)
        task.deamon = True
        task.start()

    def create_note_list(self):
        _db = ARMDataBase()
        data = []
        iterat = 0
        for i in range(len(self.note_list_ui.sAWContent_groups.children())):
            it = self.note_list_ui.sAWContent_groups.children()[i]
            if it.objectName().startswith("clb_") and it.isChecked():
                _sql = "SELECT id_group, group_name, id_prog, class FROM groups WHERE id_group=" + it.objectName().split("_")[-1]
                groups_info = _db.query(_sql)[0]
                _sql = "SELECT prog_name FROM programs WHERE id_prog=" + str(groups_info[2])
                prog_name = _db.query(_sql)[0][0]
                data.append([[], groups_info[1], groups_info[3], prog_name, self.note_list_ui.checkBox_and_prog.isChecked()])
                _sql = "SELECT student_name, id_student FROM students WHERE id_group=" + str(groups_info[0])
                studs_list = _db.query(_sql)
                j = 0
                for stud in studs_list:
                    if self.note_list_ui.checkBox_and_prog.isChecked():
                        _sql = "SELECT id_sub FROM subs_in_studs WHERE status=1 AND id_student=" + str(stud[1])
                        sub_id = _db.query(_sql)
                        studs_subs = []
                        for id in sub_id:
                            _sql = "SELECT sub_name FROM subjects WHERE id_sub=" + str(id[0])
                            studs_subs.append(_db.query(_sql)[0][0])
                        _subs = ''
                        i1 = 1
                        for subs in studs_subs:
                            if i1 < len(sub_id):
                                _subs += subs + ", "
                            else:
                                _subs += subs
                            i1 += 1
                        data[iterat][0].append([stud[0], _subs])
                    else:
                        data[iterat][0].append([stud[0]])
                iterat += 1
        _db.close()

        path = os.getcwd() + r"/Документы/Записки/"
        filename = f"Служебка списки №000000.docx"
        desk_list_dir = os.listdir(path)
        indexes_list = []
        for doc_in_dir in desk_list_dir:
            start_index = doc_in_dir.find('№')
            if start_index:
                try:
                    indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
                except Exception:
                    pass
        copy_index = 0
        while copy_index in indexes_list:
            copy_index += 1
            str_copy_index = str(copy_index)
            while len(str_copy_index) < 6:
                str_copy_index = "0" + str_copy_index
            filename = f"Служебка списки №{str_copy_index}.docx"

        self.docx_creator.close()
        set_doc_warning("Отправлено",
                        'Документ будет сохранен в служебные записки.\n'
                        'Вы сможете найти его во вкладке "Документы"->"Служебные записки"\n'
                        'Имя документа:\n' +
                        filename)

        thread_list = []
        task = threading.Thread(target=NoteListCreate(), args=(data,))
        thread_list.append(task)
        task.deamon = True
        task.start()

    def create_contract(self):
        _db = ARMDataBase()
        if str(self.contract_ui.comboBox_prog.currentData()) != 'None':
            _sql = "SELECT prog_range FROM programs WHERE id_prog=" + str(self.contract_ui.comboBox_prog.currentData())
            prog_range = _db.query(_sql)[0][0]
        else:
            prog_range = '4'
        data = [
            {
            "head": self.contract_ui.comboBox_head.currentText(),
            "program": self.contract_ui.comboBox_prog.currentText(),
            "class": self.contract_ui.lEdit_class.text(),
            "date_start": self.contract_ui.dateEdit_date_start.date().toString('dd.MM.yyyy'),
            "date_end": self.contract_ui.dateEdit_date_end.date().toString('dd.MM.yyyy'),
            "manager_cpui": self.contract_ui.comboBox_manager_cpui.currentText(),
            "ls": self.contract_ui.comboBox_head_ls.currentText(),
            "fullname": self.contract_ui.lEdit_fullname.text(),
            "fullname_parent": self.contract_ui.lEdit_fullname_parent.text(),
            "date_birthday": self.contract_ui.dateEdit_birthday.date().toString('dd.MM.yyyy'),
            "passport_date": self.contract_ui.dateEdit_passport_date.date().toString('dd.MM.yyyy'),
            "passport_parent_date": self.contract_ui.dateEdit_passport_parent_date.date().toString('dd.MM.yyyy'),
            "male": self.contract_ui.radioButton_gender_male.isChecked(),
            "address": self.contract_ui.lEdit_address.text(),
            "phone": self.contract_ui.lEdit_phone.text(),
            "phone_parent": self.contract_ui.lEdit_phone_parent.text(),
            "mail": self.contract_ui.lEdit_mail.text(),
            "mail_parent": self.contract_ui.lEdit_mail_parent.text(),
            "uinst": self.contract_ui.lEdit_uinst.text(),
            "passport_seria": self.contract_ui.lEdit_seria.text(),
            "passport_number": self.contract_ui.lEdit_number.text(),
            "passport_seria_parent": self.contract_ui.lEdit_seria_parent.text(),
            "passport_number_parent": self.contract_ui.lEdit_number_parent.text(),
            "passport_whom": self.contract_ui.lEdit_passport_whom.text(),
            "passport_parent_whom": self.contract_ui.lEdit_passport_parent_whom.text(),
            "prog_range": prog_range,
            },
            []
        ]

        for i in range(len(self.contract_ui.sAWContent_subs.children())):
            it = self.contract_ui.sAWContent_subs.children()[i]
            if it.objectName().startswith("clb_") and it.isChecked():
                _sql = "SELECT id_sub, sub_name, sub_hours_need FROM subjects WHERE id_sub=" + it.objectName().split("_")[-1]
                sub_info = _db.query(_sql)
                j = 0
                for sub in sub_info:
                    data[1].append([sub[0], sub[1], sub[2]])

        _db.close()

        path = os.getcwd() + r"/Документы/Договора/"
        filename = f"№000000"
        desk_list_dir = os.listdir(path)
        indexes_list = []
        for doc_in_dir in desk_list_dir:
            start_index = doc_in_dir.find('№')
            if '№' in doc_in_dir:
                try:
                    if not doc_in_dir.startswith('~$'):
                        indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
                except Exception:
                    print('Exception')
        copy_index = 0
        while copy_index in indexes_list:
            copy_index += 1
            str_copy_index = str(copy_index)
            while len(str_copy_index) < 6:
                str_copy_index = "0" + str_copy_index
            filename = f"№{str_copy_index}"

        self.docx_creator.close()
        set_doc_warning("Отправлено",
                        'Документы будут сохранены в договора.\n'
                        'Вы сможете найти их во вкладке "Документы"->"Договора"\n'
                        'Номер документов:\n' +
                        filename)

        thread_list = []
        task = threading.Thread(target=ContractCreate(), args=(data,))
        thread_list.append(task)
        task.deamon = True
        task.start()

    def settings_window(self, war_icon=":/sfu_logo.ico"):
        def save_settings():
            pc.set_option('checking', str(int(settings_win.check_mail.isChecked())))
            pc.set_option('login', settings_win.lEdit_mail.text())
            pc.set_option('password', settings_win.lEdit_password.text())
            pc.set_option('mail', settings_win.lEdit_service.text())
            pc.set_option('sender', settings_win.lEdit_sender.text())
            path_for_save = settings_win.lEdit_path.text()
            if path_for_save.endswith('/') or path_for_save.endswith('\\'):
                pc.set_option('path_for_save_letters', settings_win.lEdit_path.text())
            else:
                pc.set_option('path_for_save_letters', settings_win.lEdit_path.text() + '/')
            pc.set_option('time_sleep', str(settings_win.spin_rate.value() * 60))
            set_win.close()
        set_win = QtWidgets.QDialog(self)
        settings_win = Ui_Settings()
        settings_win.setupUi(set_win)
        set_win.setWindowTitle('Settings')
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(war_icon), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        set_win.setWindowIcon(icon)

        settings_win.check_mail.setChecked(int(pc.get_option('checking')))
        settings_win.lEdit_mail.setText(pc.get_option('login'))
        settings_win.lEdit_password.setText(pc.get_option('password'))
        settings_win.lEdit_service.setText(pc.get_option('mail'))
        settings_win.lEdit_sender.setText(pc.get_option('sender'))
        settings_win.lEdit_path.setText(pc.get_option('path_for_save_letters'))
        settings_win.spin_rate.setValue(int(int(pc.get_option('time_sleep')) / 60))

        settings_win.btn_save.clicked.connect(lambda: save_settings())
        settings_win.btn_back.clicked.connect(lambda: set_win.close())
        set_win.exec_()


class OutlayCreate:
    def __call__(self, outlay_data):
        create_outlay_doc(outlay_data)


class DecreeEnrollmentCreate:
    def __call__(self, decree_enr_data):
        create_decree_enrollment_doc(decree_enr_data)


class NotePassesCreate:
    def __call__(self, data):
        create_note_passes_doc(data)


class NotePasswordCreate:
    def __call__(self, data):
        create_note_password_doc(data)


class NoteListCreate:
    def __call__(self, data):
        create_note_list_doc(data)


class ContractCreate:
    def __call__(self, data):
        create_contract(data)


def create_outlay_doc(outlay_data):
    path = os.getcwd() + r"/Документы/Сметы/"

    filename = f"Смета {outlay_data[4]['class']} класс на курсы {outlay_data[4]['program']}," \
               f" {outlay_data[4]['date_start'][-2:]}-{outlay_data[4]['date_end'][-4:]} №000000.docx"
    desk_list_dir = os.listdir(path)
    indexes_list = []
    for doc_in_dir in desk_list_dir:
        start_index = doc_in_dir.find('№')
        if start_index:
            try:
                indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
            except Exception:
                pass
    copy_index = 0
    while copy_index in indexes_list:
        copy_index += 1
        str_copy_index = str(copy_index)
        while len(str_copy_index) < 6:
            str_copy_index = "0" + str_copy_index
        filename = f"Смета {outlay_data[4]['class']} класс на курсы {outlay_data[4]['program']}," \
                   f" {outlay_data[4]['date_start'][-2:]}-{outlay_data[4]['date_end'][-4:]} №{str_copy_index}.docx"

    doc = docx.Document()

    # START DOC / FIRST PAGE
    date_confirm = ""
    if outlay_data[4]['date_confirm'][3:5] == "01":
        date_confirm = 'января'
    elif outlay_data[4]['date_confirm'][3:5] == "02":
        date_confirm = 'февраля'
    elif outlay_data[4]['date_confirm'][3:5] == "03":
        date_confirm = 'марта'
    elif outlay_data[4]['date_confirm'][3:5] == "04":
        date_confirm = 'апреля'
    elif outlay_data[4]['date_confirm'][3:5] == "05":
        date_confirm = 'мая'
    elif outlay_data[4]['date_confirm'][3:5] == "06":
        date_confirm = 'июня'
    elif outlay_data[4]['date_confirm'][3:5] == "07":
        date_confirm = 'июля'
    elif outlay_data[4]['date_confirm'][3:5] == "08":
        date_confirm = 'августа'
    elif outlay_data[4]['date_confirm'][3:5] == "09":
        date_confirm = 'сентября'
    elif outlay_data[4]['date_confirm'][3:5] == "10":
        date_confirm = 'октября'
    elif outlay_data[4]['date_confirm'][3:5] == "11":
        date_confirm = 'ноября'
    elif outlay_data[4]['date_confirm'][3:5] == "12":
        date_confirm = 'декабря'

    doc.sections[0].page_height = docx.shared.Cm(29.7)
    doc.sections[0].page_width = docx.shared.Cm(21)
    doc.sections[0].top_margin = docx.shared.Cm(2)
    doc.sections[0].right_margin = docx.shared.Cm(1.5)
    doc.sections[0].left_margin = docx.shared.Cm(3)
    doc.sections[0].bottom_margin = docx.shared.Cm(1.25)

    # HEADER
    tab_head = doc.add_table(rows=4, cols=2)
    tab_head.alignment = WD_TABLE_ALIGNMENT.RIGHT
    cell = tab_head.cell(0, 0)
    cell.merge(tab_head.cell(0, 1))
    cell.text = "УТВЕРЖДАЮ:"
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cell.add_paragraph(" ")
    cell.paragraphs[1].runs[0].font.name = "Times New Roman"
    cell.paragraphs[1].runs[0].font.size = docx.shared.Pt(14)
    cell.paragraphs[1].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[1].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    tab_head.rows[0].height = docx.shared.Cm(1.1)
    set_cell_margins(cell, start=0)
    cell.width = docx.shared.Cm(7.1)

    cell = tab_head.cell(1, 0)
    cell.merge(tab_head.cell(1, 1))
    cell.text = "Директор ХТИ – филиала СФУ"
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cell.add_paragraph(" ")
    cell.paragraphs[1].runs[0].font.name = "Times New Roman"
    cell.paragraphs[1].runs[0].font.size = docx.shared.Pt(14)
    cell.paragraphs[1].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[1].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    tab_head.rows[1].height = docx.shared.Cm(1.1)
    set_cell_margins(cell, start=0)
    cell.width = docx.shared.Cm(7.1)

    cell = tab_head.cell(2, 0)
    cell.text = " "
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    set_cell_margins(cell, start=0)
    cell.width = docx.shared.Cm(2.1)
    set_cell_border(
        cell,
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )

    cell = tab_head.cell(2, 1)

    cell.text = f"{outlay_data[4]['head'].split(' ')[-2][:1]}. " \
                f"{outlay_data[4]['head'].split(' ')[-1][:1]}. " \
                f"{outlay_data[4]['head'].split(' ')[-3]}"

    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    tab_head.rows[2].height = docx.shared.Cm(0.55)
    set_cell_margins(cell, start=0)
    cell.width = docx.shared.Cm(5)

    cell = tab_head.cell(3, 0)
    cell.merge(tab_head.cell(3, 1))
    cell.text = " "
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    cell.add_paragraph(f"«")
    cell.paragraphs[1].runs[0].font.name = "Times New Roman"
    cell.paragraphs[1].runs[0].font.size = docx.shared.Pt(14)

    cell.paragraphs[1].add_run(f"{outlay_data[4]['date_confirm'][:2]}")
    cell.paragraphs[1].runs[1].font.name = "Times New Roman"
    cell.paragraphs[1].runs[1].font.size = docx.shared.Pt(14)
    cell.paragraphs[1].runs[1].font.underline = True

    cell.paragraphs[1].add_run(f"» ")
    cell.paragraphs[1].runs[2].font.name = "Times New Roman"
    cell.paragraphs[1].runs[2].font.size = docx.shared.Pt(14)

    cell.paragraphs[1].add_run(f"{date_confirm} ")
    cell.paragraphs[1].runs[3].font.name = "Times New Roman"
    cell.paragraphs[1].runs[3].font.size = docx.shared.Pt(14)
    cell.paragraphs[1].runs[3].font.underline = True

    cell.paragraphs[1].add_run(f"{outlay_data[4]['date_confirm'][6:]}г.")
    cell.paragraphs[1].runs[4].font.name = "Times New Roman"
    cell.paragraphs[1].runs[4].font.size = docx.shared.Pt(14)

    cell.paragraphs[1].paragraph_format.space_after = docx.shared.Pt(0)
    cell.paragraphs[1].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    tab_head.rows[1].height = docx.shared.Cm(1.1)
    set_cell_margins(cell, start=0)
    cell.width = docx.shared.Cm(7.1)

    par = doc.add_paragraph(" ")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.paragraph_format.space_after = docx.shared.Pt(0)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    par = doc.add_paragraph(" ")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.paragraph_format.space_after = docx.shared.Pt(0)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    par = doc.add_paragraph(f"СМЕТА")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_paragraph("доходов и расходов на проведение курсов")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.paragraph_format.space_after = docx.shared.Pt(0)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_paragraph(
        f"«{outlay_data[4]['program']}» ({outlay_data[4]['class']} класс)")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.paragraph_format.space_after = docx.shared.Pt(0)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_paragraph(
        f"(с {outlay_data[4]['date_start']} г. по {outlay_data[4]['date_end']} г.)")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.paragraph_format.space_after = docx.shared.Pt(0)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_paragraph(" ")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.paragraph_format.space_after = docx.shared.Pt(0)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # TABLE OUTLAY
    tab_outlay = doc.add_table(rows=7, cols=4, style='Table Grid')

    add_cell("Код", tab_outlay.cell(0, 0), tab_outlay.rows[0], 1.69, 0.53)
    add_cell("Наименование предметных статей", tab_outlay.cell(0, 1), tab_outlay.rows[0], 8.08, 0.53)
    add_cell("Отношение к затратам, %", tab_outlay.cell(0, 2), tab_outlay.rows[0], 4.26, 0.53)
    add_cell("Сумма, руб.", tab_outlay.cell(0, 3), tab_outlay.rows[0], 2.6, 0.53)

    profit = 0.0
    cost = 0.0
    for i in range(outlay_data[4]["count"]):
        profit += outlay_data[i]["studs"] * outlay_data[i]["price"]
        cost += (outlay_data[i]["hours"] * outlay_data[i]["tax"]) * 1.302
    if len(str(round(profit, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(" ", tab_outlay.cell(1, 0), tab_outlay.rows[1], 1.69, 0.53)
    add_cell("общий доход",
             tab_outlay.cell(1, 1), tab_outlay.rows[1], 8.08, 0.53, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(" ", tab_outlay.cell(1, 2), tab_outlay.rows[1], 4.26, 0.53)
    add_cell(str(round(profit, 2)).replace(".", ",") + zero, tab_outlay.cell(1, 3), tab_outlay.rows[1], 2.6, 0.53,
             font_size=12)

    add_cell(" ", tab_outlay.cell(2, 0), tab_outlay.rows[2], 1.69, 0.53)
    add_cell("расходы, всего в том числе:",
             tab_outlay.cell(2, 1), tab_outlay.rows[2], 8.08, 0.53, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell("100,00", tab_outlay.cell(2, 2), tab_outlay.rows[2], 4.26, 0.53, font_size=12)
    add_cell(str(round(profit, 2)).replace(".", ",") + zero, tab_outlay.cell(2, 3), tab_outlay.rows[2], 2.6, 0.53,
             font_size=12)

    try:
        otfot = cost / profit * 100
    except ZeroDivisionError:
        otfot = 0.0
    if len(str(round(otfot, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(" ", tab_outlay.cell(3, 0), tab_outlay.rows[3], 1.69, 0.53)
    add_cell("вознаграждение за образовательные услуги гражданско-правового характера и "
                "начисление страховых взносов во внебюджетные фонды",
             tab_outlay.cell(3, 1), tab_outlay.rows[3], 8.08, 0.53, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(str(round(otfot, 2)).replace(".", ",") + zero,
             tab_outlay.cell(3, 2), tab_outlay.rows[3], 4.26, 0.53, font_size=12)
    if len(str(round(cost, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(str(round(cost, 2)).replace(".", ",") + zero,
             tab_outlay.cell(3, 3), tab_outlay.rows[3], 2.6, 0.53, font_size=12)

    public_service = 0.05 * profit
    if len(str(round(public_service, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(" ", tab_outlay.cell(4, 0), tab_outlay.rows[4], 1.69, 0.53)
    add_cell("коммунальные услуги",
             tab_outlay.cell(4, 1), tab_outlay.rows[4], 8.08, 0.53, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell("5,00", tab_outlay.cell(4, 2), tab_outlay.rows[4], 4.26, 0.53, font_size=12)
    add_cell(str(round(public_service, 2)).replace(".", ",") + zero,
             tab_outlay.cell(4, 3), tab_outlay.rows[4], 2.6, 0.53, font_size=12)

    inst_cost = 0.18 * profit
    if len(str(round(inst_cost, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(" ", tab_outlay.cell(5, 0), tab_outlay.rows[5], 1.69, 0.53)
    add_cell("общеинститутские расходы",
             tab_outlay.cell(5, 1), tab_outlay.rows[5], 8.08, 0.53, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell("18,00", tab_outlay.cell(5, 2), tab_outlay.rows[5], 4.26, 0.53, font_size=12)
    add_cell(str(round(inst_cost, 2)).replace(".", ",") + zero,
             tab_outlay.cell(5, 3), tab_outlay.rows[5], 2.6, 0.53, font_size=12)

    div_cost_ratio = 100.0 - otfot - 5.0 - 18.0
    div_cost = div_cost_ratio * profit / 100
    if len(str(round(div_cost_ratio, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(" ", tab_outlay.cell(6, 0), tab_outlay.rows[6], 1.69, 0.53)
    add_cell("расходы подразделения",
             tab_outlay.cell(6, 1), tab_outlay.rows[6], 8.08, 0.53, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(str(round(div_cost_ratio, 2)).replace(".", ",") + zero,
             tab_outlay.cell(6, 2), tab_outlay.rows[6], 4.26, 0.53, font_size=12)
    if len(str(round(div_cost, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(str(round(div_cost, 2)).replace(".", ",") + zero,
             tab_outlay.cell(6, 3), tab_outlay.rows[6], 2.6, 0.53, font_size=12)

    par = doc.add_paragraph(" ")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)

    tab_approve = doc.add_table(rows=6, cols=2)

    add_cell("Зав. ЦПЮИ", tab_approve.cell(0, 0), tab_approve.rows[1], 11.94, 1.12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(f"{outlay_data[4]['manager_cpui'].split(' ')[-2][:1]}. "
                f"{outlay_data[4]['manager_cpui'].split(' ')[-1][:1]}. "
                f"{outlay_data[4]['manager_cpui'].split(' ')[-3]}",
             tab_approve.cell(0, 1), tab_approve.rows[0], 4.94, 1.12, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_cell(" ", tab_approve.cell(1, 0), tab_approve.rows[1], 11.94, 1.12)
    add_cell(" ", tab_approve.cell(1, 1), tab_approve.rows[1], 4.94, 1.12)

    add_cell("Согласовано:", tab_approve.cell(2, 0), tab_approve.rows[2], 11.94, 1.12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(" ", tab_approve.cell(2, 1), tab_approve.rows[2], 4.94, 1.12)

    add_cell("Гл. бухгалтер", tab_approve.cell(3, 0), tab_approve.rows[3], 11.94, 1.12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(f"{outlay_data[4]['bookkeeper'].split(' ')[-2][:1]}. "
                f"{outlay_data[4]['bookkeeper'].split(' ')[-1][:1]}. "
                f"{outlay_data[4]['bookkeeper'].split(' ')[-3]}",
             tab_approve.cell(3, 1), tab_approve.rows[3], 4.94, 1.12, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_cell(" ", tab_approve.cell(4, 0), tab_approve.rows[4], 11.94, 1.12)
    add_cell(" ", tab_approve.cell(4, 1), tab_approve.rows[4], 4.94, 1.12)

    add_cell("Зав. ПФС", tab_approve.cell(5, 0), tab_approve.rows[5], 11.94, 1.12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(f"{outlay_data[4]['pfc'].split(' ')[-2][:1]}. "
                f"{outlay_data[4]['pfc'].split(' ')[-1][:1]}. "
                f"{outlay_data[4]['pfc'].split(' ')[-3]}",
             tab_approve.cell(5, 1), tab_approve.rows[5], 4.94, 1.12, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # SECOND PAGE
    doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)

    doc.sections[1].orientation = docx.enum.section.WD_ORIENTATION.PORTRAIT
    doc.sections[1].page_height = docx.shared.Cm(21)
    doc.sections[1].page_width = docx.shared.Cm(29.7)
    doc.sections[1].top_margin = docx.shared.Cm(3)
    doc.sections[1].right_margin = docx.shared.Cm(2)
    doc.sections[1].left_margin = docx.shared.Cm(1.25)
    doc.sections[1].bottom_margin = docx.shared.Cm(1.5)

    par = doc.add_paragraph("Пояснение к смете")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_paragraph("Планирование доходов")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    plan_profit = doc.add_table(rows=2 + outlay_data[4]['count'], cols=4, style='Table Grid')
    plan_profit.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_cell(f"Предмет",
             plan_profit.cell(0, 0), plan_profit.rows[0], 4.25, 1.31, font_size=12)
    add_cell(f"Стоимость, руб.",
             plan_profit.cell(0, 1), plan_profit.rows[0], 4.25, 1.31, font_size=12)
    add_cell(f" ",
             plan_profit.cell(0, 2), plan_profit.rows[0], 4.25, 1.31, font_size=12)
    plan_profit.cell(0, 2).add_paragraph(f"Количество слушателей, чел.")
    plan_profit.cell(0, 2).paragraphs[-1].runs[0].font.name = "Times New Roman"
    plan_profit.cell(0, 2).paragraphs[-1].runs[0].font.size = docx.shared.Pt(12)
    plan_profit.cell(0, 2).paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_cell(f" ",
             plan_profit.cell(0, 3), plan_profit.rows[0], 4.25, 1.31, font_size=12)
    plan_profit.cell(0, 3).add_paragraph(f"Планируемый доход, руб.")
    plan_profit.cell(0, 3).paragraphs[-1].runs[0].font.name = "Times New Roman"
    plan_profit.cell(0, 3).paragraphs[-1].runs[0].font.size = docx.shared.Pt(12)
    plan_profit.cell(0, 3).paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_sub in range(outlay_data[4]['count']):
        sub_profit = float(outlay_data[row_sub]['studs'] * outlay_data[row_sub]['price'])
        if len(str(round(sub_profit, 2)).split(".")[-1]) < 2:
            zero = "0"
        else:
            zero = ""
        add_cell(outlay_data[row_sub]['subject'],
                 plan_profit.cell(row_sub + 1, 0), plan_profit.rows[row_sub + 1], 4.25, 0.45, font_size=12)
        add_cell(str(outlay_data[row_sub]['price']),
                 plan_profit.cell(row_sub + 1, 1), plan_profit.rows[row_sub + 1], 4.25, 0.45, font_size=12)
        add_cell(str(outlay_data[row_sub]['studs']),
                 plan_profit.cell(row_sub + 1, 2), plan_profit.rows[row_sub + 1], 4.25, 0.45, font_size=12)
        add_cell(str(round(sub_profit, 2)).replace(".", ",") + zero,
                 plan_profit.cell(row_sub + 1, 3), plan_profit.rows[row_sub + 1], 4.25, 0.45, font_size=12)

    if len(str(round(profit, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell("Итого",
             plan_profit.cell(outlay_data[4]['count'] + 1, 0), plan_profit.rows[outlay_data[4]['count'] + 1],
             4.25, 0.45, font_size=12)
    add_cell(" ",
             plan_profit.cell(outlay_data[4]['count'] + 1, 1), plan_profit.rows[outlay_data[4]['count'] + 1],
             4.25, 0.45, font_size=12)
    add_cell(" ",
             plan_profit.cell(outlay_data[4]['count'] + 1, 2), plan_profit.rows[outlay_data[4]['count'] + 1],
             4.25, 0.45, font_size=12)
    add_cell(str(round(profit, 2)).replace(".", ",") + zero,
             plan_profit.cell(outlay_data[4]['count'] + 1, 3), plan_profit.rows[outlay_data[4]['count'] + 1],
             4.25, 0.45, font_size=12, bold=True)

    # THIRD TABLE
    par = doc.add_paragraph(" ")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_paragraph("Распределение и оплата часов")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    payment_for_hours = doc.add_table(rows=3 + outlay_data[4]['count'], cols=9, style='Table Grid')
    payment_for_hours.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_cell(" ", payment_for_hours.cell(0, 0), payment_for_hours.rows[0], 3.53, 0.93, font_size=12)
    add_cell("ФИО", payment_for_hours.cell(0, 1), payment_for_hours.rows[0], 4.0, 0.93, font_size=12)
    add_cell("Ауд. занятия, час.", payment_for_hours.cell(0, 2), payment_for_hours.rows[0], 2.25, 0.93, font_size=12)
    add_cell("Всего, час", payment_for_hours.cell(0, 3), payment_for_hours.rows[0], 1.75, 0.93, font_size=12)
    add_cell("Стоимость 1 час., руб.", payment_for_hours.cell(0, 4), payment_for_hours.rows[0], 3, 0.93, font_size=12)
    add_cell("Сумма, руб.", payment_for_hours.cell(0, 5), payment_for_hours.rows[0], 2.75, 0.93, font_size=12)
    add_cell("Начисление на оплату труда, %", payment_for_hours.cell(0, 6), payment_for_hours.rows[0],
             3.5, 0.93, font_size=12)
    add_cell("Сумма, руб.", payment_for_hours.cell(0, 7), payment_for_hours.rows[0], 2.5, 0.93, font_size=12)
    add_cell("Всего, руб.", payment_for_hours.cell(0, 8), payment_for_hours.rows[0], 3, 0.93, font_size=12)

    add_cell("ППС", payment_for_hours.cell(1, 0), payment_for_hours.rows[0], 3.53, 0.45, font_size=12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(" ", payment_for_hours.cell(1, 1), payment_for_hours.rows[0], 4.0, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(1, 2), payment_for_hours.rows[0], 2.25, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(1, 3), payment_for_hours.rows[0], 1.75, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(1, 4), payment_for_hours.rows[0], 3, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(1, 5), payment_for_hours.rows[0], 2.75, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(1, 6), payment_for_hours.rows[0], 3.5, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(1, 7), payment_for_hours.rows[0], 2.5, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(1, 8), payment_for_hours.rows[0], 3, 0.45, font_size=12)

    for row_sub in range(outlay_data[4]['count']):
        try:
            teacher = f"{outlay_data[row_sub]['teacher'].split(' ')[1][0:1]}. " \
                      f"{outlay_data[row_sub]['teacher'].split(' ')[2][0:1]}. " \
                      f"{outlay_data[row_sub]['teacher'].split(' ')[0]}"
        except IndexError:
            teacher = " "
        tax = outlay_data[row_sub]['tax']
        hours = outlay_data[row_sub]['hours']
        price_before_ofot = float(hours * tax)
        if len(str(round(price_before_ofot, 2)).split(".")[-1]) < 2:
            zero = "0"
        else:
            zero = ""
        add_cell("Преподаватель", payment_for_hours.cell(row_sub + 2, 0), payment_for_hours.rows[row_sub + 2],
                 3.53, 0.45, font_size=12, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_cell(teacher, payment_for_hours.cell(row_sub + 2, 1), payment_for_hours.rows[row_sub + 2],
                 4.0, 0.45, font_size=12)
        add_cell(str(hours), payment_for_hours.cell(row_sub + 2, 2), payment_for_hours.rows[row_sub + 2],
                 2.25, 0.45, font_size=12)
        add_cell(str(hours), payment_for_hours.cell(row_sub + 2, 3), payment_for_hours.rows[row_sub + 2],
                 1.75, 0.45, font_size=12)
        add_cell(str(tax), payment_for_hours.cell(row_sub + 2, 4), payment_for_hours.rows[row_sub + 2],
                 3, 0.45, font_size=12)
        add_cell(str(round(price_before_ofot, 2)).replace(".", ",") + zero, payment_for_hours.cell(row_sub + 2, 5),
                 payment_for_hours.rows[row_sub + 2], 2.75, 0.45, font_size=12)
        add_cell("30,2", payment_for_hours.cell(row_sub + 2, 6), payment_for_hours.rows[row_sub + 2],
                 3.5, 0.45, font_size=12)
        price_ofot = float(price_before_ofot * 0.302)
        if len(str(round(price_ofot, 2)).split(".")[-1]) < 2:
            zero = "0"
        else:
            zero = ""
        add_cell(str(round(price_ofot, 2)).replace(".", ",") + zero, payment_for_hours.cell(row_sub + 2, 7),
                 payment_for_hours.rows[row_sub + 2], 2.5, 0.45, font_size=12)
        price_after_ofot = float(price_ofot + price_before_ofot)
        if len(str(round(price_after_ofot, 2)).split(".")[-1]) < 2:
            zero = "0"
        else:
            zero = ""
        add_cell(str(round(price_after_ofot, 2)).replace(".", ",") + zero, payment_for_hours.cell(row_sub + 2, 8),
                 payment_for_hours.rows[row_sub + 2], 3, 0.45, font_size=12)

    last_row = outlay_data[4]['count'] + 2
    add_cell(" ", payment_for_hours.cell(last_row, 0),
             payment_for_hours.rows[last_row], 3.53, 0.45, font_size=12)
    add_cell("Всего ППС", payment_for_hours.cell(last_row, 1), payment_for_hours.rows[last_row], 4.0, 0.45,
             font_size=12, bold=True)
    add_cell(" ", payment_for_hours.cell(last_row, 2), payment_for_hours.rows[last_row], 2.25, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(last_row, 3), payment_for_hours.rows[last_row], 1.75, 0.45, font_size=12)
    add_cell(" ", payment_for_hours.cell(last_row, 4), payment_for_hours.rows[last_row], 3, 0.45, font_size=12)
    sum_price_before_ofot = 0.0
    for i in range(outlay_data[4]['count']):
        sum_price_before_ofot += float(outlay_data[i]['hours'] * outlay_data[i]['tax'])
    if len(str(round(sum_price_before_ofot, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(str(round(sum_price_before_ofot, 2)).replace(".", ",") + zero, payment_for_hours.cell(last_row, 5),
             payment_for_hours.rows[last_row], 2.75, 0.45, font_size=12, bold=True)
    add_cell(" ", payment_for_hours.cell(last_row, 6), payment_for_hours.rows[last_row], 3.5, 0.45, font_size=12)
    sum_price_ofot = float(sum_price_before_ofot * 0.302)
    if len(str(round(sum_price_ofot, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(str(round(sum_price_ofot, 2)).replace(".", ",") + zero, payment_for_hours.cell(last_row, 7),
             payment_for_hours.rows[last_row], 2.5, 0.45, font_size=12, bold=True)
    sum_after_ofot = float(sum_price_ofot + sum_price_before_ofot)
    if len(str(round(sum_after_ofot, 2)).split(".")[-1]) < 2:
        zero = "0"
    else:
        zero = ""
    add_cell(str(round(sum_after_ofot, 2)).replace(".", ",") + zero, payment_for_hours.cell(last_row, 8),
             payment_for_hours.rows[last_row], 3, 0.45, font_size=12, bold=True)

    # SECOND HEADERS
    par = doc.add_paragraph(" ")
    par.runs[0].font.name = "Times New Roman"
    par.runs[0].font.size = docx.shared.Pt(14)
    par.paragraph_format.space_after = docx.shared.Pt(0)

    tab_approve = doc.add_table(rows=3, cols=2)

    add_cell("Зав. ЦПЮИ", tab_approve.cell(0, 0), tab_approve.rows[1], 11, 1.12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(f"{outlay_data[4]['manager_cpui'].split(' ')[-2][:1]}. "
                f"{outlay_data[4]['manager_cpui'].split(' ')[-1][:1]}. "
                f"{outlay_data[4]['manager_cpui'].split(' ')[-3]}",
             tab_approve.cell(0, 1), tab_approve.rows[0], 6, 1.12, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_cell("Согласовано:", tab_approve.cell(1, 0), tab_approve.rows[1], 11, 1.12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(" ", tab_approve.cell(1, 1), tab_approve.rows[1], 6, 1.12)

    add_cell("Зав. ПФС", tab_approve.cell(2, 0), tab_approve.rows[2], 11, 1.12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(f"{outlay_data[4]['pfc'].split(' ')[-2][:1]}. "
                f"{outlay_data[4]['pfc'].split(' ')[-1][:1]}. "
                f"{outlay_data[4]['pfc'].split(' ')[-3]}",
             tab_approve.cell(2, 1), tab_approve.rows[2], 6, 1.12, text_alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # SETTINGS FOR DOCX
    properties = doc.core_properties
    properties.author = "ЦПЮИ ХТИ"
    # DOCX SAVE
    doc.save(path + filename)
    return path, filename


def add_cell(text, __cell, __rows, _width, _height, height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST, font_size=14,
             font_name="Times New Roman",
             text_alignment=WD_ALIGN_PARAGRAPH.CENTER, cell_alignment=WD_ALIGN_VERTICAL.CENTER,
             line_spacing=WD_LINE_SPACING.SINGLE, __space_after=0, underline=False, bold=False, first_line_indent=0):
    __cell.text = text
    __cell.paragraphs[0].runs[0].font.name = font_name
    __cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(font_size)
    __cell.paragraphs[0].paragraph_format.line_spacing_rule = line_spacing
    __cell.paragraphs[0].runs[0].font.underline = underline
    __cell.paragraphs[0].runs[0].font.bold = bold
    __cell.paragraphs[0].alignment = text_alignment
    __cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(__space_after)
    __cell.paragraphs[0].paragraph_format.first_line_indent = docx.shared.Cm(first_line_indent)
    __cell.vertical_alignment = cell_alignment
    __cell.width = docx.shared.Cm(_width)
    __rows.height_rule = height_rule
    __rows.height = docx.shared.Cm(_height)


def add_par(document, text, font_name="Times New Roman", font_size=12, space_after=0,
            line_spacing_rule=WD_LINE_SPACING.SINGLE, text_aligment=WD_ALIGN_PARAGRAPH.CENTER,
            underline=False, bold=False, italic=False, first_line_indent=0):
    par = document.add_paragraph(text)
    par.runs[0].font.name = font_name
    par.runs[0].font.size = docx.shared.Pt(font_size)
    par.paragraph_format.space_after = docx.shared.Pt(space_after)
    par.paragraph_format.line_spacing = 1.15
    par.paragraph_format.line_spacing_rule = line_spacing_rule
    par.paragraph_format.first_line_indent = docx.shared.Cm(first_line_indent)
    par.alignment = text_aligment
    par.runs[0].font.underline = underline
    par.runs[0].font.bold = bold
    par.runs[0].font.italic = italic


def create_decree_enrollment_doc(decree_enr_data):
    path = os.getcwd() + r"/Документы/Приказы/"
    filename = f"Приказ на зачисление {decree_enr_data[0]['program']}, " \
               f"{decree_enr_data[0]['date_start'][-2:]}-{decree_enr_data[0]['date_end'][-2:]}, " \
               f"{decree_enr_data[0]['group']} №000000.docx"
    desk_list_dir = os.listdir(path)
    indexes_list = []
    for doc_in_dir in desk_list_dir:
        start_index = doc_in_dir.find('№')
        if start_index:
            try:
                indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
            except Exception:
                pass
    copy_index = 0
    while copy_index in indexes_list:
        copy_index += 1
        str_copy_index = str(copy_index)
        while len(str_copy_index) < 6:
            str_copy_index = "0" + str_copy_index
        filename = f"Приказ на зачисление {decree_enr_data[0]['program']}, " \
                   f"{decree_enr_data[0]['date_start'][-2:]}-{decree_enr_data[0]['date_end'][-2:]}, " \
                   f"{decree_enr_data[0]['group']} №{str_copy_index}.docx"

    doc = docx.Document()

    # START DOC / FIRST PAGE

    doc.sections[0].page_height = docx.shared.Cm(29.7)
    doc.sections[0].page_width = docx.shared.Cm(21)
    doc.sections[0].top_margin = docx.shared.Cm(1.5)
    doc.sections[0].right_margin = docx.shared.Cm(1.5)
    doc.sections[0].left_margin = docx.shared.Cm(3)
    doc.sections[0].bottom_margin = docx.shared.Cm(1.25)

    # HEADER
    add_par(doc, "Министерство науки и высшего образования РФ", line_spacing_rule=WD_LINE_SPACING.MULTIPLE, italic=True)
    add_par(doc, "Федеральное государственное автономное образовательное учреждение",
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE, italic=True)
    add_par(doc, "высшего образования", line_spacing_rule=WD_LINE_SPACING.MULTIPLE, italic=True)
    add_par(doc, "«СИБИРСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ»",
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE, bold=True, italic=True)
    add_par(doc, "ХАКАССКИЙ ТЕХНИЧЕСКИЙ ИНСТИТУТ –", line_spacing_rule=WD_LINE_SPACING.MULTIPLE, bold=True, italic=True)
    add_par(doc, "филиал ФГАОУ ВО «Сибирский федеральный университет»",
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE, italic=True)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, italic=True)
    add_par(doc, "ПРИКАЗ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, italic=True, bold=True)

    tab_num_decree = doc.add_table(rows=1, cols=3)
    tab_num_decree.alignment = WD_TABLE_ALIGNMENT.RIGHT
    add_cell(" ", tab_num_decree.cell(0, 0), tab_num_decree.rows[0], 3.47, 0.45, font_size=12)
    set_cell_border(
        tab_num_decree.cell(0, 0),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )
    add_cell("№", tab_num_decree.cell(0, 1), tab_num_decree.rows[0], 0.79, 0.45, font_size=12)
    add_cell(" ", tab_num_decree.cell(0, 2), tab_num_decree.rows[0], 2.19, 0.45, font_size=12, height_rule=WD_ROW_HEIGHT_RULE.EXACTLY)
    set_cell_border(
        tab_num_decree.cell(0, 2),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )

    add_par(doc, " ")
    add_par(doc, "По основной", text_aligment=WD_ALIGN_PARAGRAPH.LEFT)
    add_par(doc, "деятельности", text_aligment=WD_ALIGN_PARAGRAPH.LEFT, underline=True)
    new_run = doc.paragraphs[-1].add_run("                                               г. Абакан")
    new_run.font.name = "Times New Roman"
    new_run.font.size = docx.shared.Pt(12)
    add_par(doc, " ")
    add_par(doc, "«О контингенте обучающихся по дополнительной образовательной программе в ЦПЮИ»",
            first_line_indent=3.5, text_aligment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_par(doc, " ")
    add_par(doc, "ПРИКАЗЫВАЮ:", text_aligment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_par(doc, " ")
    add_par(
        doc,
        "1. Зачислить в ЦПЮИ с {0} по {1} года на {2}-х месячную дополнительную образовательную программу «{3}» "
        "(группа {4}) следующих обучающихся, оплативших дополнительные услуги и заключивших "
        "соответствующий договор:".format(
            decree_enr_data[0]['date_start'],
            decree_enr_data[0]['date_end'],
            decree_enr_data[0]['prog_range'],
            decree_enr_data[0]['program'],
            decree_enr_data[0]['group']
        ),
        text_aligment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=1.25)
    add_par(doc, " ")
    add_par(doc, " ")
    add_par(doc, " ", font_size=9)
    studs = []
    for i in range(len(decree_enr_data[1])):
        studs.append(decree_enr_data[1][i])
    try:
        studs.sort(key=lambda x: x[0])
    except IndexError:
        pass
    for i in range(len(studs)):
        try:
            add_par(
                doc,
                f"{str(i+1)}. {studs[i][0]} ({studs[i][1]});",
                first_line_indent=3.5,
                text_aligment=WD_ALIGN_PARAGRAPH.LEFT
            )
        except IndexError:
            pass
    doc.paragraphs[-1].runs[0].text = doc.paragraphs[-1].runs[0].text.replace(";", ".")
    add_par(doc, " ")
    add_par(doc, " ")
    add_par(
        doc,
        f"2. Зав. канцелярией {decree_enr_data[0]['office'].split(' ')[-3]} "
        f"{decree_enr_data[0]['office'].split(' ')[-2][:1]}. "
        f"{decree_enr_data[0]['office'].split(' ')[-1][:1]}. ознакомить с приказом всех поименованных в нем лиц.",
        text_aligment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=1.25
    )
    add_par(doc, " ")
    add_par(
        doc,
        f"3. Контроль за исполнением приказа возложить на зав. ЦПЮИ "
        f"{decree_enr_data[0]['manager_cpui'].split(' ')[-3]} "
        f"{decree_enr_data[0]['manager_cpui'].split(' ')[-2][:1]}. "
        f"{decree_enr_data[0]['manager_cpui'].split(' ')[-1][:1]}.",
        text_aligment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=1.25
    )
    add_par(doc, " ")
    add_par(doc, " ")
    add_par(doc, " ")
    add_par(doc, " ")
    tab_head = doc.add_table(rows=1, cols=2)
    tab_head.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_cell("   Директор", tab_head.cell(0, 0), tab_head.rows[0], 6, 0.4, font_size=12,
             text_alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_cell(
        f"{decree_enr_data[0]['head'].split(' ')[-2][:1]}. "
        f"{decree_enr_data[0]['head'].split(' ')[-1][:1]}. "
        f"{decree_enr_data[0]['head'].split(' ')[-3]}",
        tab_head.cell(0, 1), tab_head.rows[0], 6.8, 0.4, font_size=12, text_alignment=WD_ALIGN_PARAGRAPH.RIGHT
    )

    # SECOND PAGE
    doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    doc.sections[-1].page_height = docx.shared.Cm(29.7)
    doc.sections[-1].page_width = docx.shared.Cm(21)
    doc.sections[-1].top_margin = docx.shared.Cm(1.5)
    doc.sections[-1].right_margin = docx.shared.Cm(1.5)
    doc.sections[-1].left_margin = docx.shared.Cm(3)
    doc.sections[-1].bottom_margin = docx.shared.Cm(1.25)
    add_par(doc, "ЛИСТ СОГЛАСОВАНИЯ", font_size=14, bold=True)
    add_par(doc, " ", font_size=14)
    add_par(doc, "приказ № _______ от _______ 20____ г.", font_size=14)
    add_par(doc, " ", font_size=14)

    tab_pfc = doc.add_table(rows=2, cols=5)
    tab_pfc.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_cell(f"зав. ПФС", tab_pfc.cell(0, 0), tab_pfc.rows[0], 3.76, 0.46, first_line_indent=-0.5)
    add_cell(f" ", tab_pfc.cell(0, 1), tab_pfc.rows[0], 1, 0.46)
    add_cell(f" ", tab_pfc.cell(0, 2), tab_pfc.rows[0], 4, 0.46)
    add_cell(f" ", tab_pfc.cell(0, 3), tab_pfc.rows[0], 0.75, 0.46)
    add_cell(f"{decree_enr_data[0]['pfc'].split(' ')[-2][:1]}. "
             f"{decree_enr_data[0]['pfc'].split(' ')[-1][:1]}. "
             f"{decree_enr_data[0]['pfc'].split(' ')[-3]}", tab_pfc.cell(0, 4), tab_pfc.rows[0], 5, 0.46)
    add_cell(f"должность лица, согласующего", tab_pfc.cell(1, 0), tab_pfc.rows[1], 3.76, 0.46,
             first_line_indent=-0.2, font_size=9, cell_alignment=WD_ALIGN_VERTICAL.TOP)
    add_cell(f"       приказ", tab_pfc.cell(1, 1), tab_pfc.rows[1], 1, 0.46, font_size=9,
             first_line_indent=-0.2, text_alignment=WD_ALIGN_PARAGRAPH.LEFT, cell_alignment=WD_ALIGN_VERTICAL.TOP)
    tab_pfc.cell(1, 0).merge(tab_pfc.cell(1, 1))
    add_cell(f"личная подпись", tab_pfc.cell(1, 2), tab_pfc.rows[1], 4, 0.46, font_size=9,
             cell_alignment=WD_ALIGN_VERTICAL.TOP)
    add_cell(f" ", tab_pfc.cell(1, 3), tab_pfc.rows[1], 0.75, 0.46)
    add_cell(f"инициалы, фамилия", tab_pfc.cell(1, 4), tab_pfc.rows[1], 5, 0.46, font_size=9,
             cell_alignment=WD_ALIGN_VERTICAL.TOP)
    set_cell_border(
        tab_pfc.cell(0, 0),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )
    set_cell_border(
        tab_pfc.cell(0, 2),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )
    set_cell_border(
        tab_pfc.cell(0, 4),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )

    add_par(doc, " ", font_size=10)

    tab_cpui = doc.add_table(rows=2, cols=5)
    tab_cpui.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_cell(f"зав. ЦПЮИ", tab_cpui.cell(0, 0), tab_cpui.rows[0], 3.76, 0.46)
    add_cell(f" ", tab_cpui.cell(0, 1), tab_cpui.rows[0], 1, 0.46)
    add_cell(f" ", tab_cpui.cell(0, 2), tab_cpui.rows[0], 4, 0.46)
    add_cell(f" ", tab_cpui.cell(0, 3), tab_cpui.rows[0], 0.75, 0.46)
    add_cell(f"{decree_enr_data[0]['manager_cpui'].split(' ')[-2][:1]}. "
             f"{decree_enr_data[0]['manager_cpui'].split(' ')[-1][:1]}. "
             f"{decree_enr_data[0]['manager_cpui'].split(' ')[-3]}", tab_cpui.cell(0, 4), tab_cpui.rows[0], 5, 0.46)
    add_cell(f"должность лица, согласующего", tab_cpui.cell(1, 0), tab_cpui.rows[1], 3.76, 0.46,
             first_line_indent=-0.2, font_size=9, cell_alignment=WD_ALIGN_VERTICAL.TOP)
    add_cell(f"       приказ", tab_cpui.cell(1, 1), tab_cpui.rows[1], 1, 0.46, font_size=9,
             first_line_indent=-0.2, text_alignment=WD_ALIGN_PARAGRAPH.LEFT, cell_alignment=WD_ALIGN_VERTICAL.TOP)
    tab_cpui.cell(1, 0).merge(tab_cpui.cell(1, 1))
    add_cell(f"личная подпись", tab_cpui.cell(1, 2), tab_cpui.rows[1], 4, 0.46, font_size=9,
             cell_alignment=WD_ALIGN_VERTICAL.TOP)
    add_cell(f" ", tab_cpui.cell(1, 3), tab_cpui.rows[1], 0.75, 0.46)
    add_cell(f"инициалы, фамилия", tab_cpui.cell(1, 4), tab_cpui.rows[1], 5, 0.46, font_size=9,
             cell_alignment=WD_ALIGN_VERTICAL.TOP)
    set_cell_border(
        tab_cpui.cell(0, 0),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )
    set_cell_border(
        tab_cpui.cell(0, 2),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )
    set_cell_border(
        tab_cpui.cell(0, 4),
        bottom={"sz": 10, "val": "single", "color": "#000000"}
    )


    # SETTINGS FOR DOCX
    properties = doc.core_properties
    properties.author = "ЦПЮИ ХТИ"
    # DOCX SAVE
    doc.save(path + filename)
    return path, filename


def create_note_passes_doc(data):
    path = os.getcwd() + r"/Документы/Записки/"
    filename = f"Служебка на пропуска {data[0]['program']}, " \
               f"{data[0]['date_start'][-2:]}-{data[0]['date_end'][-2:]}, " \
               f"{data[0]['group']} №000000.docx"
    desk_list_dir = os.listdir(path)
    indexes_list = []
    for doc_in_dir in desk_list_dir:
        start_index = doc_in_dir.find('№')
        if start_index:
            try:
                indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
            except Exception:
                pass
    copy_index = 0
    while copy_index in indexes_list:
        copy_index += 1
        str_copy_index = str(copy_index)
        while len(str_copy_index) < 6:
            str_copy_index = "0" + str_copy_index
        filename = f"Служебка на пропуска {data[0]['program']}, " \
                   f"{data[0]['date_start'][-2:]}-{data[0]['date_end'][-2:]}, " \
                   f"{data[0]['group']} №{str_copy_index}.docx"

    doc = docx.Document()

    # START DOC / FIRST PAGE

    doc.sections[0].page_height = docx.shared.Cm(29.7)
    doc.sections[0].page_width = docx.shared.Cm(21)
    doc.sections[0].top_margin = docx.shared.Cm(2)
    doc.sections[0].right_margin = docx.shared.Cm(1.5)
    doc.sections[0].left_margin = docx.shared.Cm(3)
    doc.sections[0].bottom_margin = docx.shared.Cm(2)

    # HEADER
    add_par(doc, "Директору  ХТИ - филиала СФУ",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_par(doc, f"{data[0]['head'].split(' ')[-3]} "
                 f"{data[0]['head'].split(' ')[-2][:1]}. "
                 f"{data[0]['head'].split(' ')[-1][:1]}.",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_par(doc, "Зав. ЦПЮИ  " + f"{data[0]['manager_cpui'].split(' ')[-3]} "
                                 f"{data[0]['manager_cpui'].split(' ')[-2][:1]}. "
                                 f"{data[0]['manager_cpui'].split(' ')[-1][:1]}.",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.RIGHT)

    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE)
    add_par(doc, "СЛУЖЕБНАЯ ЗАПИСКА",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.LEFT)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE)
    add_par(doc, f"{data[0]['date']} г.",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.LEFT)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE)
    add_par(doc, f"В связи с открытием дополнительной образовательной программы «{data[0]['program']}»"
                 f" в ЦПЮИ с {data[0]['date_start']} г. по {data[0]['date_end']} г. прошу, для обучающихся"
                 f" в группе {data[0]['group']}, изготовить пропуска согласно списку:",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=1.25)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE)

    studs = []
    for i in range(len(data[1])):
        studs.append(data[1][i])
    try:
        studs.sort(key=lambda x: x[0])
    except IndexError:
        pass
    for i in range(len(studs)):
        try:
            add_par(
                doc,
                f"{str(i + 1)}. {studs[i][0]} ({studs[i][1]});",
                first_line_indent=4.5,
                text_aligment=WD_ALIGN_PARAGRAPH.LEFT
            )
        except IndexError:
            pass
    doc.paragraphs[-1].runs[0].text = doc.paragraphs[-1].runs[0].text.replace(";", ".")
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE)
    add_par(doc, f"Зав. ЦПЮИ                                                                   "
                 f"{data[0]['manager_cpui'].split(' ')[-2][:1]}. "
                 f"{data[0]['manager_cpui'].split(' ')[-1][:1]}. "
                 f"{data[0]['manager_cpui'].split(' ')[-3]}",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.LEFT)

    # SETTINGS FOR DOCX
    properties = doc.core_properties
    properties.author = "ЦПЮИ ХТИ"
    # DOCX SAVE
    doc.save(path + filename)
    return path, filename


def create_note_password_doc(data):
    path = os.getcwd() + r"/Документы/Записки/"
    filename = f"Служебка на пароли {data[0]['program']}, " \
               f"{data[0]['group']} №000000.docx"
    desk_list_dir = os.listdir(path)
    indexes_list = []
    for doc_in_dir in desk_list_dir:
        start_index = doc_in_dir.find('№')
        if start_index:
            try:
                indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
            except Exception:
                pass
    copy_index = 0
    while copy_index in indexes_list:
        copy_index += 1
        str_copy_index = str(copy_index)
        while len(str_copy_index) < 6:
            str_copy_index = "0" + str_copy_index
        filename = f"Служебка на пароли {data[0]['program']}, " \
                   f"{data[0]['group']} №{str_copy_index}.docx"

    doc = docx.Document()

    # START DOC / FIRST PAGE

    doc.sections[0].page_height = docx.shared.Cm(29.7)
    doc.sections[0].page_width = docx.shared.Cm(21)
    doc.sections[0].top_margin = docx.shared.Cm(2)
    doc.sections[0].right_margin = docx.shared.Cm(1.5)
    doc.sections[0].left_margin = docx.shared.Cm(3)
    doc.sections[0].bottom_margin = docx.shared.Cm(2)

    # HEADER
    add_par(
        doc,
        "Руководителю департамента",
        first_line_indent=8.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=14
    )
    add_par(
        doc,
        "информационных технологий ",
        first_line_indent=8.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=14
    )
    add_par(
        doc,
        f"{data[0]['depo'].split(' ')[-3]} "
        f"{data[0]['depo'].split(' ')[-2][:1]}. "
        f"{data[0]['depo'].split(' ')[-1][:1]}.",
        first_line_indent=8.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=14
    )
    add_par(
        doc,
        f"зав. кафедрой ПИМиЕД",
        first_line_indent=8.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=14
    )
    add_par(
        doc,
        f"ХТИ – филиала СФУ",
        first_line_indent=8.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=14
    )
    add_par(
        doc,
        f"{data[0]['manager_cpui'].split(' ')[-3]} "
        f"{data[0]['manager_cpui'].split(' ')[-2][:1]}. "
        f"{data[0]['manager_cpui'].split(' ')[-1][:1]}.",
        first_line_indent=8.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=14
    )


    add_par(doc, " ", font_size=14)
    add_par(doc, "СЛУЖЕБНАЯ ЗАПИСКА",
            text_aligment=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=14)
    add_par(doc, " ", font_size=14)
    add_par(doc, f"{data[0]['date']} г.",
            text_aligment=WD_ALIGN_PARAGRAPH.LEFT, font_size=14)
    add_par(doc, " ", font_size=14)
    add_par(doc, f"Прошу создать учетные записи для входа на еКурсы (https://e.sfu-kras.ru) слушателям "
                 f"Центра подготовки юного инженера ХТИ – филиала СФУ с доступом до {data[0]['date_before']}:",
            line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, text_aligment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=1.25,
            font_size=14)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE, font_size=14)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE, font_size=14)

    studs = []
    for i in range(len(data[1])):
        studs.append(data[1][i])
    try:
        studs.sort(key=lambda x: x[0])
    except IndexError:
        pass
    for i in range(len(studs)):
        try:
            add_par(
                doc,
                f"{str(i + 1)}. {studs[i][0]} ({studs[i][1]});",
                first_line_indent=1.25,
                text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=14
            )
        except IndexError:
            pass
    doc.paragraphs[-1].runs[0].text = doc.paragraphs[-1].runs[0].text.replace(";", ".")
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE, font_size=14)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE, font_size=14)
    add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, font_size=14)

    tab_manager_cpui = doc.add_table(rows=1, cols=2)
    add_cell(
        'Зав. кафедрой ПИМиЕД ХТИ – филилала СФУ',
        tab_manager_cpui.cell(0, 0),
        tab_manager_cpui.rows[0],
        6.5, 1.4, text_alignment=WD_ALIGN_PARAGRAPH.LEFT
    )
    add_cell(
        f"{data[0]['manager_cpui'].split(' ')[-2][:1]}. "
        f"{data[0]['manager_cpui'].split(' ')[-1][:1]}. "
        f"{data[0]['manager_cpui'].split(' ')[-3]} ",
        tab_manager_cpui.cell(0, 1),
        tab_manager_cpui.rows[0],
        10.44, 1.4, text_alignment=WD_ALIGN_PARAGRAPH.RIGHT
    )

    # SETTINGS FOR DOCX
    properties = doc.core_properties
    properties.author = "ЦПЮИ ХТИ"
    # DOCX SAVE
    doc.save(path + filename)
    return path, filename


def create_note_list_doc(data):
    path = os.getcwd() + r"/Документы/Записки/"
    filename = f"Служебка списки №000000.docx"
    desk_list_dir = os.listdir(path)
    indexes_list = []
    for doc_in_dir in desk_list_dir:
        start_index = doc_in_dir.find('№')
        if start_index:
            try:
                indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
            except Exception:
                pass
    copy_index = 0
    while copy_index in indexes_list:
        copy_index += 1
        str_copy_index = str(copy_index)
        while len(str_copy_index) < 6:
            str_copy_index = "0" + str_copy_index
        filename = f"Служебка списки №{str_copy_index}.docx"

    doc = docx.Document()

    # START DOC / FIRST PAGE

    doc.sections[0].page_height = docx.shared.Cm(29.7)
    doc.sections[0].page_width = docx.shared.Cm(21)
    doc.sections[0].top_margin = docx.shared.Cm(2)
    doc.sections[0].right_margin = docx.shared.Cm(1.5)
    doc.sections[0].left_margin = docx.shared.Cm(3)
    doc.sections[0].bottom_margin = docx.shared.Cm(2)

    # HEADER
    add_par(
        doc,
        "СПИСКИ ОБУЧАЮЩИХСЯ",
        text_aligment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        font_size=14
    )
    add_par(
        doc,
        " ",
        text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=14
    )

    for group in data:
        add_par(
            doc,
            f"Группа: {group[1]}:",
            text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=10,
            font_size=14
        )

        studs = []
        for i in range(len(group[0])):
            studs.append(group[0][i])
        try:
            studs.sort(key=lambda x: x[0])
        except IndexError:
            pass

        for i in range(len(studs)):
            try:
                add_par(
                    doc,
                    f"{str(i + 1)}. {studs[i][0]} ({studs[i][1]});" if group[4] else f"{str(i + 1)}. {studs[i][0]};",
                    first_line_indent=1.25,
                    text_aligment=WD_ALIGN_PARAGRAPH.LEFT,
                    font_size=14
                )
            except IndexError:
                pass
        doc.paragraphs[-1].runs[0].text = doc.paragraphs[-1].runs[0].text.replace(";", ".")
        add_par(doc, " ", line_spacing_rule=WD_LINE_SPACING.SINGLE, font_size=14)

    # SETTINGS FOR DOCX
    properties = doc.core_properties
    properties.author = "ЦПЮИ ХТИ"
    # DOCX SAVE
    doc.save(path + filename)
    return path, filename


def create_contract(data):
    path = os.getcwd() + r"/Документы/Договора/"
    dir = os.getcwd()
    filename = f"№000000"
    desk_list_dir = os.listdir(path)
    indexes_list = []
    for doc_in_dir in desk_list_dir:
        start_index = doc_in_dir.find('№')
        if '№' in doc_in_dir:
            try:
                if not doc_in_dir.startswith('~$'):
                    indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
            except Exception:
                print('Exception')
    copy_index = 0
    while copy_index in indexes_list:
        copy_index += 1
        str_copy_index = str(copy_index)
        while len(str_copy_index) < 6:
            str_copy_index = "0" + str_copy_index
        filename = f"№{str_copy_index}"

    cla = ''
    for i in range(len(data[1])):
        if i + 1 < len(data[1]):
            cla += f"по предмету {data[1][i][1].lower()} - {str(data[1][i][2])} часов ({str(int(int(data[1][i][2]) / 2))} занятий), "
        else:
            cla += f"по предмету {data[1][i][1].lower()} - {str(data[1][i][2])} часов ({str(int(int(data[1][i][2]) / 2))} занятий)."

    doc = DocxTemplate(dir + "\\ШАБЛОНЫ\\Данные.docx")
    context = {
        'fullname': data[0]['fullname'],
        'fullname_parent': data[0]['fullname_parent'],
        'daybirth': data[0]['date_birthday'],
        'gender': 'Мужской' if data[0]['male'] else 'Женский',
        'address': data[0]['address'],
        'phone': data[0]['phone'],
        'phone_parent': data[0]['phone_parent'],
        'mail': data[0]['mail'],
        'mail_parent': data[0]['mail_parent'],
        'uinst': data[0]['uinst'],
        'passport_seria': data[0]['passport_seria'],
        'passport_number': data[0]['passport_number'],
        'passport_whom': data[0]['passport_whom'],
        'passport_date': data[0]['passport_date'],
        'passport_parent_date': data[0]['passport_parent_date'],
        'psp': data[0]['passport_seria_parent'],
        'pnp': data[0]['passport_number_parent'],
        'passport_parent_whom': data[0]['passport_parent_whom'],
        'cla': cla,
    }

    doc.render(context)
    doc.save(path + filename + f" Данные {data[0]['fullname']}.docx")

    doc = DocxTemplate(dir + "\\ШАБЛОНЫ\\Договор.docx")
    context = {
        'fullname': data[0]['fullname'],
        'fullname_parent': data[0]['fullname_parent'],
        'address': data[0]['address'],
        'phone': data[0]['phone'],
        'phone_par': data[0]['phone_parent'],
        'seria_par': data[0]['passport_seria_parent'],
        'passno_par': data[0]['passport_number_parent'],
        'seria': data[0]['passport_seria'],
        'passno': data[0]['passport_number'],
        'who_par': data[0]['passport_parent_whom'],
        'when_par': data[0]['passport_parent_date'],
        'who': data[0]['passport_whom'],
        'when': data[0]['passport_date'],
        'prog_range': data[0]['prog_range'],
        'date_end': data[0]['date_end'],
        'date_start': data[0]['date_start'],
        'prog_name': data[0]['program'],
        'cla': cla,
        'manager_cpui':
            f"{data[0]['manager_cpui'].split(' ')[-2][:1]}. "
            f"{data[0]['manager_cpui'].split(' ')[-1][:1]}. "
            f"{data[0]['manager_cpui'].split(' ')[-3]}",
        'head':
            f"{data[0]['head'].split(' ')[-2][:1]}. "
            f"{data[0]['head'].split(' ')[-1][:1]}. "
            f"{data[0]['head'].split(' ')[-3]}",
        'ls':
            f"{data[0]['ls'].split(' ')[-2][:1]}. "
            f"{data[0]['ls'].split(' ')[-1][:1]}. "
            f"{data[0]['ls'].split(' ')[-3]}",
    }
    doc.render(context)
    doc.save(path + filename + f" Договор {data[0]['fullname']}.docx")

    doc = DocxTemplate(dir + "\\ШАБЛОНЫ\\Заявление.docx")
    context = {
        'firstname': data[0]['fullname'],
        'address': data[0]['address'],
        'phone': data[0]['phone'],
        'mail': data[0]['mail'],
        'place': data[0]['uinst'],
        'date_end': data[0]['date_end'],
        'date_start': data[0]['date_start'],
        'program': data[0]['program'],
        'cla': cla,
        'manager_cpui':
            f"{data[0]['manager_cpui'].split(' ')[-2][:1]}. "
            f"{data[0]['manager_cpui'].split(' ')[-1][:1]}. "
            f"{data[0]['manager_cpui'].split(' ')[-3]}",
        'head':
            f"{data[0]['head'].split(' ')[-2][:1]}. "
            f"{data[0]['head'].split(' ')[-1][:1]}. "
            f"{data[0]['head'].split(' ')[-3]}",
    }
    doc.render(context)
    doc.save(path + filename + f" Заявление {data[0]['fullname']}.docx")

    doc = DocxTemplate(dir + "\\ШАБЛОНЫ\\Согласие на обработку персональных данных родителей.docx")
    context = {
        'firstname_pred': data[0]['fullname_parent'],
        'address': data[0]['address'],
        'seria_pred': data[0]['passport_seria'],
        'passno_pred': data[0]['passport_number'],
        'who_pred': data[0]['passport_parent_whom'],
        'when_pred': data[0]['passport_parent_date'],
    }
    doc.render(context)
    doc.save(path + filename + f" Согласие на обработку персональных данных родителей {data[0]['fullname']}.docx")

    return path, filename


def create_timetable(sub):
    thread_list = []
    task = threading.Thread(target=TimetableCreate(), args=(sub,))
    thread_list.append(task)
    task.deamon = True
    task.start()


class TimetableCreate:
    def __call__(self, sub):
        create_timetable_doc(sub)


def check_new_messages():
    thread_list = []
    task = threading.Thread(target=CheckNewMessage(), args=())
    thread_list.append(task)
    task.deamon = True
    task.name = "MailCheckerARM"
    task.start()


class CheckNewMessage:
    def __call__(self):
        while threading.main_thread().is_alive():
            try:
                if pc.get_option('login') != "None" \
                and pc.get_option('login') != "" \
                and pc.get_option('sender') != "" \
                and pc.get_option('password') != "None" \
                and pc.get_option('password') != "" \
                and pc.get_option('checking') != "0" \
                and pc.get_option('sender') != "None":
                    check_new_message()
            except Exception:
                print(f"Ошибка в вызове check_new_message()")
            t = 0
            while (t < 60 \
                   and threading.main_thread().is_alive()) \
            or (t < int(pc.get_option("time_sleep")) \
                and threading.main_thread().is_alive()):
                t += 5
                sleep(5)


def check_new_message():
    def set_name_doc(_folder):
        def select_doc_name(doc, path_doc, _folder):
            if doc.paragraphs[0].runs[0].text.startswith("Договор"):
                os.rename(
                    path_doc,
                    f"{pc.get_option('path_for_save_letters')}{_folder}\\Договор .docx"
                )
            elif doc.paragraphs[0].runs[0].text.startswith(
                    "                                                                       "):
                os.rename(
                    path_doc,
                    f"{pc.get_option('path_for_save_letters')}{_folder}\\Заявление .docx"
                )
            elif doc.paragraphs[0].runs[0].text.startswith("Конфиденциально"):
                os.rename(
                    path_doc,
                    f"{pc.get_option('path_for_save_letters')}{_folder}\\СОГЛАСИЕ .docx"
                )
            elif doc.paragraphs[0].runs[0].text.startswith("ФИО:"):
                os.rename(
                    path_doc,
                    f"{pc.get_option('path_for_save_letters')}{_folder}\\Данные .docx"
                )
            else:
                os.rename(
                    path_doc,
                    f"{pc.get_option('path_for_save_letters')}{_folder}\\Неизвестный документ.docx"
                )

        select_doc_name(
            docx.Document(f"{pc.get_option('path_for_save_letters')}{_folder}\\1.docx"),
            f"{pc.get_option('path_for_save_letters')}{_folder}\\1.docx",
            _folder
        )
        select_doc_name(
            docx.Document(f"{pc.get_option('path_for_save_letters')}{_folder}\\2.docx"),
            f"{pc.get_option('path_for_save_letters')}{_folder}\\2.docx",
            _folder
        )
        select_doc_name(
            docx.Document(f"{pc.get_option('path_for_save_letters')}{_folder}\\3.docx"),
            f"{pc.get_option('path_for_save_letters')}{_folder}\\3.docx",
            _folder
        )
        select_doc_name(
            docx.Document(f"{pc.get_option('path_for_save_letters')}{_folder}\\4.docx"),
            f"{pc.get_option('path_for_save_letters')}{_folder}\\4.docx",
            _folder
        )

        doc = docx.Document(f"{pc.get_option('path_for_save_letters')}{_folder}\\Данные .docx")
        stud_name = doc.paragraphs[0].runs[-1].text
        try:
            os.rename(
                f"{pc.get_option('path_for_save_letters')}{_folder}\\Договор .docx",
                f"{pc.get_option('path_for_save_letters')}{_folder}\\Договор {stud_name}.docx"
            )
        except FileNotFoundError:
            print(f"Ошибка в вызове os.rename() в set_name_doc()")
        try:
            os.rename(
                f"{pc.get_option('path_for_save_letters')}{_folder}\\Заявление .docx",
                f"{pc.get_option('path_for_save_letters')}{_folder}\\Заявление {stud_name}.docx"
            )
        except FileNotFoundError:
            print(f"Ошибка в вызове os.rename() в set_name_doc()")
        try:
            os.rename(
                f"{pc.get_option('path_for_save_letters')}{_folder}\\СОГЛАСИЕ .docx",
                f"{pc.get_option('path_for_save_letters')}{_folder}\\СОГЛАСИЕ {stud_name}.docx"
            )
        except FileNotFoundError:
            print(f"Ошибка в вызове os.rename() в set_name_doc()")
        try:
            os.rename(
                f"{pc.get_option('path_for_save_letters')}{_folder}\\Данные .docx",
                f"{pc.get_option('path_for_save_letters')}{_folder}\\Данные {stud_name}.docx"
            )
        except FileNotFoundError:
            print(f"Ошибка в вызове os.rename() в set_name_doc()")
        return stud_name

    mail = MailConnect()
    email_list = mail.check_list_message().decode("utf-8")
    new_letters = mail.parse_new_messages(email_list)
    folders = []
    for letter in new_letters:
        folders.append(mail.take_files(letter))
    mail.close()
    for folder in folders:
        try:
            name_stud = set_name_doc(folder)
            os.rename(
                f"{pc.get_option('path_for_save_letters')}{folder}",
                f"{pc.get_option('path_for_save_letters')}{folder} {name_stud}"
            )
        except Exception:
            print(f"Ошибка в вызове os.rename() в check_new_message() в цикле 'for folder in folders:'")


def open_file(command):
    thread_list = []
    task = threading.Thread(target=OpenFile(), args=(command,))
    thread_list.append(task)
    task.deamon = True
    task.start()


class OpenFile:
    def __call__(self, command):
        open_file_function(command)


def open_file_function(command):
    call(f'{command}', shell=True)


def create_timetable_doc(_sub):
    path = os.getcwd() + r"/Документы/Расписания/"

    _db = ARMDataBase('arm_db.db')

    _sql = "SELECT sub_ttable, sub_name, id_prog FROM subjects WHERE id_sub=" + _sub
    timetable = _db.query(_sql)

    _sql = "SELECT group_name FROM groups WHERE id_prog=" + str(timetable[0][2])
    try:
        group_name = _db.query(_sql)[0][0]
    except IndexError:
        group_name = "Нет группы с этой программой"
    except Exception:
        group_name = "Ошибка загрузки группы"

    _sql = "SELECT id_student FROM subs_in_studs WHERE id_sub=" + _sub + " AND status='1'"
    students_q = _db.query(_sql)
    students = []

    filename = "Расписание " + group_name + " " + timetable[0][1] + " №000000.docx"
    desk_list_dir = os.listdir(path)
    indexes_list = []
    for doc_in_dir in desk_list_dir:
        start_index = doc_in_dir.find('№')
        if start_index:
            try:
                indexes_list.append(int(doc_in_dir[start_index + 1: start_index + 7]))
            except Exception:
                pass
    copy_index = 0
    while copy_index in indexes_list:
        copy_index += 1
        str_copy_index = str(copy_index)
        while len(str_copy_index) < 6:
            str_copy_index = "0" + str_copy_index
        filename = f"Расписание {group_name} {timetable[0][1]} №{str_copy_index}.docx"

    for i in range(len(students_q)):
        students.append([str(students_q[i][0])])
        _sql = "SELECT student_name FROM students WHERE id_student=" + str(students_q[i][0])
        students[i].append(_db.query(_sql)[0][0])

    _db.close()

    students.sort(key=lambda x: x[1])

    parse_timetable = []
    if timetable[0][0] is not None and timetable[0][0] != '':
        for date in timetable[0][0].split(","):
            parse_timetable.append(date)
        for i in range(len(parse_timetable)):
            parse_timetable[i] = parse_timetable[i].split("|")
            parse_timetable[i][0] = datetime.datetime.strptime(parse_timetable[i][0], "%d.%m.%Y")
            if parse_timetable[i][0].weekday() == 0:
                parse_timetable[i].append('Понедельник')
            elif parse_timetable[i][0].weekday() == 1:
                parse_timetable[i].append('Вторник')
            elif parse_timetable[i][0].weekday() == 2:
                parse_timetable[i].append('Среда')
            elif parse_timetable[i][0].weekday() == 3:
                parse_timetable[i].append('Четверг')
            elif parse_timetable[i][0].weekday() == 4:
                parse_timetable[i].append('Пятница')
            elif parse_timetable[i][0].weekday() == 5:
                parse_timetable[i].append('Суббота')
            elif parse_timetable[i][0].weekday() == 6:
                parse_timetable[i].append('Воскресенье')

            if parse_timetable[i][0].strftime("%m") == "01":
                parse_timetable[i].append('Января')
            elif parse_timetable[i][0].strftime("%m") == "02":
                parse_timetable[i].append('Февраля')
            elif parse_timetable[i][0].strftime("%m") == "03":
                parse_timetable[i].append('Марта')
            elif parse_timetable[i][0].strftime("%m") == "04":
                parse_timetable[i].append('Апреля')
            elif parse_timetable[i][0].strftime("%m") == "05":
                parse_timetable[i].append('Мая')
            elif parse_timetable[i][0].strftime("%m") == "06":
                parse_timetable[i].append('Июня')
            elif parse_timetable[i][0].strftime("%m") == "07":
                parse_timetable[i].append('Июля')
            elif parse_timetable[i][0].strftime("%m") == "08":
                parse_timetable[i].append('Августа')
            elif parse_timetable[i][0].strftime("%m") == "09":
                parse_timetable[i].append('Сентября')
            elif parse_timetable[i][0].strftime("%m") == "10":
                parse_timetable[i].append('Октября')
            elif parse_timetable[i][0].strftime("%m") == "11":
                parse_timetable[i].append('Ноября')
            elif parse_timetable[i][0].strftime("%m") == "12":
                parse_timetable[i].append('Декабря')

    doc = docx.Document()

    doc.sections[-1].orientation = docx.enum.section.WD_ORIENTATION.PORTRAIT
    doc.sections[-1].page_height = docx.shared.Cm(21)
    doc.sections[-1].page_width = docx.shared.Cm(29.7)
    doc.sections[-1].top_margin = docx.shared.Cm(1.3)
    doc.sections[-1].right_margin = docx.shared.Cm(1)
    doc.sections[-1].left_margin = docx.shared.Cm(1)
    doc.sections[-1].bottom_margin = docx.shared.Cm(1)

    doc.add_paragraph(group_name + " " + timetable[0][1])
    doc.paragraphs[0].runs[0].bold = True
    doc.paragraphs[0].runs[0].font.name = "Times New Roman"
    doc.paragraphs[0].runs[0].font.size = docx.shared.Pt(14)

    tabs_c = str(len(parse_timetable) / 22).split(".")

    if tabs_c[1] != "0":
        tabs_c = int(tabs_c[0]) + 1
    else:
        tabs_c = int(tabs_c[0])

    if len(parse_timetable) - (22 * (tabs_c - 1)) <= 10:
        len_date = 27
    else:
        len_date = 22

    tabs_c = str(len(parse_timetable) / len_date).split(".")

    if tabs_c[1] != "0":
        tabs_c = int(tabs_c[0]) + 1
    else:
        tabs_c = int(tabs_c[0])

    table_timetable_list = []
    j = 0
    for i in range(tabs_c):
        if tabs_c == 1:
            table_timetable = doc.add_table(rows=1 + len(students), cols=2 + len(parse_timetable), style='Table Grid')
            table_timetable_list.append(table_timetable)
        elif i + 1 != tabs_c:
            table_timetable = doc.add_table(rows=1 + len(students), cols=2 + len_date, style='Table Grid')
            table_timetable_list.append(table_timetable)
            if len(students) > 13:
                par = doc.add_paragraph('_')
                par.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

                par = doc.add_paragraph(group_name + " " + timetable[0][1])
                par.runs[0].bold = True
                par.runs[0].font.name = "Times New Roman"
                par.runs[0].font.size = docx.shared.Pt(14)
            elif len(students) < 13 and j % 2 == 1:
                par = doc.add_paragraph('_')
                par.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

                par = doc.add_paragraph(group_name + " " + timetable[0][1])
                par.runs[0].bold = True
                par.runs[0].font.name = "Times New Roman"
                par.runs[0].font.size = docx.shared.Pt(14)
            elif len(students) < 13 and j % 2 == 0:
                par = doc.add_paragraph('_')
                par.runs[0].font.size = docx.shared.Pt(1)
                par.paragraph_format.space_after = docx.shared.Pt(0)
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif len(students) == 13 and j % 2 == 0:
                par = doc.add_paragraph('_')
                par.runs[0].font.size = docx.shared.Pt(1)
                par.paragraph_format.space_after = docx.shared.Pt(0)
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif len(students) == 13 and j % 2 == 1:
                par = doc.add_paragraph('_')
                par.runs[0].font.size = docx.shared.Pt(1)
                par.paragraph_format.space_after = docx.shared.Pt(0)
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                par = doc.add_paragraph(group_name + " " + timetable[0][1])
                par.runs[0].bold = True
                par.runs[0].font.name = "Times New Roman"
                par.runs[0].font.size = docx.shared.Pt(14)
        else:
            table_timetable = doc.add_table(rows=1 + len(students),
                                            cols=2 + len(parse_timetable) - (len_date * (tabs_c - 1)),
                                            style='Table Grid')
            table_timetable_list.append(table_timetable)
        j += 1

    for i in range(len(table_timetable_list)):
        for row in range(len(table_timetable_list[i].rows)):
            for col in range(len(table_timetable_list[i].columns)):
                cell = table_timetable_list[i].cell(row, col)
                if row == 0 and col == 0:
                    cell.text = "№\nп\\п"
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(10)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
                    cell.width = docx.shared.Cm(0.89)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif row == 0 and col == 1:
                    cell.text = "ФИО"
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(10)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
                    cell.width = docx.shared.Cm(5.5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif row > 0 and col == 0:
                    cell.text = str(row) + "."
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(10)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
                    cell.width = docx.shared.Cm(0.89)
                    table_timetable_list[i].rows[row].height = docx.shared.Cm(0.4)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif row > 0 and col == 1:
                    cell.text = students[row - 1][1]
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(10)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
                    cell.width = docx.shared.Cm(5.8)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                elif row == 0 and col > 1:
                    cell.text = parse_timetable[0][0].strftime("%d")[1:] \
                        if parse_timetable[0][0].strftime("%d").startswith("0") \
                        else parse_timetable[0][0].strftime("%d")
                    cell.text += " " + parse_timetable[0][3]
                    parse_timetable.pop(0)
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(10)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
                    cell.paragraphs[0].paragraph_format.left_indent = docx.shared.Cm(0.2)
                    cell.paragraphs[0].paragraph_format.right_indent = docx.shared.Cm(0.2)
                    cell.width = docx.shared.Cm(0.81)
                    table_timetable_list[i].rows[row].height = docx.shared.Cm(2.5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_vertical_cell_direction(cell, "btLr")
                else:
                    cell.text = ""
                    cell.paragraphs[0].runs[0].font.name = "Times New Roman"
                    cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(10)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
                    cell.width = docx.shared.Cm(0.81)
                    table_timetable_list[i].rows[row].height = docx.shared.Cm(0.4)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    properties = doc.core_properties
    properties.author = "ЦПЮИ ХТИ"

    doc.save(path + filename)
    return path, filename


# Set font vertical direction in docx
def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


# Clearing edit list
def clear_list(children_list):
    for i in children_list:
        if i.objectName().startswith('clb_'):
            i.setObjectName('deleteLater_this_clb')
            i.deleteLater()
            i.setAttribute(55, 1)
            i.close()


# Clearing edit groupBox
def clear_group_box(group_box):
    for i in group_box:
        if i.objectName().startswith('el_') or i.objectName().startswith('grB_'):
            i.setAttribute(55, 1)
            i.close()


# Clearing widget
def clear_widget(widget):
    for i in widget:
        if i.objectName().startswith('widget_'):
            i.setObjectName('deleteLater_this_widget')
            i.deleteLater()
            i.setAttribute(55, 1)
            i.close()
            del i


# Func for print docs
def print_doc(filepath, filename):
    f = '"' + filepath + filename + '"'
    win32api.ShellExecute(0, "printto", f, '"%s"' % win32print.GetDefaultPrinter(), ".", 0)


# Func for warning
def set_doc_warning(war_name, war_text, war_icon=":/sfu_logo.ico"):
    _set_doc_warning = QMessageBox()
    _set_doc_warning.setWindowTitle(war_name)
    _set_doc_warning.setText(war_text)
    _set_doc_warning.setIcon(QMessageBox.Warning)
    icon = QtGui.QIcon()
    icon.addPixmap(QtGui.QPixmap(war_icon), QtGui.QIcon.Normal, QtGui.QIcon.Off)
    _set_doc_warning.setWindowIcon(icon)
    _set_doc_warning.exec_()


# Class for MailConnect
class MailConnect:
    def __init__(
            self,
            login=pc.get_option("login"),
            password=pc.get_option("password"),
            email_address=pc.get_option("mail")
    ):
        self.mail = imaplib.IMAP4_SSL(email_address)
        self.mail.login(login, password)

    def check_list_message(self):
        self.mail.select("inbox")
        result, data = self.mail.uid('search', None, "ALL")
        return data[0]

    def check_from(self, uid_message):
        latest_email_uid = uid_message
        result, data = self.mail.uid('fetch', latest_email_uid, '(RFC822)')
        raw_email = data[0][1]
        try:
            message_info = email.message_from_string(raw_email.decode("UTF-8"))
        except UnicodeDecodeError:
            message_info = email.message_from_string(raw_email.decode("cp1251"))
        except Exception:
            pass
        return email.utils.parseaddr(message_info['From'])[1]

    def parse_new_messages(self, _email_list):
        old_list = pc.get_option("letters").split()
        el = _email_list.split()
        new_messages = []
        for i in old_list:
            if i in el:
                el.remove(i)
        if el:
            for i in el:
                old_list.append(i)
            new_list = ""
            for i in old_list:
                new_list += i + " "
            new_list = new_list[:-1]
            pc.set_option("letters", new_list)
            for i in el:
                if self.check_from(i) == pc.get_option("sender"):
                    new_messages.append(i)
        return new_messages

    def take_files(self, uid_message):
        result, data = self.mail.uid('fetch', uid_message, '(RFC822)')
        raw_email = data[0][1]
        _email_message = email.message_from_bytes(raw_email)
        message_path = "000000"
        if not os.path.exists(f"{pc.get_option('path_for_save_letters')}"):
            os.mkdir(f"{pc.get_option('path_for_save_letters')}")
        folders = os.listdir(pc.get_option('path_for_save_letters'))
        for i in range(len(folders)):
            folders[i] = folders[i][:6]
        while message_path in folders:
            message_path = int(message_path)
            message_path += 1
            message_path = str(message_path)
            while len(message_path) < 6:
                message_path = "0" + message_path
        if not os.path.exists(f"{pc.get_option('path_for_save_letters')}{message_path}"):
            os.mkdir(f"{pc.get_option('path_for_save_letters')}{message_path}")
        if _email_message.is_multipart():
            cpui_mail_files = 0
            for part in _email_message.walk():
                filename = part.get_filename()
                if filename:
                    cpui_mail_files += 1
                    with open(
                            f"{pc.get_option('path_for_save_letters')}{message_path}\\{str(cpui_mail_files)}.docx",
                            'wb') as new_file:
                        new_file.write(part.get_payload(decode=True))
        return message_path

    def close(self, close_mode=False):
        self.mail.close()
        if not close_mode:
            self.mail.logout()


# Func for main window start
def main_win_start():
    app = QtWidgets.QApplication([])
    application = MainWindow()
    application.show()
    sys.exit(app.exec())


# Main func
def main():
    check_new_messages()
    main_win_start()


# Start application if her main
if __name__ == "__main__":
    main()
